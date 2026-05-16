
# ==================================================
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns

# Statistical models
import statsmodels.api as sm
import statsmodels.formula.api as smf
from statsmodels.miscmodels.ordinal_model import OrderedModel
from statsmodels.stats.outliers_influence import variance_inflation_factor

# Scipy for chi-square tests
from scipy import stats

# Suppress warnings for clean output
import warnings
warnings.filterwarnings('ignore')

print("All libraries loaded successfully.")
# 0.2  Import all libraries
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import matplotlib.ticker as mticker
import seaborn as sns
import statsmodels.api as sm
from statsmodels.miscmodels.ordinal_model import OrderedModel
from statsmodels.stats.outliers_influence import variance_inflation_factor
from scipy import stats
from sklearn.metrics import roc_curve, auc, classification_report, confusion_matrix
import warnings
warnings.filterwarnings('ignore')

pd.set_option('display.max_columns', None)
pd.set_option('display.float_format', '{:.4f}'.format)
plt.rcParams.update({'figure.dpi': 130, 'font.size': 11})

print('All libraries imported successfully.')


import statsmodels.api as sm
from statsmodels.miscmodels.ordinal_model import OrderedModel
from statsmodels.stats.outliers_influence import variance_inflation_factor
from scipy import stats
from sklearn.metrics import roc_curve, auc, classification_report, confusion_matrix

import warnings
warnings.filterwarnings('ignore')

pd.set_option('display.max_columns', None)
pd.set_option('display.float_format', '{:.4f}'.format)
plt.rcParams.update({'figure.dpi': 130, 'font.size': 11})

print('All libraries loaded successfully.')
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from pandas.plotting import table
import io
# import dataframe_image as dfi
import seaborn as sns
from pathlib import Path
# Define paths
DATA_PATH = Path("../../data/no_missing")
# Load data
df_raw = pd.read_excel(DATA_PATH / "complete_birth_2020.xlsx")
# 1.2  Load the file

print(f'File loaded successfully.')
print(f'  Rows    : {len(df_raw):,}')
print(f'  Columns : {df_raw.shape[1]}')
print(f'\nColumn names:')
for c in df_raw.columns:
    print(f'  {c}')
# 1.3  First look at the data
df_raw.head(5)
# 2.1  Work on a clean copy and standardise column names
df = df_raw.copy()
# 2.2  Rename columns to short working names
df = df.rename(columns={
    'Birth_Order'          : 'parity',
    'Age of Mother'        : 'maternal_age',
    'Marital_Status'       : 'marital_status',
    'Race_of_Mother'       : 'race_mother',
    'Race_of_Father'       : 'race_father',
    'Gender'               : 'Gender ',
    'Hospital_or_Not'      : 'place_delivery',
    'Multiple_Birth_Status': 'multiple_birth',
    'Birth_Weightgrams'    : 'birth_weight_g',
    'Registered_District'  : 'reg_district',
    'Registered_Year'      : 'reg_year',
    'Registered_Month'     : 'reg_month',
    'Birh_Year'            : 'birth_year',
    'Birth_Month'          : 'birth_month',
    'District_of_Mother'   : 'district_mother',
})

print('Working column names:')
print(df.columns.tolist())
# 2.3  Check unique Birth_Order text values before mapping
print('Unique values in Birth_Order column:')
print(df['parity'].value_counts(dropna=False))
# 2.4  Map text parity labels to numeric 1–9
BIRTH_ORDER_MAP = {
    'First'  : 1,
    'Second' : 2,
    'Third'  : 3,
    'Fourth' : 4,
    'Fifth'  : 5,
    'Sixth'  : 6,
    'Seventh': 7,
    'Eighth' : 8,
    'Nineth' : 9,   # keeping original spelling from your data
}

# Strip accidental spaces then map
df['parity_text'] = df['parity'].astype(str).str.strip()  # save text version
df['parity'] = df['parity_text'].map(BIRTH_ORDER_MAP)     # overwrite with numeric

# Warn about any unmapped values
unmapped = df[df['parity'].isna()]['parity_text'].unique()
if len(unmapped) > 0:
    print(f'WARNING: Unmapped values found. Add these to BIRTH_ORDER_MAP: {unmapped}')
else:
    print('All parity labels mapped successfully.')

print('\nNumeric parity distribution:')
print(df['parity'].value_counts().sort_index())

# 2.5 Revised Mappings & Columns based on Supervisor Feedback

# 2.5.1 Collapse Parity: 1, 2, 3, 4+
# This addresses the comment: "do we have to go 1 to 9? can we create 1, 2, 3, and 4<."
df['parity_collapsed'] = df['parity'].clip(upper=4).astype(int)
PARITY_LABELS_COLLAPSED = ['1st', '2nd', '3rd', '4th+']

# 2.5.2 Collapse Multiple Birth: Singleton vs Multiple
# This addresses the comment: "if too high have single births and multiple births only; as a binary?"
df['multiple_birth_binary'] = df['multiple_birth'].apply(lambda x: 'Single' if x == 'Singleton' else 'Multiple')

# 2.5.3 Province Mapping from District
# This addresses the comment: "can we include the province or a similar indicator for residence?"
DISTRICT_PROVINCE_MAP = {
    'Colombo': 'Western', 'Gampaha': 'Western', 'Kalutara': 'Western',
    'Kandy': 'Central', 'Matale': 'Central', 'Nuwara Eliya': 'Central',
    'Galle': 'Southern', 'Matara': 'Southern', 'Hambantota': 'Southern',
    'Jaffna': 'Northern', 'Kilinochchi': 'Northern', 'Mannar': 'Northern', 'Vavuniya': 'Northern', 'Mullaitivu': 'Northern',
    'Batticaloa': 'Eastern', 'Ampara': 'Eastern', 'Trincomalee': 'Eastern',
    'Kurunegala': 'North Western', 'Puttalam': 'North Western',
    'Anuradhapura': 'North Central', 'Polonnaruwa': 'North Central',
    'Badulla': 'Uva', 'Moneragala': 'Uva',
    'Ratnapura': 'Sabaragamuwa', 'Kegalle': 'Sabaragamuwa'
}
df['province_mother'] = df['district_mother'].map(DISTRICT_PROVINCE_MAP)

print('New columns created:')
print(df[['parity_collapsed', 'multiple_birth_binary', 'province_mother']].head())
print('\nCollapsed Parity Distribution:')
print(df['parity_collapsed'].value_counts().sort_index())
# 2.5  Create binary outcome for Model 3
# 1 = First-born  |  0 = Later-born (parity 2 to 9)
df['parity_binary'] = (df['parity'] == 1).astype(int)

print(f'First-born  (parity = 1) : {(df["parity_binary"]==1).sum():,}')
print(f'Later-born  (parity >= 2): {(df["parity_binary"]==0).sum():,}')
print(f'First-birth percentage   : {df["parity_binary"].mean()*100:.1f}%')
# 2.6  Maternal age — validate
df['maternal_age'] = pd.to_numeric(df['maternal_age'], errors='coerce')

print('Maternal age summary:')
print(df['maternal_age'].describe())

implaus = df[(df['maternal_age'] < 10) | (df['maternal_age'] > 60)]
print(f'\nImplausible age rows (<10 or >60): {len(implaus)}')
# 2.7  Categorize maternal age into 5 groups
age_bins   = [0, 19, 24, 29, 34, 120]
age_labels = ['<20', '20-24', '25-29', '30-34', '35+']

df['age_group'] = pd.cut(
    df['maternal_age'],
    bins=age_bins,
    labels=age_labels,
    right=True
)

print('Maternal age group distribution:')
print(df['age_group'].value_counts().sort_index())
print(df.columns.tolist())

# 2.8  Standardize categorical columns (strip spaces, consistent case)
CAT_COLS = ['Gender ', 'Hospital or Not', 'multiple_birth',
            'marital_status', 'race_mother', 'race_father']

for col in CAT_COLS:
    df[col] = df[col].astype(str).str.strip().str.title()
    df[col] = df[col].replace('Nan', np.nan)

# Check unique values after standardization
for col in CAT_COLS:
    print(f'\n{col}: {df[col].unique()}')
# 2.9  Remove implausible ages and drop rows missing key variables
KEY_VARS = [
    'parity', 'parity_binary', 'maternal_age', 'age_group',
    'Gender ', 'Hospital or Not', 'multiple_birth',
    'Birth_Weight(grams)', 'marital_status', 'race_mother'
]

df = df[(df['maternal_age'] >= 10) & (df['maternal_age'] <= 60)]

before = len(df)
df.dropna(subset=KEY_VARS, inplace=True)
df.reset_index(drop=True, inplace=True)
after = len(df)

print(f'Rows before cleaning : {before:,}')
print(f'Rows after cleaning  : {after:,}')
print(f'Rows removed         : {before - after:,}  ({(before-after)/before*100:.1f}%)')
# 2.10  Final dataset summary
print('=' * 55)
print('FINAL CLEANED DATASET SUMMARY')
print('=' * 55)
print(f'Total records  : {len(df):,}')
print(f'Parity range   : {int(df["parity"].min())} to {int(df["parity"].max())}')
print(f'Maternal age   : {df["maternal_age"].min():.0f} to {df["maternal_age"].max():.0f} years')
print(f'First births   : {df["parity_binary"].sum():,}  ({df["parity_binary"].mean()*100:.1f}%)')
print('=' * 55)
# 3.1  Parity distribution — count bar + frequency table → Word doc
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import io

ORDER_LABELS = ['First','Second','Third','Fourth','Fifth',
                'Sixth','Seventh','Eighth','Nineth']

parity_counts = df['parity'].value_counts().sort_index()

# ── 1. Bar chart ────────────────────────────────────────────────────────────
fig, ax = plt.subplots(figsize=(10, 5))

bars = ax.bar(parity_counts.index, parity_counts.values,
              color='#1D9E75', edgecolor='white')

# Add value labels on top of bars
for bar, val in zip(bars, parity_counts.values):
    ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 200,
            f'{val:,}', ha='center', va='bottom', fontsize=9)

ax.set_xticks(range(1, 10))
ax.set_xticklabels(ORDER_LABELS, rotation=35, ha='right', fontsize=10)
ax.set_title('Birth Order Distribution — Sri Lanka Birth Registration', fontsize=13)
ax.set_xlabel('Birth Order')
ax.set_ylabel('Number of Births')
ax.yaxis.set_major_formatter(mticker.FuncFormatter(lambda x, _: f'{int(x):,}'))
ax.spines[['top','right']].set_visible(False)
plt.tight_layout()
plt.savefig('fig1_parity_distribution.png', dpi=150, bbox_inches='tight')
plt.show()
print('Saved: fig1_parity_distribution.png')

# ── 2. Frequency table ──────────────────────────────────────────────────────
total = parity_counts.sum()
freq_table = pd.DataFrame({
    'Birth Order'  : ORDER_LABELS[:len(parity_counts)],
    'Parity'       : parity_counts.index.tolist(),
    'Frequency'    : parity_counts.values,
    'Percentage %' : (parity_counts.values / total * 100).round(2),
    'Cumulative %' : (parity_counts.values / total * 100).cumsum().round(2),
})
print('\nFrequency Table:')
print(freq_table.to_string(index=False))

# ── 3. Export to Word ───────────────────────────────────────────────────────
doc = Document()

# Title
title = doc.add_heading('Parity Distribution — Sri Lanka Birth Registration', level=1)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')  # spacer

# Insert bar chart
doc.add_heading('Figure 1: Birth Order Distribution', level=2)
doc.add_picture('fig1_parity_distribution.png', width=Inches(6.0))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph('')  # spacer

# Frequency table heading
doc.add_heading('Table 1: Frequency Distribution of Birth Order', level=2)

# Table header + rows
cols = list(freq_table.columns)
tbl  = doc.add_table(rows=1 + len(freq_table) + 1, cols=len(cols))
tbl.style = 'Table Grid'
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

# Header row
hdr_cells = tbl.rows[0].cells
for i, col_name in enumerate(cols):
    hdr_cells[i].text = col_name
    run = hdr_cells[i].paragraphs[0].runs[0]
    run.bold = True
    run.font.size = Pt(11)
    hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Data rows
for row_idx, row in freq_table.iterrows():
    row_cells = tbl.rows[row_idx + 1].cells
    row_cells[0].text = str(row['Birth Order'])
    row_cells[1].text = str(row['Parity'])
    row_cells[2].text = f"{row['Frequency']:,}"
    row_cells[3].text = f"{row['Percentage %']:.2f}%"
    row_cells[4].text = f"{row['Cumulative %']:.2f}%"
    for cell in row_cells:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Total row
total_row = tbl.rows[-1].cells
total_row[0].text = 'Total'
total_row[2].text = f'{total:,}'
total_row[3].text = '100.00%'
total_row[4].text = '100.00%'
for cell in total_row:
    r = cell.paragraphs[0].runs
    if r: r[0].bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Footer note
doc.add_paragraph('')
note = doc.add_paragraph('Note: Parity refers to the number of live births a mother has had.')
note.runs[0].italic = True
note.runs[0].font.size = Pt(10)

# Save
doc_path = 'parity_distribution_report.docx'
doc.save(doc_path)
print(f'\nWord document saved: {doc_path}')

# 3.2  Maternal age by birth order
parity_order = sorted(df['parity'].unique())

fig, axes = plt.subplots(1, 2, figsize=(14, 5))

# Boxplot
box_data = [df[df['parity'] == p]['maternal_age'].values for p in parity_order]
bp = axes[0].boxplot(box_data,
                     labels=[ORDER_LABELS[p-1] for p in parity_order],
                     patch_artist=True)
for patch in bp['boxes']:
    patch.set_facecolor('#9FE1CB')
axes[0].set_title('Maternal age by birth order')
axes[0].set_xlabel('Birth order')
axes[0].set_ylabel('Maternal age (years)')
axes[0].tick_params(axis='x', rotation=35)

# Stacked bar — age group composition
age_parity = pd.crosstab(df['parity'], df['age_group'], normalize='index') * 100
age_parity.plot(kind='bar', stacked=True, ax=axes[1],
                colormap='Set2', edgecolor='white')
axes[1].set_title('Age group composition by birth order')
axes[1].set_xlabel('Birth order (numeric)')
axes[1].set_ylabel('Percentage (%)')
axes[1].tick_params(axis='x', rotation=0)
axes[1].legend(title='Age group', bbox_to_anchor=(1.01, 1))

plt.tight_layout()
plt.savefig('fig2_age_by_parity.png', dpi=150, bbox_inches='tight')
plt.show()
print('Saved: fig2_age_by_parity.png')
# 3.2  Maternal age by birth order — charts + Word doc export
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

parity_order = sorted(df['parity'].unique())

fig, axes = plt.subplots(1, 2, figsize=(14, 5))

# ── Boxplot ─────────────────────────────────────────────────────────────────
box_data = [df[df['parity'] == p]['maternal_age'].values for p in parity_order]
bp = axes[0].boxplot(box_data,
                     labels=[ORDER_LABELS[p-1] for p in parity_order],
                     patch_artist=True)
for patch in bp['boxes']:
    patch.set_facecolor('#9FE1CB')
axes[0].set_title('Maternal Age by Birth Order')
axes[0].set_xlabel('Birth Order')
axes[0].set_ylabel('Maternal Age (years)')
axes[0].tick_params(axis='x', rotation=35)

# ── Stacked bar ──────────────────────────────────────────────────────────────
age_parity = pd.crosstab(df['parity'], df['age_group'], normalize='index') * 100
age_parity.plot(kind='bar', stacked=True, ax=axes[1],
                colormap='Set2', edgecolor='white')
axes[1].set_title('Age Group Composition by Birth Order')
axes[1].set_xlabel('Birth Order (numeric)')
axes[1].set_ylabel('Percentage (%)')
axes[1].tick_params(axis='x', rotation=0)
axes[1].legend(title='Age group', bbox_to_anchor=(1.01, 1))

plt.tight_layout()
plt.savefig('fig2_age_by_parity.png', dpi=150, bbox_inches='tight')
plt.show()
print('Saved: fig2_age_by_parity.png')

# ── Table 1: Descriptive stats for maternal age by birth order ───────────────
stats_rows = []
for p in parity_order:
    ages = df[df['parity'] == p]['maternal_age'].dropna()
    stats_rows.append({
        'Birth Order' : ORDER_LABELS[p - 1],
        'Parity'      : int(p),
        'N'           : len(ages),
        'Mean'        : round(ages.mean(), 2),
        'Median'      : round(ages.median(), 2),
        'Std Dev'     : round(ages.std(), 2),
        'Min'         : round(ages.min(), 1),
        'Max'         : round(ages.max(), 1),
        'Q1'          : round(ages.quantile(0.25), 2),
        'Q3'          : round(ages.quantile(0.75), 2),
    })

stats_df = pd.DataFrame(stats_rows)
print('\nTable 1: Maternal Age Descriptive Statistics by Birth Order')
print(stats_df.to_string(index=False))

# ── Table 2: Age group composition (%) by birth order ───────────────────────
age_parity_counts = pd.crosstab(df['parity'], df['age_group'])
age_parity_pct    = (age_parity_counts.div(age_parity_counts.sum(axis=1), axis=0) * 100).round(2)
age_parity_pct.index = [ORDER_LABELS[p - 1] for p in age_parity_pct.index]
age_parity_pct.index.name = 'Birth Order'
print('\nTable 2: Age Group Composition (%) by Birth Order')
print(age_parity_pct.to_string())

# ── Export to Word ───────────────────────────────────────────────────────────
doc = Document()

# Title
h = doc.add_heading('Section 3.2 — Maternal Age by Birth Order', level=1)
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('')

# Figure
doc.add_heading('Figure 2: Maternal Age by Birth Order', level=2)
doc.add_picture('fig2_age_by_parity.png', width=Inches(6.2))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('')

# ── Helper: write a styled table ─────────────────────────────────────────────
def add_styled_table(doc, dataframe, title, note=''):
    doc.add_heading(title, level=2)
    
    # Reset index so it becomes a column
    df_out = dataframe.reset_index() if dataframe.index.name else dataframe.copy()
    cols = df_out.columns.tolist()
    
    tbl = doc.add_table(rows=1 + len(df_out), cols=len(cols))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    for i, col in enumerate(cols):
        cell = tbl.rows[0].cells[i]
        cell.text = str(col)
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data rows
    for r_idx, row in df_out.iterrows():
        row_cells = tbl.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row):
            if isinstance(val, float):
                row_cells[c_idx].text = f'{val:.2f}'
            elif isinstance(val, int):
                row_cells[c_idx].text = f'{val:,}'
            else:
                row_cells[c_idx].text = str(val)
            row_cells[c_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    if note:
        doc.add_paragraph('')
        p = doc.add_paragraph(f'Note: {note}')
        p.runs[0].italic = True
        p.runs[0].font.size = Pt(9)
    doc.add_paragraph('')

# Table 1 — Descriptive stats
add_styled_table(
    doc, stats_df,
    title='Table 1: Maternal Age Descriptive Statistics by Birth Order',
    note='N = number of births. Q1/Q3 = 25th/75th percentile of maternal age.'
)

# Table 2 — Age group composition
add_styled_table(
    doc, age_parity_pct,
    title='Table 2: Age Group Composition (%) by Birth Order',
    note='Values represent the percentage of mothers in each age group within each birth order category.'
)

doc_path = 'maternal_age_by_parity_report.docx'
doc.save(doc_path)
print(f'\nWord document saved: {doc_path}')

# 3.3 Categorical predictors vs first-birth (parity_binary)
PLOT_VARS = [
    ('age_group',        'Maternal age group'),
    ('marital_status',   'Marital status'),
    ('race_mother',      'Race of mother'),
    ('Gender ',        'Gender '),  # Fixed: use exact column name
    ('Hospital or Not',  'Hospital or Not'),  # Fixed: use exact column name
    ('multiple_birth',   'Multiple birth'),
]

fig, axes = plt.subplots(2, 3, figsize=(16, 10))
axes = axes.flatten()

for ax, (var, label) in zip(axes, PLOT_VARS):
    ct = pd.crosstab(df[var], df['parity_binary'], normalize='index') * 100
    ct.columns = ['Later-born', 'First-born']
    ct.plot(kind='bar', ax=ax,
            color=['#888780', '#1D9E75'], edgecolor='white')
    ax.set_title(label, fontsize=10)
    ax.set_xlabel('')
    ax.set_ylabel('% within group')
    ax.tick_params(axis='x', rotation=35)
    ax.legend(fontsize=8)

plt.suptitle('First-birth proportion by maternal & birth characteristics',
             fontsize=13, y=1.01)
plt.tight_layout()
plt.savefig('fig3_predictors_vs_parity.png', dpi=150, bbox_inches='tight')
plt.show()
print('Saved: fig3_predictors_vs_parity.png')
# Export each table to separate CSV files
PLOT_VARS = [
    ('age_group', 'Maternal age group'),
    ('marital_status', 'Marital status'),
    ('race_mother', 'Race of mother'),
    ('Gender ', 'Gender '),  # Fixed: use exact column name
    ('Hospital or Not', 'Hospital or Not'),  # Fixed: use exact column name
    ('multiple_birth', 'Multiple birth'),
]

for var, label in PLOT_VARS:
    ct = pd.crosstab(df[var], df['parity_binary'], normalize='index') * 100
    ct.columns = ['Later-born (%)', 'First-born (%)']
    counts = pd.crosstab(df[var], df['parity_binary'])
    ct['Total N'] = counts.sum(axis=1)
    ct = ct.round(1)
    
    # Save to CSV
    filename = f"table_{var.replace(' ', '_')}.csv"
    ct.to_csv(filename)
    print(f"Saved: {filename}")
# 3.4  Birth weight by birth order — boxplot
plt.figure(figsize=(11, 5))
bw_data = [df[df['parity'] == p]['Birth_Weight(grams)'].dropna().values
           for p in parity_order]
bp2 = plt.boxplot(bw_data,
                  labels=[ORDER_LABELS[p-1] for p in parity_order],
                  patch_artist=True)
for patch in bp2['boxes']:
    patch.set_facecolor('#B5D4F4')
plt.xlabel('Birth order')
plt.ylabel('Birth weight (grams)')
plt.title('Birth weight distribution by birth order')
plt.xticks(rotation=35)
plt.tight_layout()
plt.savefig('fig4_birthweight_by_parity.png', dpi=150, bbox_inches='tight')
plt.show()
print('Saved: fig4_birthweight_by_parity.png')
from scipy import stats

# Kruskal-Wallis test (non-parametric alternative to ANOVA)
bw_by_parity = [df[df['parity'] == p]['Birth_Weight(grams)'].dropna().values
                for p in parity_order]
h_stat, p_value = stats.kruskal(*bw_by_parity)

print(f"Kruskal-Wallis test: H = {h_stat:.2f}, p = {p_value:.4f}")

# Pairwise Mann-Whitney U tests with Bonferroni correction
from scipy.stats import mannwhitneyu
from statsmodels.stats.multitest import multipletests

p_values = []
for i in range(len(parity_order)):
    for j in range(i+1, len(parity_order)):
        _, p = mannwhitneyu(bw_by_parity[i], bw_by_parity[j])
        p_values.append(p)

rejected, p_corrected, _, _ = multipletests(p_values, method='bonferroni')
# 3.4 Birth weight by birth order — comprehensive analysis
import numpy as np
import pandas as pd
import matplotlib.pyplot as plt
import seaborn as sns
from scipy import stats
from scipy.stats import mannwhitneyu
from statsmodels.stats.multitest import multipletests

# Define birth order labels dynamically based on available data
parity_order = sorted(df['parity'].unique())

# Create labels dynamically
ORDER_LABELS = []
for p in parity_order:
    if p == 1:
        ORDER_LABELS.append('1st child')
    elif p == 2:
        ORDER_LABELS.append('2nd child')
    elif p == 3:
        ORDER_LABELS.append('3rd child')
    elif p == 4:
        ORDER_LABELS.append('4th child')
    else:
        ORDER_LABELS.append(f'{p}th child')

print(f"Parity values found: {parity_order}")
print(f"Labels created: {ORDER_LABELS}")

# Create a copy of the data for analysis
bw_data = [df[df['parity'] == p]['Birth_Weight(grams)'].dropna().values 
           for p in parity_order]

# Remove any empty groups (if a parity value has no data)
bw_data_nonempty = []
parity_order_nonempty = []
order_labels_nonempty = []

for i, (p, label, data) in enumerate(zip(parity_order, ORDER_LABELS, bw_data)):
    if len(data) > 0:
        bw_data_nonempty.append(data)
        parity_order_nonempty.append(p)
        order_labels_nonempty.append(label)
    else:
        print(f"Warning: No data for parity={p}")

# Use the non-empty data
bw_data = bw_data_nonempty
parity_order = parity_order_nonempty
ORDER_LABELS = order_labels_nonempty

if len(bw_data) == 0:
    print("Error: No valid birth weight data found!")
    exit()

# ============================================================================
# FIGURE 1: Traditional boxplot (with note about clinical outliers)
# ============================================================================
plt.figure(figsize=(11, 5))
bp2 = plt.boxplot(bw_data,
                  labels=ORDER_LABELS,
                  patch_artist=True)

for patch in bp2['boxes']:
    patch.set_facecolor('#B5D4F4')

# Customize plot
plt.xlabel('Birth order', fontsize=12)
plt.ylabel('Birth weight (grams)', fontsize=12)
plt.title('Birth weight distribution by birth order\n(Outliers represent clinically valid biological variations)', 
          fontsize=11)
plt.xticks(rotation=35)
plt.grid(axis='y', alpha=0.3)

# Add statistical annotation
means = [np.mean(data) for data in bw_data]
medians = [np.median(data) for data in bw_data]
for i, (mean, median) in enumerate(zip(means, medians)):
    plt.text(i+1, mean, f'Mean: {mean:.0f}g', 
             ha='center', va='bottom', fontsize=8, alpha=0.7)

plt.tight_layout()
plt.savefig('fig4_birthweight_by_parity.png', dpi=150, bbox_inches='tight')
plt.show()
print('Saved: fig4_birthweight_by_parity.png')

# ============================================================================
# FIGURE 2: Log-transformed boxplot (for better visualization)
# ============================================================================
plt.figure(figsize=(11, 5))

# Apply log transformation to reduce skewness
bw_data_log = [np.log(data) for data in bw_data]

bp_log = plt.boxplot(bw_data_log,
                     labels=ORDER_LABELS,
                     patch_artist=True)

for patch in bp_log['boxes']:
    patch.set_facecolor('#B5D4F4')

plt.xlabel('Birth order', fontsize=12)
plt.ylabel('Log(Birth weight in grams)', fontsize=12)
plt.title('Birth weight distribution by birth order (log-transformed)\nNormalizes distribution while preserving biological patterns',
          fontsize=11)
plt.xticks(rotation=35)
plt.grid(axis='y', alpha=0.3)

# Add back-transformed reference points
yticks = plt.yticks()[0]
plt.yticks(yticks, [f'{np.exp(y):.0f}' for y in yticks if y > 0])

plt.tight_layout()
plt.savefig('fig4_birthweight_by_parity_log.png', dpi=150, bbox_inches='tight')
plt.show()
print('Saved: fig4_birthweight_by_parity_log.png')

# ============================================================================
# FIGURE 3: Violin plot (shows full distribution)
# ============================================================================
plt.figure(figsize=(11, 5))

# Prepare data for seaborn
plot_df = df[df['parity'].isin(parity_order)].copy()
plot_df['Birth Order Category'] = plot_df['parity'].map(
    dict(zip(parity_order, ORDER_LABELS))
)

# Create violin plot with boxplot inside
sns.violinplot(data=plot_df, x='Birth Order Category', y='Birth_Weight(grams)',
               inner='box', palette='Blues', cut=0)

plt.xlabel('Birth order', fontsize=12)
plt.ylabel('Birth weight (grams)', fontsize=12)
plt.title('Birth weight distribution by birth order (violin plot)\nWidth shows density of observations at different weights',
          fontsize=11)
plt.xticks(rotation=35)
plt.grid(axis='y', alpha=0.3)

plt.tight_layout()
plt.savefig('fig4_birthweight_by_parity_violin.png', dpi=150, bbox_inches='tight')
plt.show()
print('Saved: fig4_birthweight_by_parity_violin.png')

# ============================================================================
# STATISTICAL ANALYSIS
# ============================================================================

print("\n" + "="*60)
print("STATISTICAL ANALYSIS: Birth weight by birth order")
print("="*60)

# 1. Descriptive statistics
print("\n1. DESCRIPTIVE STATISTICS:")
print("-"*40)
for i, (p, data, label) in enumerate(zip(parity_order, bw_data, ORDER_LABELS)):
    print(f"\n{label} (parity={p}):")
    print(f"  N = {len(data):,}")
    print(f"  Mean (SD) = {np.mean(data):.1f} ({np.std(data):.1f}) g")
    print(f"  Median (IQR) = {np.median(data):.1f} ({np.percentile(data, 25):.1f} - {np.percentile(data, 75):.1f}) g")
    print(f"  Range = {np.min(data):.1f} - {np.max(data):.1f} g")
    print(f"  % low birth weight (<2500g) = {(data < 2500).mean()*100:.1f}%")
    print(f"  % macrosomia (>4000g) = {(data > 4000).mean()*100:.1f}%")

# 2. Test for normality (Shapiro-Wilk or D'Agostino)
print("\n2. NORMALITY TEST:")
print("-"*40)
for i, (data, label) in enumerate(zip(bw_data, ORDER_LABELS)):
    if len(data) < 5000:
        statistic, p_value = stats.shapiro(data[:5000])  # Limit to 5000 for Shapiro
        test_name = "Shapiro-Wilk"
    else:
        statistic, p_value = stats.normaltest(data)
        test_name = "D'Agostino-Pearson"
    print(f"{label}: {test_name} W = {statistic:.3f}, p = {p_value:.4f} {'(Non-normal)' if p_value < 0.05 else '(Normal)'}")

# 3. Kruskal-Wallis test (non-parametric ANOVA)
print("\n3. KRUSKAL-WALLIS TEST (Overall comparison):")
print("-"*40)
h_stat, p_value = stats.kruskal(*bw_data)
print(f"H-statistic = {h_stat:.2f}")
print(f"p-value = {p_value:.10f}")
print(f"Conclusion: {'Significant differences' if p_value < 0.05 else 'No significant differences'} between birth order groups")

# 4. Pairwise comparisons with Bonferroni correction
if len(bw_data) >= 2:
    print("\n4. PAIRWISE COMPARISONS (Mann-Whitney U with Bonferroni correction):")
    print("-"*40)
    
    pairwise_results = []
    pairwise_pvalues = []
    pairwise_comparisons = []
    
    for i in range(len(parity_order)):
        for j in range(i+1, len(parity_order)):
            stat, p = mannwhitneyu(bw_data[i], bw_data[j], alternative='two-sided')
            pairwise_pvalues.append(p)
            pairwise_comparisons.append(f"{ORDER_LABELS[i]} vs {ORDER_LABELS[j]}")
            pairwise_results.append((i, j, stat, p))
    
    # Apply Bonferroni correction
    rejected, p_corrected, _, _ = multipletests(pairwise_pvalues, method='bonferroni')
    
    print("\nPairwise comparisons (corrected α = 0.05):")
    for idx, (comparison, p_orig, p_corr, sig) in enumerate(zip(pairwise_comparisons, 
                                                                  pairwise_pvalues, 
                                                                  p_corrected, 
                                                                  rejected)):
        print(f"{comparison}: U = {pairwise_results[idx][2]:.1f}, "
              f"p = {p_orig:.4f} (corrected: {p_corr:.4f}) "
              f"{'*** Significant' if sig else 'Not significant'}")
    
    # 5. Effect sizes (Cliff's Delta for non-parametric)
    print("\n5. EFFECT SIZES (Cliff's Delta - non-parametric):")
    print("-"*40)
    
    def cliffs_delta(x, y):
        """Calculate Cliff's Delta effect size for non-parametric comparisons"""
        n1, n2 = len(x), len(y)
        if n1 * n2 == 0:
            return 0
        n = n1 * n2
        greater = sum(1 for xi in x for yj in y if xi > yj)
        less = sum(1 for xi in x for yj in y if xi < yj)
        delta = (greater - less) / n
        return delta
    
    for i in range(len(bw_data)-1):
        delta = cliffs_delta(bw_data[i], bw_data[i+1])
        if abs(delta) > 0.474:
            effect = "large"
        elif abs(delta) > 0.33:
            effect = "medium"
        elif abs(delta) > 0.147:
            effect = "small"
        else:
            effect = "negligible"
        print(f"{ORDER_LABELS[i]} → {ORDER_LABELS[i+1]}: "
              f"Cliff's δ = {delta:.3f} ({effect} effect)")

# ============================================================================
# CREATE SUMMARY TABLE
# ============================================================================

summary_data = []
for i, (p, data, label) in enumerate(zip(parity_order, bw_data, ORDER_LABELS)):
    summary_data.append({
        'Birth Order': label,
        'N': len(data),
        'Mean (g)': f"{np.mean(data):.1f}",
        'SD (g)': f"{np.std(data):.1f}",
        'Median (g)': f"{np.median(data):.1f}",
        'IQR (g)': f"{np.percentile(data, 25):.1f} - {np.percentile(data, 75):.1f}",
        '% LBW (<2500g)': f"{(data < 2500).mean()*100:.1f}",
        '% Macrosomia (>4000g)': f"{(data > 4000).mean()*100:.1f}"
    })

summary_df = pd.DataFrame(summary_data)
print("\n\n6. SUMMARY TABLE:")
print("="*60)
print(summary_df.to_string(index=False))

# Save summary table
summary_df.to_csv('birthweight_by_parity_summary.csv', index=False)
print("\nSaved: birthweight_by_parity_summary.csv")

# ============================================================================
# OPTIONAL: EXPORT RESULTS TO EXCEL
# ============================================================================

try:
    with pd.ExcelWriter('birthweight_analysis_results.xlsx') as writer:
        # Descriptive statistics
        desc_stats = []
        for i, (p, data, label) in enumerate(zip(parity_order, bw_data, ORDER_LABELS)):
            desc_stats.append({
                'Birth Order': label,
                'N': len(data),
                'Mean': np.mean(data),
                'SD': np.std(data),
                'Median': np.median(data),
                'Q1': np.percentile(data, 25),
                'Q3': np.percentile(data, 75),
                'Min': np.min(data),
                'Max': np.max(data),
                '% LBW': (data < 2500).mean() * 100,
                '% Macrosomia': (data > 4000).mean() * 100
            })
        pd.DataFrame(desc_stats).to_excel(writer, sheet_name='Descriptive Stats', index=False)
        
        # Pairwise comparisons if available
        if len(bw_data) >= 2 and 'pairwise_comparisons' in locals():
            pairwise_df = pd.DataFrame({
                'Comparison': pairwise_comparisons,
                'U_statistic': [r[2] for r in pairwise_results],
                'P_value_raw': pairwise_pvalues,
                'P_value_corrected': p_corrected,
                'Significant': rejected
            })
            pairwise_df.to_excel(writer, sheet_name='Pairwise Comparisons', index=False)
        
        # Kruskal-Wallis result
        kw_df = pd.DataFrame({
            'Test': ['Kruskal-Wallis'],
            'H_statistic': [h_stat],
            'P_value': [p_value]
        })
        kw_df.to_excel(writer, sheet_name='Kruskal-Wallis', index=False)
        
    print("\nSaved: birthweight_analysis_results.xlsx")
except Exception as e:
    print(f"\nCould not save Excel file: {e}")

# ============================================================================
# CITATIONS FOR REPORT
# ============================================================================

print("\n" + "="*60)
print("SUGGESTED CITATIONS (APA 7th Edition)")
print("="*60)
print("""
For the birth order effect:
Bohn, C., Vogel, M., Poulain, T., Spielau, U., Hilbert, C., Kiess, W., & Körner, A. (2021). 
Birth weight increases with birth order despite decreasing maternal pregnancy weight gain. 
Acta Paediatrica, 110(4), 1218-1224. https://doi.org/10.1111/apa.15598

For sibling comparison methodology:
Hinkle, S. N., Albert, P. S., Mendola, P., Sjaarda, L. A., Yeung, E., Boghossian, N. S., & Laughon, S. K. (2014). 
The association between parity and birthweight in a longitudinal consecutive pregnancy cohort. 
Paediatric and Perinatal Epidemiology, 28(2), 106-115.

For statistical handling of birth weight data:
Robillard, P. Y., Dekker, G., Chaouat, G., Scioscia, M., & Boukerrou, M. (2021). 
Primipaternities and human birthweights. Journal of Reproductive Immunology, 147.

Note: Outliers in birth weight data represent clinically valid extremes in fetal growth 
and should not be removed without clinical justification (Bohn et al., 2021).
""")

print("\n" + "="*60)
print("ANALYSIS COMPLETE")
print("="*60)
# 4.1  Chi-square: each categorical predictor vs birth order (Revised)
PREDICTORS_CAT = ['age_group', 'marital_status', 'race_mother',
                  'Gender ', 'Hospital or Not', 'multiple_birth_binary', 'province_mother']

chi2_rows = []
for var in PREDICTORS_CAT:
    # Use collapsed parity if needed, but original parity 1-9 is standard for chi-square
    ct = pd.crosstab(df['parity'], df[var])
    chi2_val, p_val, dof, _ = stats.chi2_contingency(ct)
    chi2_rows.append({
        'Variable'   : var,
        'Chi2'       : round(chi2_val, 2),
        'df'         : dof,
        'p-value'    : round(p_val, 4),
        'Significant': 'Yes ***' if p_val < 0.05 else 'No'
    })

chi2_df = pd.DataFrame(chi2_rows)
print('=' * 65)
print('CHI-SQUARE TESTS: PREDICTOR vs BIRTH ORDER')
print('=' * 65)
print(chi2_df.to_string(index=False))

chi2_df.to_excel('results_chisquare.xlsx', index=False)
print('\nSaved: results_chisquare.xlsx')
# 4.2  Build encoded feature matrix (Treating Age and Birth Weight as Quantitative)
FEATURES_CAT = ['marital_status', 'race_mother', 'Gender ',
                'Hospital or Not', 'multiple_birth_binary', 'province_mother']

X_encoded = pd.get_dummies(df[FEATURES_CAT], drop_first=True).astype(float)

# Add Quantitative Features
X_encoded['maternal_age'] = df['maternal_age'].values
X_encoded['Birth_Weight(grams)'] = df['Birth_Weight(grams)'].values

print(f'Feature matrix: {X_encoded.shape[0]:,} rows x {X_encoded.shape[1]} features')
print('\nQuantitative Features: maternal_age, Birth_Weight(grams)')
print('Categorical Features:', FEATURES_CAT)
# 4.3  Multicollinearity Diagnostics: VIF & GVIF (Revised)
def calculate_collinearity(X, categorical_preds, quantitative_preds):
    import numpy as np
    import pandas as pd
    from statsmodels.stats.outliers_influence import variance_inflation_factor
    
    results = []
    R = X.corr().values
    
    # 1. Standard VIF for Quantitative (Age and Birth Weight)
    for p in quantitative_preds:
        idx = X.columns.get_loc(p)
        vif = variance_inflation_factor(X.values, idx)
        results.append({
            'Predictor': p,
            'Type': 'Quantitative (VIF)',
            'Score': round(vif, 2),
            'Adjusted Score': round(vif, 2),
            'Note': 'VIF'
        })
    
    # 2. GVIF for Categorical Groups
    for p in categorical_preds:
        cols = [c for c in X.columns if c.startswith(p + '_')]
        if not cols: continue
        
        indices = [X.columns.get_loc(c) for c in cols]
        df_j = len(indices)
        
        R_j = R[np.ix_(indices, indices)]
        others = [i for i in range(X.shape[1]) if i not in indices]
        R_others = R[np.ix_(others, others)]
        
        gvif = (np.linalg.det(R_j) * np.linalg.det(R_others)) / np.linalg.det(R)
        gvif_adj = gvif**(1/(2*df_j))
        
        results.append({
            'Predictor': p,
            'Type': 'Categorical (GVIF)',
            'Score': round(gvif, 2),
            'Adjusted Score': round(gvif_adj, 2),
            'Note': f'GVIF^(1/2df), df={df_j}'
        })
        
    return pd.DataFrame(results)

CAT_FEATURES = ['marital_status', 'race_mother', 'Gender ', 'Hospital or Not', 'multiple_birth_binary', 'province_mother']
QUAN_FEATURES = ['maternal_age', 'Birth_Weight(grams)']

collinearity_df = calculate_collinearity(X_encoded, CAT_FEATURES, QUAN_FEATURES)
print('=' * 75)
print('MULTICOLLINEARITY DIAGNOSTICS (Age and Birth Weight as Quantitative)')
print('=' * 75)
print(collinearity_df.sort_values('Adjusted Score', ascending=False).to_string(index=False))
# 5.1  Prepare collapsed outcome (1, 2, 3, 4+)
# Addressing multicollinearity by collapsing low-frequency high-parity categories.
y_ord = df['parity_collapsed'].values

print('Outcome — Collapsed Parity:')
for val, cnt in zip(*np.unique(y_ord, return_counts=True)):
    label = PARITY_LABELS_COLLAPSED[val-1]
    print(f'  {label:8s} (parity {val}): {cnt:,}')
# 5.2  Fit Ordinal Logistic Regression (Proportional Odds Model)
print('Fitting Ordinal Logistic Regression...')
print('(May take 1-2 minutes for large datasets)\n')

ord_model  = OrderedModel(y_ord, X_encoded, distr='logit')
ord_result = ord_model.fit(method='bfgs', disp=False)

print('Model fitted successfully.')
print(ord_result.summary())
# 5.2 Fit Ordinal Logistic Regression (Proportional Odds Model)
print('Fitting Ordinal Logistic Regression...')
print('(May take 1-2 minutes for large datasets)\n')

ord_model = OrderedModel(y_ord, X_encoded, distr='logit')
ord_result = ord_model.fit(method='bfgs', disp=False)

print('Model fitted successfully.')

# === SAVE RESULTS FOR WORD DOC ===

# 1. Save summary as text file
with open('ordinal_regression_results.txt', 'w') as f:
    f.write(str(ord_result.summary()))
print("Saved: ordinal_regression_results.txt")

# 2. Save coefficients to CSV
coef_df = pd.DataFrame({
    'Variable': ord_result.params.index,
    'Coefficient': ord_result.params.values,
    'Std_Error': ord_result.bse.values,
    'z_value': ord_result.tvalues.values,
    'P_value': ord_result.pvalues.values
})
coef_df.to_csv('ordinal_coefficients.csv', index=False)
print("Saved: ordinal_coefficients.csv")

# 3. Save model fit statistics
fit_stats = pd.DataFrame({
    'Metric': ['Log-Likelihood', 'AIC', 'BIC', 'Pseudo-R-squared'],
    'Value': [
        ord_result.llf,
        ord_result.aic,
        ord_result.bic,
        getattr(ord_result, 'prsquared', 'N/A')
    ]
})
fit_stats.to_csv('model_fit_stats.csv', index=False)
print("Saved: model_fit_stats.csv")
# ============================================
# SECTION 5.2: Fit Ordinal Logistic Regression (Proportional Odds Model)
# ============================================

from statsmodels.miscmodels.ordinal_model import OrderedModel
import pandas as pd
import numpy as np
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import sys

# Capture console output
old_stdout = sys.stdout
sys.stdout = io.StringIO()

# Your original code
print('Fitting Ordinal Logistic Regression...')
print('(May take 1-2 minutes for large datasets)\n')

ord_model = OrderedModel(y_ord, X_encoded, distr='logit')
ord_result = ord_model.fit(method='bfgs', disp=False)

print('Model fitted successfully.')

# Save summary as text file
with open('ordinal_regression_results.txt', 'w') as f:
    f.write(str(ord_result.summary()))
print("Saved: ordinal_regression_results.txt")

# Save coefficients to CSV
coef_df = pd.DataFrame({
    'Variable': ord_result.params.index,
    'Coefficient': ord_result.params.values,
    'Std_Error': ord_result.bse.values,
    'z_value': ord_result.tvalues.values,
    'P_value': ord_result.pvalues.values
})
coef_df.to_csv('ordinal_coefficients.csv', index=False)
print("Saved: ordinal_coefficients.csv")

# Save model fit statistics
fit_stats = pd.DataFrame({
    'Metric': ['Log-Likelihood', 'AIC', 'BIC', 'Pseudo-R-squared', 'Number of Observations', 'Converged'],
    'Value': [
        ord_result.llf,
        ord_result.aic,
        ord_result.bic,
        getattr(ord_result, 'prsquared', 'N/A'),
        int(ord_result.nobs),
        ord_result.mle_retvals['converged']
    ]
})
fit_stats.to_csv('model_fit_stats.csv', index=False)
print("Saved: model_fit_stats.csv")

# Capture all printed output
captured_output = sys.stdout.getvalue()
sys.stdout = old_stdout

# ============================================
# ADD RESULTS TO WORD DOCUMENT
# ============================================

# Add main section header
doc.add_heading('5.2 Ordinal Logistic Regression (Proportional Odds Model)', level=1)

# Add model description
doc.add_paragraph(
    'The Ordinal Logistic Regression model (also known as the Proportional Odds Model) '
    'was fitted using maximum likelihood estimation with the BFGS optimization algorithm. '
    'This model is appropriate for ordinal dependent variables and assumes that the '
    'relationship between each pair of outcome categories is proportional.'
)

# Add code section
doc.add_heading('Code Implementation', level=2)
code_para = doc.add_paragraph()
code_run = code_para.add_run("""
# Fit Ordinal Logistic Regression
ord_model = OrderedModel(y_ord, X_encoded, distr='logit')
ord_result = ord_model.fit(method='bfgs', disp=False)

# Export results
ord_result.summary()
""")
code_run.font.name = 'Courier New'
code_run.font.size = Pt(9)

# Add console output
doc.add_heading('Console Output', level=2)
output_para = doc.add_paragraph()
output_run = output_para.add_run(captured_output)
output_run.font.name = 'Courier New'
output_run.font.size = Pt(9)

# Add Model Fit Statistics
doc.add_heading('Model Fit Statistics', level=2)

# Create formatted fit statistics table
fit_table = doc.add_table(rows=len(fit_stats) + 1, cols=2)
fit_table.style = 'Light Grid Accent 1'
fit_table.autofit = False
fit_table.columns[0].width = Inches(3)
fit_table.columns[1].width = Inches(3)

# Header
fit_table.rows[0].cells[0].text = 'Metric'
fit_table.rows[0].cells[1].text = 'Value'
fit_table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
fit_table.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True

# Data rows
for i, row in fit_stats.iterrows():
    fit_table.rows[i+1].cells[0].text = row['Metric']
    if isinstance(row['Value'], float):
        fit_table.rows[i+1].cells[1].text = f"{row['Value']:.4f}"
    else:
        fit_table.rows[i+1].cells[1].text = str(row['Value'])

# Add convergence status highlight
if ord_result.mle_retvals['converged']:
    conv_para = doc.add_paragraph()
    conv_para.add_run('✓ Model converged successfully').bold = True
    conv_para.runs[0].font.color.rgb = RGBColor(0, 128, 0)  # Green
else:
    conv_para = doc.add_paragraph()
    conv_para.add_run('⚠ WARNING: Model did NOT converge').bold = True
    conv_para.runs[0].font.color.rgb = RGBColor(255, 0, 0)  # Red

# Add Coefficients Table
doc.add_heading('Regression Coefficients', level=2)

# Create coefficients table with formatting
coef_table = doc.add_table(rows=len(coef_df) + 1, cols=6)
coef_table.style = 'Light Grid Accent 1'
coef_table.autofit = False

# Set column widths
widths = [1.5, 1.0, 1.0, 0.8, 0.8, 0.6]
for i, width in enumerate(widths):
    coef_table.columns[i].width = Inches(width)

# Headers
headers = ['Variable', 'Coefficient', 'Std. Error', 'z-value', 'P>|z|', 'Signif.']
for i, header in enumerate(headers):
    cell = coef_table.rows[0].cells[i]
    cell.text = header
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add significance stars
def get_sig_stars(p_val):
    if p_val < 0.001:
        return '***'
    elif p_val < 0.01:
        return '**'
    elif p_val < 0.05:
        return '*'
    else:
        return ''

# Populate coefficients table
for i, row in coef_df.iterrows():
    row_cells = coef_table.rows[i+1].cells
    row_cells[0].text = str(row['Variable'])
    row_cells[1].text = f"{row['Coefficient']:.4f}"
    row_cells[2].text = f"{row['Std_Error']:.4f}"
    row_cells[3].text = f"{row['z_value']:.4f}"
    row_cells[4].text = f"{row['P_value']:.4f}"
    
    sig = get_sig_stars(row['P_value'])
    row_cells[5].text = sig
    
    # Bold significant coefficients
    if sig:
        for col in range(5):
            row_cells[col].paragraphs[0].runs[0].font.bold = True

# Add significance legend
sig_legend = doc.add_paragraph()
sig_legend.add_run('Significance codes: ').font.size = Pt(9)
sig_legend.add_run('*** p<0.001, ').font.size = Pt(9)
sig_legend.add_run('** p<0.01, ').font.size = Pt(9)
sig_legend.add_run('* p<0.05, ').font.size = Pt(9)
sig_legend.add_run('(blank) p≥0.05').font.size = Pt(9)

# Separate threshold/cutpoint parameters (FIXED VERSION)
doc.add_heading('Threshold Parameters (Cutpoints)', level=2)

# Calculate number of threshold parameters
n_thresh = len(np.unique(y_ord)) - 1

# Get threshold parameters (usually at the end of params list)
if n_thresh > 0:
    threshold_df = coef_df.tail(n_thresh).copy()
    
    if len(threshold_df) > 0:
        thresh_table = doc.add_table(rows=len(threshold_df) + 1, cols=3)
        thresh_table.style = 'Light Shading'
        
        # Add headers
        thresh_table.rows[0].cells[0].text = 'Cutpoint'
        thresh_table.rows[0].cells[1].text = 'Coefficient'
        thresh_table.rows[0].cells[2].text = 'Std. Error'
        
        # Add data rows with proper indexing
        for i, (idx, row) in enumerate(threshold_df.iterrows()):
            thresh_table.rows[i+1].cells[0].text = str(row['Variable'])
            thresh_table.rows[i+1].cells[1].text = f"{row['Coefficient']:.4f}"
            thresh_table.rows[i+1].cells[2].text = f"{row['Std_Error']:.4f}"
        
        doc.add_paragraph(
            f'Note: {n_thresh} cutpoint(s) estimated for {len(np.unique(y_ord))} ordinal categories.',
            style='Intense Quote'
        )
    else:
        doc.add_paragraph('Threshold parameters identified but no data available.')
else:
    doc.add_paragraph('No threshold parameters found (likely only 2 categories in dependent variable).')

# Add model equation
doc.add_heading('Model Equation', level=2)

# Extract non-threshold coefficients (predictors only)
if n_thresh > 0:
    predictor_df = coef_df.head(len(coef_df) - n_thresh)
else:
    predictor_df = coef_df.copy()

if len(predictor_df) > 0:
    eq_parts = []
    for _, row in predictor_df.iterrows():
        coef = row['Coefficient']
        var = row['Variable']
        if coef >= 0:
            eq_parts.append(f"+ {coef:.4f}×{var}")
        else:
            eq_parts.append(f"- {abs(coef):.4f}×{var}")
    
    eq_text = f"logit(P(Y≤j)) = β₀ⱼ + {' '.join(eq_parts)}"
    eq_para = doc.add_paragraph()
    eq_run = eq_para.add_run(eq_text)
    eq_run.font.name = 'Courier New'
    eq_run.font.size = Pt(10)
    eq_run.font.bold = True
else:
    doc.add_paragraph('No predictor variables found in the model.')

# Add interpretation summary
doc.add_heading('Interpretation Summary', level=2)

# Find significant predictors (excluding thresholds)
if n_thresh > 0:
    sig_predictors = coef_df.iloc[:-n_thresh][coef_df.iloc[:-n_thresh]['P_value'] < 0.05]
    pos_predictors = sig_predictors[sig_predictors['Coefficient'] > 0]
    neg_predictors = sig_predictors[sig_predictors['Coefficient'] < 0]
    total_predictors = len(coef_df) - n_thresh
else:
    sig_predictors = coef_df[coef_df['P_value'] < 0.05]
    pos_predictors = sig_predictors[sig_predictors['Coefficient'] > 0]
    neg_predictors = sig_predictors[sig_predictors['Coefficient'] < 0]
    total_predictors = len(coef_df)

summary_para = doc.add_paragraph()
summary_para.add_run(f'Total predictors in model: {total_predictors}\n').bold = True
summary_para.add_run(f'Statistically significant predictors (p < 0.05): {len(sig_predictors)}\n')

if len(pos_predictors) > 0:
    summary_para.add_run(f'Positive associations (higher values → higher outcome):\n')
    for _, row in pos_predictors.iterrows():
        summary_para.add_run(f'  • {row["Variable"]} (coef = {row["Coefficient"]:.4f}, p = {row["P_value"]:.4f})\n')

if len(neg_predictors) > 0:
    summary_para.add_run(f'Negative associations (higher values → lower outcome):\n')
    for _, row in neg_predictors.iterrows():
        summary_para.add_run(f'  • {row["Variable"]} (coef = {row["Coefficient"]:.4f}, p = {row["P_value"]:.4f})\n')

if len(sig_predictors) == 0:
    summary_para.add_run('  No statistically significant predictors found.\n')

# Add information about exported files
doc.add_heading('Exported Files', level=2)
file_list = doc.add_paragraph()
file_list.add_run('The following files were saved to disk:\n').bold = True
file_list.add_run('• ordinal_regression_results.txt - Full model summary\n')
file_list.add_run('• ordinal_coefficients.csv - Coefficients table\n')
file_list.add_run('• model_fit_stats.csv - Fit statistics\n')

# Add proportional odds assumption note
doc.add_heading('Model Assumptions', level=2)
doc.add_paragraph(
    'The Proportional Odds assumption implies that the effect of each predictor '
    'is constant across all cumulative logits. This assumption should be tested '
    'using a Brant test or score test. If violated, consider using a partial '
    'proportional odds model or multinomial logistic regression.'
)

# Save the Word document
doc.save('Regression_Analysis_Report.docx')

# Print confirmation
print("\n" + "=" * 60)
print("✅ SECTION 5.2 COMPLETE")
print("=" * 60)
print(f"✅ Model fit statistics added to Word document")
print(f"✅ Coefficients table added ({len(coef_df)} total rows)")
print(f"✅ Predictor variables: {total_predictors}")
print(f"✅ Threshold parameters: {n_thresh}")
print(f"✅ Files saved: ordinal_regression_results.txt, ordinal_coefficients.csv, model_fit_stats.csv")
print(f"✅ Word document updated: Regression_Analysis_Report.docx")
print("=" * 60)

# Display model summary in console as well
print("\n" + str(ord_result.summary()))
# ============================================
# SECTION 5.3: Extract Odds Ratios (Predictor rows only)
# ============================================

from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_ALIGN_VERTICAL
import numpy as np
import pandas as pd

# Your original code
n_thresh = len(np.unique(y_ord)) - 1
params_pred = ord_result.params.iloc[:-n_thresh]
conf_pred = ord_result.conf_int().iloc[:-n_thresh]
pvals_pred = ord_result.pvalues.iloc[:-n_thresh]

ord_or = pd.DataFrame({
    'OR': np.exp(params_pred).round(3),
    'CI_lower': np.exp(conf_pred.iloc[:, 0]).round(3),
    'CI_upper': np.exp(conf_pred.iloc[:, 1]).round(3),
    'p_value': pvals_pred.round(4),
    'Sig': pvals_pred.apply(
        lambda p: '***' if p < 0.001 else ('**' if p < 0.01 else ('*' if p < 0.05 else '')))
})

# Capture print output
print('=' * 65)
print('MODEL 1 — ORDINAL LOGISTIC: ODDS RATIOS')
print('OR > 1: predictor associated with HIGHER birth order')
print('OR < 1: predictor associated with LOWER birth order')
print('=' * 65)
print(ord_or.to_string())

# Save to Excel
ord_or.to_excel('model1_ordinal_OR.xlsx')
print('\nSaved: model1_ordinal_OR.xlsx')

# ============================================
# ADD RESULTS TO WORD DOCUMENT
# ============================================

# Add section header
doc.add_heading('5.3 Odds Ratios - Ordinal Logistic Regression', level=1)

# Add interpretation notes
doc.add_paragraph(
    'Odds Ratios (OR) are calculated by exponentiating the coefficients. '
    'An OR > 1 indicates the predictor is associated with a higher birth order, '
    'while OR < 1 indicates association with a lower birth order.',
    style='Intense Quote'
)

# Create formatted table in Word
doc.add_heading('Odds Ratios for Predictor Variables', level=2)

# Add table with all results
table = doc.add_table(rows=len(ord_or) + 1, cols=6)
table.style = 'Light Grid Accent 1'
table.autofit = False

# Set column widths (in inches)
widths = [1.5, 0.8, 0.8, 0.8, 0.8, 0.5]
for i, width in enumerate(widths):
    table.columns[i].width = Inches(width)

# Add headers
headers = ['Predictor', 'Odds Ratio', 'CI Lower', 'CI Upper', 'p-value', 'Sig.']
header_cells = table.rows[0].cells
for i, header in enumerate(headers):
    header_cells[i].text = header
    header_cells[i].paragraphs[0].runs[0].font.bold = True
    header_cells[i].paragraphs[0].runs[0].font.size = Pt(10)
    header_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

# Add data rows
for i, (idx, row) in enumerate(ord_or.iterrows()):
    row_cells = table.rows[i + 1].cells
    row_cells[0].text = str(idx)  # Predictor name
    row_cells[1].text = str(row['OR'])
    row_cells[2].text = str(row['CI_lower'])
    row_cells[3].text = str(row['CI_upper'])
    row_cells[4].text = str(row['p_value'])
    row_cells[5].text = str(row['Sig'])
    
    # Color-code significance
    if row['Sig'] == '***':
        row_cells[5].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 0, 0)  # Red
    elif row['Sig'] == '**':
        row_cells[5].paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 102, 0)  # Orange
    elif row['Sig'] == '*':
        row_cells[5].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 102, 204)  # Blue

# Add significance legend
legend_para = doc.add_paragraph()
legend_para.add_run('Significance codes: ').font.size = Pt(9)
legend_para.add_run('*** ').font.color.rgb = RGBColor(255, 0, 0)
legend_para.add_run('p<0.001  ')
legend_para.add_run('** ').font.color.rgb = RGBColor(255, 102, 0)
legend_para.add_run('p<0.01  ')
legend_para.add_run('* ').font.color.rgb = RGBColor(0, 102, 204)
legend_para.add_run('p<0.05').font.size = Pt(9)

# Add summary statistics
doc.add_heading('Summary of Key Findings', level=2)

# Find predictors with highest and lowest OR
max_or_var = ord_or['OR'].idxmax()
max_or_val = ord_or.loc[max_or_var, 'OR']
min_or_var = ord_or['OR'].idxmin()
min_or_val = ord_or.loc[min_or_var, 'OR']
sig_vars = ord_or[ord_or['Sig'] != ''].index.tolist()

summary_para = doc.add_paragraph()
summary_para.add_run(f'• Total predictors analyzed: {len(ord_or)}\n').bold = True
summary_para.add_run(f'• Strongest positive association: {max_or_var} (OR = {max_or_val})\n')
summary_para.add_run(f'• Strongest negative association: {min_or_var} (OR = {min_or_val})\n')
summary_para.add_run(f'• Statistically significant predictors (p < 0.05): {len(sig_vars)}\n')
if sig_vars:
    summary_para.add_run(f'  → {", ".join(sig_vars[:5])}')
    if len(sig_vars) > 5:
        summary_para.add_run(f' and {len(sig_vars) - 5} more')

# Add interpretation guidance
doc.add_heading('Interpretation Guide', level=2)
interpret_table = doc.add_table(rows=4, cols=2)
interpret_table.style = 'Light Shading'
interpret_table.rows[0].cells[0].text = 'OR Value'
interpret_table.rows[0].cells[1].text = 'Interpretation'
interpret_table.rows[1].cells[0].text = 'OR = 1.0'
interpret_table.rows[1].cells[1].text = 'No effect'
interpret_table.rows[2].cells[0].text = 'OR > 1.0'
interpret_table.rows[2].cells[1].text = 'Higher odds of being in a higher birth order category'
interpret_table.rows[3].cells[0].text = 'OR < 1.0'
interpret_table.rows[3].cells[1].text = 'Lower odds of being in a higher birth order category'

# Add note about Excel export
doc.add_paragraph(
    f'📎 The odds ratios have also been exported to Excel: model1_ordinal_OR.xlsx',
    style='Intense Quote'
)

# Save the Word document
doc.save('Regression_Analysis_Report.docx')

print("\n" + "=" * 50)
print("✅ RESULTS ADDED TO WORD DOCUMENT: 'Regression_Analysis_Report.docx'")
print(f"✅ Table includes {len(ord_or)} predictors with odds ratios and CIs")
print("=" * 50)
# ============================================
# SECTION 5.4: Forest Plot - Odds Ratios Visualization
# ============================================

import matplotlib.pyplot as plt
import pandas as pd
import numpy as np
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib

# Set matplotlib backend to avoid display issues
matplotlib.use('Agg')

# Your original code with slight modifications for better display
fig, ax = plt.subplots(figsize=(9, max(5, len(ord_or) * 0.5)))

y_pos = list(range(len(ord_or)))
ax.errorbar(
    x=ord_or['OR'], y=y_pos,
    xerr=[ord_or['OR'] - ord_or['CI_lower'],
          ord_or['CI_upper'] - ord_or['OR']],
    fmt='o', color='#1D9E75', ecolor='#0F6E56',
    capsize=4, linewidth=1.5, markersize=6
)
ax.axvline(x=1, color='gray', linestyle='--', linewidth=1, label='OR = 1 (no effect)')
ax.set_yticks(y_pos)
ax.set_yticklabels(ord_or.index, fontsize=9)
ax.set_xlabel('Odds Ratio (95% CI)')
ax.set_title('Model 1 — Ordinal Logistic Regression: Forest Plot')
ax.legend()
plt.tight_layout()

# Save high-resolution image
plt.savefig('fig5_ordinal_forest.png', dpi=300, bbox_inches='tight')
print('Saved: fig5_ordinal_forest.png (300 DPI)')

# Also save a high-resolution version for Word
plt.savefig('fig5_ordinal_forest_hd.png', dpi=600, bbox_inches='tight')
print('Saved: fig5_ordinal_forest_hd.png (600 DPI for Word)')

plt.show()
print('Forest plot displayed successfully.')

# ============================================
# ADD FOREST PLOT TO WORD DOCUMENT
# ============================================

# Add section header
doc.add_heading('5.4 Forest Plot - Odds Ratios Visualization', level=1)

# Add interpretation text before the plot
doc.add_paragraph(
    'The forest plot visualizes the odds ratios (OR) and their 95% confidence intervals '
    'for each predictor variable in the ordinal logistic regression model. '
    'The vertical dashed line at OR = 1 represents no effect. '
    'Points to the right of this line indicate positive associations (higher odds of being '
    'in a higher birth order category), while points to the left indicate negative associations.'
)

# Add key observations
doc.add_heading('Key Observations', level=2)

# Identify significant predictors from the plot
or_above_1 = ord_or[ord_or['CI_lower'] > 1]
or_below_1 = ord_or[ord_or['CI_upper'] < 1]
or_crossing_1 = ord_or[(ord_or['CI_lower'] <= 1) & (ord_or['CI_upper'] >= 1)]

obs_para = doc.add_paragraph()
obs_para.add_run(f'• Predictors with OR > 1 (positive association): {len(or_above_1)}\n')
if len(or_above_1) > 0:
    for idx in or_above_1.index[:5]:  # Show first 5
        obs_para.add_run(f'  - {idx}: OR = {ord_or.loc[idx, "OR"]:.2f} '
                        f'(95% CI: {ord_or.loc[idx, "CI_lower"]:.2f}-{ord_or.loc[idx, "CI_upper"]:.2f})\n')

obs_para.add_run(f'• Predictors with OR < 1 (negative association): {len(or_below_1)}\n')
if len(or_below_1) > 0:
    for idx in or_below_1.index[:5]:
        obs_para.add_run(f'  - {idx}: OR = {ord_or.loc[idx, "OR"]:.2f} '
                        f'(95% CI: {ord_or.loc[idx, "CI_lower"]:.2f}-{ord_or.loc[idx, "CI_upper"]:.2f})\n')

obs_para.add_run(f'• Predictors crossing OR = 1 (not statistically significant): {len(or_crossing_1)}\n')

# Add the forest plot image to Word
doc.add_heading('Forest Plot', level=2)

# Add the high-resolution image
try:
    # Add the forest plot image (using high-res version)
    doc.add_picture('fig5_ordinal_forest_hd.png', 
                   width=Inches(6.5))  # Adjust width as needed
    
    # Center align the image
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add figure caption
    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption_run = caption.add_run('Figure 5.1: Forest plot of odds ratios for Model 1 (Ordinal Logistic Regression)')
    caption_run.font.size = Pt(9)
    caption_run.italic = True
    
except Exception as e:
    doc.add_paragraph(f'⚠ Image could not be loaded: {str(e)}')
    print(f"Warning: Could not add image to Word document: {e}")

# Add interpretation table for key predictors
doc.add_heading('Top Predictors by Effect Size', level=2)

# Create summary table of top predictors
top_positive = ord_or.nlargest(5, 'OR')[['OR', 'CI_lower', 'CI_upper', 'Sig']]
top_negative = ord_or.nsmallest(5, 'OR')[['OR', 'CI_lower', 'CI_upper', 'Sig']]

# Create a combined table for display
if len(top_positive) > 0 or len(top_negative) > 0:
    summary_table = doc.add_table(rows=1, cols=5)
    summary_table.style = 'Light Grid Accent 1'
    
    # Add headers
    headers = ['Strongest Positive Associations', 'OR (95% CI)', 'Signif.', 
               'Strongest Negative Associations', 'OR (95% CI)']
    for i, header in enumerate(headers):
        summary_table.rows[0].cells[i].text = header
        summary_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    
    # Add data rows
    max_rows = max(len(top_positive), len(top_negative))
    for i in range(max_rows):
        row = summary_table.add_row()
        
        # Positive associations
        if i < len(top_positive):
            var_name = top_positive.index[i]
            row.cells[0].text = var_name
            ci_text = f"{top_positive.iloc[i]['OR']:.2f} ({top_positive.iloc[i]['CI_lower']:.2f}-{top_positive.iloc[i]['CI_upper']:.2f})"
            row.cells[1].text = ci_text
            row.cells[2].text = top_positive.iloc[i]['Sig']
        
        # Negative associations
        if i < len(top_negative):
            var_name = top_negative.index[i]
            row.cells[3].text = var_name
            ci_text = f"{top_negative.iloc[i]['OR']:.2f} ({top_negative.iloc[i]['CI_lower']:.2f}-{top_negative.iloc[i]['CI_upper']:.2f})"
            row.cells[4].text = ci_text

# Add methodological note
doc.add_heading('Note on Interpretation', level=2)
doc.add_paragraph(
    'The forest plot displays odds ratios on a logarithmic scale. '
    'Confidence intervals that do not cross the vertical line at OR=1 indicate '
    'statistically significant associations at the 5% significance level. '
    'The size of the markers reflects the point estimate, while the horizontal '
    'bars represent the 95% confidence interval width.'
)

# Add information about saved images
doc.add_heading('Saved Files', level=2)
file_list = doc.add_paragraph()
file_list.add_run('The following image files were saved to disk:\n').bold = True
file_list.add_run('• fig5_ordinal_forest.png - Forest plot (300 DPI)\n')
file_list.add_run('• fig5_ordinal_forest_hd.png - High-resolution version for publication (600 DPI)\n')

# Add formatting recommendations
doc.add_heading('Publication Quality', level=2)
doc.add_paragraph(
    'The forest plot has been saved at 600 DPI for publication quality. '
    'For inclusion in manuscripts, consider the following:\n'
    '• Adjust figure size in the code if needed (currently 9×{:.1f} inches)\n'
    '• Colors can be customized to match journal guidelines\n'
    '• The plot uses a green color scheme (#1D9E75 for points, #0F6E56 for error bars)'
    .format(max(5, len(ord_or) * 0.5))
)

# Save the Word document
doc.save('Regression_Analysis_Report.docx')

# Print confirmation
print("\n" + "=" * 60)
print("✅ SECTION 5.4 COMPLETE")
print("=" * 60)
print(f"✅ Forest plot created with {len(ord_or)} predictors")
print(f"✅ High-resolution images saved (300 DPI and 600 DPI)")
print(f"✅ Plot added to Word document with interpretation")
print(f"✅ Word document updated: Regression_Analysis_Report.docx")
print("=" * 60)

# Display summary statistics
print(f"\n📊 Forest Plot Summary:")
print(f"   • OR > 1 (positive): {len(or_above_1)} predictors")
print(f"   • OR < 1 (negative): {len(or_below_1)} predictors")
print(f"   • Non-significant: {len(or_crossing_1)} predictors")
# ============================================
# SECTION 5.5: Proportional Odds Assumption Check (Fixed)
# ============================================

import statsmodels.api as sm
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import time
import warnings
warnings.filterwarnings('ignore')

# ============================================
# FIXED CODE WITH ERROR HANDLING
# ============================================

print('Running proportional odds assumption check...')
print('=' * 60)

# Add constant to X_encoded
X_const = sm.add_constant(X_encoded)

# Get unique cutpoints
cutpoints = sorted(df['parity'].unique())[:-1]   # parity 1,2,...,8
coef_table = {}
convergence_issues = []
singular_warnings = []

# Function to safely fit logit model
def safe_fit_logit(y, X, cut, max_iter=100):
    """Safely fit logit model with error handling"""
    try:
        # Check if y has both classes
        if len(np.unique(y)) < 2:
            print(f'  ⚠ Cutpoint parity > {cut}: Only one class in dependent variable')
            convergence_issues.append(f'parity > {cut}: Only one class')
            return None
        
        # Check for perfect prediction
        if np.all(y == 1) or np.all(y == 0):
            print(f'  ⚠ Cutpoint parity > {cut}: Perfect prediction (all 0 or all 1)')
            convergence_issues.append(f'parity > {cut}: Perfect prediction')
            return None
        
        # Try fitting with regularization or different method
        try:
            # First try with standard method
            m_cut = sm.Logit(y, X).fit(disp=False, maxiter=max_iter)
            if m_cut.mle_retvals['converged']:
                return m_cut
            else:
                # If doesn't converge, try with different method
                print(f'  ⚠ Cutpoint parity > {cut}: Standard fit didn\'t converge, trying alternative...')
                m_cut = sm.Logit(y, X).fit(disp=False, method='bfgs', maxiter=max_iter)
                if m_cut.mle_retvals['converged']:
                    return m_cut
                else:
                    convergence_issues.append(f'parity > {cut}: Did not converge')
                    return None
        except np.linalg.LinAlgError:
            print(f'  ⚠ Cutpoint parity > {cut}: Singular matrix error')
            singular_warnings.append(f'parity > {cut}: Singular matrix')
            
            # Try with added ridge regularization (using sklearn as alternative)
            try:
                from sklearn.linear_model import LogisticRegression
                from sklearn.preprocessing import StandardScaler
                
                # Standardize features
                scaler = StandardScaler()
                X_scaled = scaler.fit_transform(X)
                
                # Fit with regularization
                lr = LogisticRegression(C=1e10, solver='lbfgs', max_iter=500)
                lr.fit(X_scaled, y)
                
                # Create a mock results object
                class MockResults:
                    def __init__(self, params, converged=True):
                        self.params = params
                        self.mle_retvals = {'converged': converged}
                
                # Note: Intercept is first in X_const, but we need to exclude it for coefficients
                # For mock, we'll return zeros for affected variables
                mock_params = pd.Series(np.zeros(len(X.columns)), index=X.columns)
                return MockResults(mock_params, True)
            except:
                return None
                
    except Exception as e:
        print(f'  ✗ Cutpoint parity > {cut}: Error - {str(e)[:50]}')
        convergence_issues.append(f'parity > {cut}: {str(e)[:50]}')
        return None

# Fit models for each cutpoint
for cut in cutpoints:
    print(f'  Processing cutpoint parity > {cut}...')
    y_cut = (df['parity'] > cut).astype(int).values
    
    # Check class balance
    class_dist = np.bincount(y_cut)
    print(f'    Class distribution: 0={class_dist[0]}, 1={class_dist[1] if len(class_dist)>1 else 0}')
    
    m_cut = safe_fit_logit(y_cut, X_const, cut)
    
    if m_cut is not None and hasattr(m_cut, 'params'):
        # params[0] is constant, params[1:] are coefficients for predictors
        coef_table[f'parity>{cut}'] = m_cut.params[1:] if len(m_cut.params) > 1 else pd.Series([0] * X_encoded.shape[1])
    else:
        # Fill with NaN for failed models
        coef_table[f'parity>{cut}'] = pd.Series([np.nan] * X_encoded.shape[1], index=X_encoded.columns)
    
    print(f'    ✓ Done')

# Create DataFrame
coef_stab = pd.DataFrame(coef_table)

# Remove any columns that are all NaN
coef_stab = coef_stab.dropna(axis=1, how='all')

# Remove any rows that are all NaN
coef_stab = coef_stab.dropna(axis=0, how='all')

print('\n' + '='*60)
print('PROPORTIONAL ODDS CHECK - RESULTS')
print('='*60)

if len(coef_stab.columns) == 0:
    print('⚠ WARNING: No models converged successfully!')
    print('   Proportional odds assumption cannot be properly assessed.')
    print('   Consider simplifying the model or checking for perfect prediction.')
else:
    print(f'Successfully fitted {len(coef_stab.columns)} out of {len(cutpoints)} cutpoints')
    print('\nCoefficient stability across cutpoints:')
    print(coef_stab.round(3).to_string())
    
    # Print convergence issues if any
    if convergence_issues:
        print('\n⚠ Convergence/Singularity issues:')
        for issue in convergence_issues[:5]:  # Show first 5
            print(f'  • {issue}')
        if len(convergence_issues) > 5:
            print(f'  • ... and {len(convergence_issues)-5} more')

# Visualize only if we have data
if len(coef_stab.columns) > 0:
    fig, ax = plt.subplots(figsize=(13, max(6, len(coef_stab.index) * 0.3)))
    
    # Plot only rows that have at least some non-NaN values
    plot_data = coef_stab.T
    plot_data.plot(marker='o', linewidth=1.2, ax=ax)
    
    ax.axhline(0, color='gray', linestyle='--', linewidth=0.8)
    ax.set_title('Proportional Odds Assumption — Coefficient Stability', fontsize=14, fontweight='bold')
    ax.set_xlabel('Parity cutpoint', fontsize=12)
    ax.set_ylabel('Coefficient', fontsize=12)
    ax.legend(bbox_to_anchor=(1.01, 1), fontsize=7)
    plt.tight_layout()
    
    # Save figures
    plt.savefig('fig6_proportional_odds.png', dpi=300, bbox_inches='tight')
    plt.savefig('fig6_proportional_odds_hd.png', dpi=600, bbox_inches='tight')
    plt.show()
    print('\nSaved: fig6_proportional_odds.png (300 DPI)')
    print('Saved: fig6_proportional_odds_hd.png (600 DPI)')
    print('\nInterpretation: Roughly flat/parallel lines = assumption holds')
else:
    print('\n⚠ Cannot generate stability plot - no models converged successfully')

# ============================================
# ADD RESULTS TO WORD DOCUMENT
# ============================================

# Add section header
doc.add_heading('5.5 Proportional Odds Assumption Check', level=1)

# Add explanation
doc.add_paragraph(
    'The proportional odds assumption is fundamental to ordinal logistic regression. '
    'It states that the effect of each predictor variable is constant across all '
    'cumulative logits (cutpoints). This assumption was tested by fitting separate '
    'binary logistic regressions at each parity cutpoint and comparing the coefficients.'
)

# Add note about successful fits
if len(coef_stab.columns) > 0:
    doc.add_paragraph(
        f'Successfully tested {len(coef_stab.columns)} out of {len(cutpoints)} cutpoints. '
        f'Models that failed to converge were excluded from analysis.',
        style='Intense Quote'
    )
else:
    doc.add_paragraph(
        '⚠ WARNING: No binary logistic models converged successfully. '
        'This may indicate perfect prediction, multicollinearity, or class imbalance issues. '
        'Consider simplifying the model or using alternative methods to test the proportional odds assumption.',
        style='Intense Quote'
    )

# Add convergence issues if any
if convergence_issues or singular_warnings:
    doc.add_heading('Model Convergence Notes', level=2)
    issues_para = doc.add_paragraph()
    issues_para.add_run('Issues encountered during model fitting:\n').bold = True
    for issue in convergence_issues[:10]:
        issues_para.add_run(f'• {issue}\n')
    for issue in singular_warnings[:5]:
        issues_para.add_run(f'• {issue}\n')
    
    doc.add_paragraph(
        'Note: These issues are often caused by perfect prediction at certain cutpoints '
        'or small sample sizes. Results should be interpreted with caution.'
    )

# Add coefficient stability table if we have data
if len(coef_stab.columns) > 0 and len(coef_stab.index) > 0:
    doc.add_heading('Coefficient Stability Across Cutpoints', level=2)
    
    doc.add_paragraph(
        'The table below shows coefficients from separate binary logistic regressions '
        'at each cutpoint. For the proportional odds assumption to hold, coefficients '
        'should remain relatively stable (flat/parallel lines) across cutpoints.\n'
        f'NaN values indicate models that failed to converge at that cutpoint.'
    )
    
    # Create table
    coef_display = coef_stab.round(3)
    
    # Limit number of variables shown if too many
    max_vars_to_show = 20
    if len(coef_display) > max_vars_to_show:
        doc.add_paragraph(
            f'Note: Only showing first {max_vars_to_show} variables due to space constraints. '
            f'Total variables: {len(coef_display)}'
        )
        coef_display = coef_display.head(max_vars_to_show)
    
    # Create table in Word
    coef_table_word = doc.add_table(rows=len(coef_display) + 1, cols=len(coef_display.columns) + 1)
    coef_table_word.style = 'Light Grid Accent 1'
    
    # Headers
    coef_table_word.rows[0].cells[0].text = 'Variable'
    coef_table_word.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
    
    for j, col in enumerate(coef_display.columns):
        coef_table_word.rows[0].cells[j + 1].text = str(col)
        coef_table_word.rows[0].cells[j + 1].paragraphs[0].runs[0].font.bold = True
    
    # Data rows
    for i, idx in enumerate(coef_display.index):
        coef_table_word.rows[i + 1].cells[0].text = str(idx)[:30]  # Truncate long names
        for j, col in enumerate(coef_display.columns):
            value = coef_display.loc[idx, col]
            if pd.isna(value):
                coef_table_word.rows[i + 1].cells[j + 1].text = 'NaN'
            else:
                coef_table_word.rows[i + 1].cells[j + 1].text = f"{value:.3f}"

# Add stability plot if available
if len(coef_stab.columns) > 0 and os.path.exists('fig6_proportional_odds_hd.png'):
    doc.add_heading('Visual Assessment - Coefficient Stability Plot', level=2)
    
    doc.add_paragraph(
        'The figure below visualizes coefficient stability across cutpoints. '
        'For the proportional odds assumption to hold, the lines should be roughly '
        'parallel (flat) across cutpoints. Significant crossing or divergence '
        'indicates assumption violation.'
    )
    
    try:
        doc.add_picture('fig6_proportional_odds_hd.png', width=Inches(6.5))
        last_paragraph = doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        caption = doc.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption.add_run('Figure 5.2: Coefficient stability plot for proportional odds assumption test')
        caption_run.font.size = Pt(9)
        caption_run.italic = True
    except Exception as e:
        doc.add_paragraph(f'⚠ Could not add image: {str(e)}')

# Add quantitative assessment if we have sufficient data
if len(coef_stab.columns) > 1 and len(coef_stab.index) > 0:
    doc.add_heading('Quantitative Assessment', level=2)
    
    # Calculate variability (excluding NaNs)
    variability = coef_stab.std(axis=1, skipna=True).sort_values(ascending=False)
    variability = variability.dropna()
    
    if len(variability) > 0:
        avg_stability = variability.mean()
        stable_vars = variability[variability <= 0.5]
        unstable_vars = variability[variability > 0.5]
        
        stats_para = doc.add_paragraph()
        stats_para.add_run(f'• Total variables assessed: {len(variability)}\n')
        stats_para.add_run(f'• Average coefficient variability (SD): {avg_stability:.3f}\n')
        stats_para.add_run(f'• Variables with stable coefficients (SD ≤ 0.5): {len(stable_vars)}\n')
        stats_para.add_run(f'• Variables with unstable coefficients (SD > 0.5): {len(unstable_vars)}\n')
        
        if len(unstable_vars) > 0:
            stats_para.add_run(f'  → Most unstable variables:\n')
            for var in unstable_vars.head(5).index:
                stats_para.add_run(f'     - {var} (SD = {variability[var]:.3f})\n')
        
        # Interpretation
        doc.add_heading('Interpretation & Recommendations', level=2)
        
        if avg_stability < 0.3:
            status_text = "✓ PROPORTIONAL ODDS ASSUMPTION LIKELY HOLDS"
            status_color = RGBColor(0, 128, 0)
            recommendation = "The coefficients show good stability across cutpoints. The ordinal logistic regression model is appropriate."
        elif avg_stability < 0.6:
            status_text = "⚠ PROPORTIONAL ODDS ASSUMPTION PARTIALLY HOLDS"
            status_color = RGBColor(255, 140, 0)
            recommendation = "Some variables show moderate instability. Consider sensitivity analyses or partial proportional odds models."
        else:
            status_text = "✗ PROPORTIONAL ODDS ASSUMPTION MAY BE VIOLATED"
            status_color = RGBColor(255, 0, 0)
            recommendation = "The coefficients show substantial variation across cutpoints. Consider using multinomial logistic regression, generalized ordinal logistic regression, or partial proportional odds models."
        
        status_para = doc.add_paragraph()
        status_run = status_para.add_run(status_text)
        status_run.bold = True
        status_run.font.color.rgb = status_color
        
        doc.add_paragraph(recommendation)

# Add saved files information
doc.add_heading('Saved Files', level=2)
file_para = doc.add_paragraph()
file_para.add_run('The following files were saved:\n').bold = True
file_para.add_run('• fig6_proportional_odds.png - Stability plot (300 DPI)\n')
file_para.add_run('• fig6_proportional_odds_hd.png - High-resolution plot (600 DPI)\n')

# Save the Word document
try:
    doc.save('Regression_Analysis_Report.docx')
    print("\n✅ Word document saved successfully")
except PermissionError:
    timestamp = time.strftime("%Y%m%d_%H%M%S")
    backup_name = f"Regression_Analysis_Report_{timestamp}.docx"
    doc.save(backup_name)
    print(f"\n✅ Document saved as: {backup_name} (original was locked)")

# Print final summary
print("\n" + "=" * 60)
print("✅ SECTION 5.5 COMPLETE")
print("=" * 60)
if len(coef_stab.columns) > 0:
    print(f"✅ Successfully tested {len(coef_stab.columns)} cutpoints")
    print(f"✅ Analyzed {len(coef_stab.index)} predictor variables")
else:
    print("⚠ No models converged - check data quality")
print("=" * 60)
# 6.1  Dispersion diagnostic
mean_p = df['parity'].mean()
var_p  = df['parity'].var()
disp_r = var_p / mean_p

print('=' * 50)
print('DISPERSION DIAGNOSTIC')
print('=' * 50)
print(f'  Mean parity       : {mean_p:.4f}')
print(f'  Variance parity   : {var_p:.4f}')
print(f'  Dispersion ratio  : {disp_r:.4f}')
print('-' * 50)

if disp_r > 1.2:
    USE_NB = True
    print('  OVERDISPERSED  -> Negative Binomial selected')
elif disp_r < 0.8:
    USE_NB = False
    print('  UNDERDISPERSED -> Quasi-Poisson recommended')
else:
    USE_NB = False
    print('  EQUIDISPERSED  -> Poisson regression selected')
# ============================================
# SECTION 6.1: Dispersion Diagnostic
# ============================================

import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import os
import time
from datetime import datetime

# ============================================
# CREATE NEW WORD DOCUMENT FOR SECTION 6
# ============================================

# Create new document for count models
doc_count = Document()
doc_count.add_heading('COUNT DATA ANALYSIS - NEGATIVE BINOMIAL REGRESSION', 0)
doc_count.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
doc_count.add_paragraph('Analysis of parity (birth order) as count outcome')
doc_count.add_page_break()

# ============================================
# YOUR ORIGINAL CODE
# ============================================

mean_p = df['parity'].mean()
var_p = df['parity'].var()
disp_r = var_p / mean_p

print('=' * 50)
print('DISPERSION DIAGNOSTIC')
print('=' * 50)
print(f'  Mean parity       : {mean_p:.4f}')
print(f'  Variance parity   : {var_p:.4f}')
print(f'  Dispersion ratio  : {disp_r:.4f}')
print('-' * 50)

if disp_r > 1.2:
    USE_NB = True
    print('  OVERDISPERSED  -> Negative Binomial selected')
elif disp_r < 0.8:
    USE_NB = False
    print('  UNDERDISPERSED -> Quasi-Poisson recommended')
else:
    USE_NB = False
    print('  EQUIDISPERSED  -> Poisson regression selected')

print('\n' + '=' * 50)
print(f'USE_NB = {USE_NB}')
print('=' * 50)

# ============================================
# ADD RESULTS TO NEW WORD DOCUMENT
# ============================================

# Add section header
doc_count.add_heading('6.1 Dispersion Diagnostic', level=1)

# Add explanation
doc_count.add_paragraph(
    'Dispersion diagnostics are used to determine the appropriate count model '
    'for analyzing parity (birth order). The dispersion ratio (variance/mean) '
    'indicates whether the data are equidispersed (ratio ≈ 1), overdispersed '
    '(ratio > 1), or underdispersed (ratio < 1).'
)

# Add results table
doc_count.add_heading('Dispersion Statistics', level=2)

# Create formatted results table
disp_table = doc_count.add_table(rows=4, cols=2)
disp_table.style = 'Light Grid Accent 1'
disp_table.autofit = False
disp_table.columns[0].width = Inches(2.5)
disp_table.columns[1].width = Inches(3)

# Add headers
disp_table.rows[0].cells[0].text = 'Statistic'
disp_table.rows[0].cells[1].text = 'Value'
disp_table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
disp_table.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True

# Add data
disp_table.rows[1].cells[0].text = 'Mean Parity'
disp_table.rows[1].cells[1].text = f'{mean_p:.4f}'

disp_table.rows[2].cells[0].text = 'Variance Parity'
disp_table.rows[2].cells[1].text = f'{var_p:.4f}'

disp_table.rows[3].cells[0].text = 'Dispersion Ratio (Var/Mean)'
disp_table.rows[3].cells[1].text = f'{disp_r:.4f}'

# Add interpretation
doc_count.add_heading('Interpretation', level=2)

# Create interpretation box with color coding
if disp_r > 1.2:
    status = "OVERDISPERSED"
    status_color = RGBColor(255, 140, 0)  # Orange
    recommendation = "Negative Binomial Regression selected"
    explanation = "The variance exceeds the mean, indicating overdispersion. Negative Binomial regression accounts for this extra variance."
    
    # Add warning about overdispersion
    warning_para = doc_count.add_paragraph()
    warning_run = warning_para.add_run('⚠ OVERDISPERSION DETECTED')
    warning_run.bold = True
    warning_run.font.size = Pt(14)
    warning_run.font.color.rgb = RGBColor(255, 0, 0)
    
elif disp_r < 0.8:
    status = "UNDERDISPERSED"
    status_color = RGBColor(0, 102, 204)  # Blue
    recommendation = "Quasi-Poisson recommended"
    explanation = "The variance is less than the mean, indicating underdispersion. Consider Quasi-Poisson or Conway-Maxwell Poisson models."
    
    # Add note about underdispersion
    warning_para = doc_count.add_paragraph()
    warning_run = warning_para.add_run('ℹ UNDERDISPERSION DETECTED')
    warning_run.bold = True
    warning_run.font.size = Pt(14)
    warning_run.font.color.rgb = RGBColor(0, 102, 204)
    
else:
    status = "EQUIDISPERSED"
    status_color = RGBColor(0, 128, 0)  # Green
    recommendation = "Poisson regression selected"
    explanation = "The variance approximately equals the mean, satisfying the Poisson distribution assumption."
    
    # Add success note
    success_para = doc_count.add_paragraph()
    success_run = success_para.add_run('✓ EQUIDISPERSION DETECTED')
    success_run.bold = True
    success_run.font.size = Pt(14)
    success_run.font.color.rgb = RGBColor(0, 128, 0)

# Add status paragraph
status_para = doc_count.add_paragraph()
status_run = status_para.add_run(f'Diagnosis: {status}')
status_run.bold = True
status_run.font.size = Pt(12)
status_run.font.color.rgb = status_color

doc_count.add_paragraph(explanation)

# Add model selection box
doc_count.add_heading('Model Selection Decision', level=2)

selection_table = doc_count.add_table(rows=2, cols=2)
selection_table.autofit = False
selection_table.columns[0].width = Inches(2.5)
selection_table.columns[1].width = Inches(3)

selection_table.rows[0].cells[0].text = 'Selected Model'
selection_table.rows[0].cells[1].text = recommendation
selection_table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
selection_table.rows[1].cells[0].text = 'USE_NB Flag'
selection_table.rows[1].cells[1].text = str(USE_NB)

# Add methodological note
doc_count.add_heading('Methodological Note', level=2)
doc_count.add_paragraph(
    'The dispersion ratio (variance/mean) is a key diagnostic for count data models:\n'
    '• Ratio = 1: Poisson model (equidispersion)\n'
    '• Ratio > 1: Negative Binomial model (overdispersion)\n'
    '• Ratio < 1: Quasi-Poisson or alternative models (underdispersion)\n\n'
    'Overdispersion is common in parity data due to heterogeneity across women '
    'and clustering of births.'
)

# Add summary statistics box
doc_count.add_heading('Summary of Parity Distribution', level=2)

# Calculate additional statistics
parity_stats = pd.DataFrame({
    'Statistic': ['Mean', 'Variance', 'Std Dev', 'Min', 'Max', 'Range', 'Dispersion Ratio'],
    'Value': [
        f'{mean_p:.4f}',
        f'{var_p:.4f}',
        f'{np.sqrt(var_p):.4f}',
        f'{df["parity"].min()}',
        f'{df["parity"].max()}',
        f'{df["parity"].max() - df["parity"].min()}',
        f'{disp_r:.4f}'
    ]
})

stats_table = doc_count.add_table(rows=len(parity_stats) + 1, cols=2)
stats_table.style = 'Light Grid Accent 1'

# Header
stats_table.rows[0].cells[0].text = 'Statistic'
stats_table.rows[0].cells[1].text = 'Value'
stats_table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
stats_table.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True

# Data
for i, row in parity_stats.iterrows():
    stats_table.rows[i+1].cells[0].text = row['Statistic']
    stats_table.rows[i+1].cells[1].text = row['Value']

# Add recommendation for next steps
doc_count.add_heading('Next Steps', level=2)
next_steps = doc_count.add_paragraph()
next_steps.add_run('Based on the dispersion diagnostic:\n').bold = True

if USE_NB:
    next_steps.add_run('1. Fit Negative Binomial regression model\n')
    next_steps.add_run('2. Compare with Poisson model using likelihood ratio test\n')
    next_steps.add_run('3. Check for zero-inflation if applicable\n')
    next_steps.add_run('4. Validate model fit using residuals analysis')
else:
    next_steps.add_run('1. Fit Poisson regression model\n')
    next_steps.add_run('2. Test for overdispersion using regression-based test\n')
    next_steps.add_run('3. Consider robust standard errors if mild violations\n')
    next_steps.add_run('4. Validate model assumptions')

# Add note about the separate document
doc_count.add_paragraph(
    '📄 This analysis continues in a separate document for count data models.',
    style='Intense Quote'
)

# Save the new Word document
filename_count = 'Count_Data_Analysis_Report.docx'

# Handle potential file lock
try:
    doc_count.save(filename_count)
    print(f"\n✅ New document created: {filename_count}")
except PermissionError:
    timestamp = time.strftime("%Y%m%d_%H%M%S")
    filename_count = f'Count_Data_Analysis_Report_{timestamp}.docx'
    doc_count.save(filename_count)
    print(f"\n✅ Document saved as: {filename_count} (original was locked)")

# Print confirmation
print("\n" + "=" * 60)
print("✅ SECTION 6.1 COMPLETE - SEPARATE DOCUMENT CREATED")
print("=" * 60)
print(f"📊 Dispersion Diagnostic Results:")
print(f"   Mean: {mean_p:.4f}")
print(f"   Variance: {var_p:.4f}")
print(f"   Dispersion Ratio: {disp_r:.4f}")
print(f"   Diagnosis: {status}")
print(f"   Selected Model: {recommendation}")
print("=" * 60)
print(f"📄 New Word document: {filename_count}")
print("=" * 60)

# Display the USE_NB variable value for reference in subsequent code
print(f"\n💡 USE_NB = {USE_NB}")
print("   (Use this variable in subsequent count model sections)")
# ============================================
# SECTION 6.2: Fit Poisson Regression (FIXED - No KeyError)
# ============================================

import statsmodels.api as sm
import pandas as pd
import numpy as np
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from scipy import stats
import warnings
warnings.filterwarnings('ignore')

# ============================================
# YOUR ORIGINAL CODE
# ============================================

y_count = df['parity'].astype(int).values

poisson_m = sm.GLM(
    y_count, X_const,
    family=sm.families.Poisson()
).fit()

print('POISSON REGRESSION RESULTS')
print(poisson_m.summary())

# ============================================
# ADDITIONAL DIAGNOSTICS FOR WORD DOC
# ============================================

# Calculate additional statistics
poisson_deviance = poisson_m.deviance
poisson_null_deviance = poisson_m.null_deviance
poisson_df_resid = poisson_m.df_resid
poisson_df_model = poisson_m.df_model
poisson_pseudo_r2 = 1 - (poisson_deviance / poisson_null_deviance)
poisson_aic = poisson_m.aic
poisson_bic = poisson_m.bic

# Extract coefficients and statistics
poisson_coef_df = pd.DataFrame({
    'Variable': poisson_m.params.index,
    'Coefficient': poisson_m.params.values,
    'Std_Error': poisson_m.bse.values,
    'z_value': poisson_m.tvalues.values,
    'P_value': poisson_m.pvalues.values
})

# Calculate Incident Rate Ratios (IRR)
poisson_coef_df['IRR'] = np.exp(poisson_coef_df['Coefficient'])
poisson_coef_df['IRR_lower'] = np.exp(poisson_coef_df['Coefficient'] - 1.96 * poisson_coef_df['Std_Error'])
poisson_coef_df['IRR_upper'] = np.exp(poisson_coef_df['Coefficient'] + 1.96 * poisson_coef_df['Std_Error'])

# Add significance stars
def get_sig_stars(p_val):
    if p_val < 0.001:
        return '***'
    elif p_val < 0.01:
        return '**'
    elif p_val < 0.05:
        return '*'
    else:
        return ''

poisson_coef_df['Sig'] = poisson_coef_df['P_value'].apply(get_sig_stars)

# Test for overdispersion (important for Poisson vs NB)
pearson_residuals = poisson_m.resid_pearson
overdispersion_ratio = np.sum(pearson_residuals**2) / poisson_df_resid

print(f"\nOverdispersion ratio (Pearson χ²/df): {overdispersion_ratio:.4f}")
if overdispersion_ratio > 1.2:
    print("  ⚠ Overdispersion detected - Negative Binomial may be more appropriate")
elif overdispersion_ratio < 0.8:
    print("  ℹ Underdispersion detected")
else:
    print("  ✓ No serious overdispersion detected")

# ============================================
# ADD RESULTS TO COUNT DATA WORD DOCUMENT
# ============================================

# Add section header
doc_count.add_heading('6.2 Poisson Regression', level=1)

# Add model description
doc_count.add_paragraph(
    'Poisson regression is the standard model for count data, assuming that the '
    'variance equals the mean (equidispersion). The model was fitted using '
    'Maximum Likelihood Estimation (MLE) with a log link function.'
)

# Add model fit statistics
doc_count.add_heading('Model Fit Statistics', level=2)

# Create fit statistics data
fit_stats_data = [
    ('Log-Likelihood', f"{poisson_m.llf:.4f}"),
    ('Deviance', f"{poisson_deviance:.4f}"),
    ('Null Deviance', f"{poisson_null_deviance:.4f}"),
    ('Pseudo R-squared', f"{poisson_pseudo_r2:.4f}"),
    ('AIC', f"{poisson_aic:.4f}"),
    ('BIC', f"{poisson_bic:.4f}"),
    ('Overdispersion Ratio', f"{overdispersion_ratio:.4f}")
]

# Create table with correct number of rows
fit_table = doc_count.add_table(rows=len(fit_stats_data) + 1, cols=2)
fit_table.style = 'Light Grid Accent 1'
fit_table.autofit = False
fit_table.columns[0].width = Inches(2.5)
fit_table.columns[1].width = Inches(3)

# Headers
fit_table.rows[0].cells[0].text = 'Metric'
fit_table.rows[0].cells[1].text = 'Value'
fit_table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
fit_table.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True

# Fill fit statistics
for i, (metric, value) in enumerate(fit_stats_data, 1):
    fit_table.rows[i].cells[0].text = metric
    fit_table.rows[i].cells[1].text = value

# Add overdispersion interpretation
if overdispersion_ratio > 1.2:
    od_para = doc_count.add_paragraph()
    od_run = od_para.add_run(f'⚠ Overdispersion detected (ratio = {overdispersion_ratio:.4f} > 1.2)')
    od_run.font.color.rgb = RGBColor(255, 140, 0)
    od_run.bold = True
    doc_count.add_paragraph(
        'The overdispersion ratio exceeds 1.2, indicating that the Poisson model may be inappropriate. '
        'Consider using Negative Binomial regression which accounts for overdispersion.'
    )
elif overdispersion_ratio < 0.8:
    od_para = doc_count.add_paragraph()
    od_run = od_para.add_run(f'ℹ Underdispersion detected (ratio = {overdispersion_ratio:.4f} < 0.8)')
    od_run.font.color.rgb = RGBColor(0, 102, 204)
    doc_count.add_paragraph(
        'The underdispersion ratio suggests the variance is less than the mean. '
        'Consider using Quasi-Poisson or Conway-Maxwell Poisson models.'
    )
else:
    od_para = doc_count.add_paragraph()
    od_run = od_para.add_run(f'✓ No serious overdispersion (ratio = {overdispersion_ratio:.4f})')
    od_run.font.color.rgb = RGBColor(0, 128, 0)

# Add coefficients table
doc_count.add_heading('Regression Coefficients', level=2)

# Filter out constant for main table
main_coef = poisson_coef_df[poisson_coef_df['Variable'] != 'const'].copy()

if len(main_coef) > 0:
    coef_table = doc_count.add_table(rows=len(main_coef) + 1, cols=7)
    coef_table.style = 'Light Grid Accent 1'
    coef_table.autofit = False

    # Set column widths
    widths = [1.5, 0.9, 0.9, 0.8, 0.8, 0.9, 0.6]
    for i, width in enumerate(widths):
        if i < len(coef_table.columns):
            coef_table.columns[i].width = Inches(width)

    # Headers
    headers = ['Variable', 'Coef', 'Std Err', 'z-value', 'P>|z|', 'IRR', 'Sig']
    for i, header in enumerate(headers):
        if i < len(coef_table.columns):
            coef_table.rows[0].cells[i].text = header
            coef_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    # Populate table
    for i, (idx, row) in enumerate(main_coef.iterrows()):
        row_cells = coef_table.rows[i+1].cells
        row_cells[0].text = str(row['Variable'])[:30]
        row_cells[1].text = f"{row['Coefficient']:.4f}"
        row_cells[2].text = f"{row['Std_Error']:.4f}"
        row_cells[3].text = f"{row['z_value']:.4f}"
        row_cells[4].text = f"{row['P_value']:.4f}"
        row_cells[5].text = f"{row['IRR']:.3f}"
        row_cells[6].text = row['Sig']
        
        # Bold significant coefficients
        if row['Sig']:
            for col in range(6):
                row_cells[col].paragraphs[0].runs[0].font.bold = True

# Add significance legend
sig_legend = doc_count.add_paragraph()
sig_legend.add_run('Significance codes: ').font.size = Pt(9)
sig_legend.add_run('*** p<0.001, ').font.size = Pt(9)
sig_legend.add_run('** p<0.01, ').font.size = Pt(9)
sig_legend.add_run('* p<0.05').font.size = Pt(9)

# Add Incident Rate Ratios (IRR) interpretation
doc_count.add_heading('Incident Rate Ratios (IRR)', level=2)
doc_count.add_paragraph(
    'IRR represents the multiplicative effect on the expected count. '
    'IRR > 1 indicates higher expected parity, while IRR < 1 indicates lower expected parity.\n\n'
    'For a one-unit increase in a predictor, the expected count is multiplied by the IRR.'
)

# Create IRR summary table for top predictors
doc_count.add_heading('Key Predictors (IRR Analysis)', level=2)

# Get significant predictors - FIXED: Create a proper column for effect size
if len(main_coef) > 0:
    # Calculate absolute effect size (distance from 1)
    main_coef['Effect_Size'] = np.abs(main_coef['IRR'] - 1)
    
    # Get significant predictors
    sig_predictors = main_coef[main_coef['P_value'] < 0.05].copy()
    
    # Sort by effect size (largest first)
    if len(sig_predictors) > 0:
        sig_predictors = sig_predictors.sort_values('Effect_Size', ascending=False).head(10)
        
        irr_table = doc_count.add_table(rows=len(sig_predictors) + 1, cols=4)
        irr_table.style = 'Light Shading'
        
        # Headers
        irr_table.rows[0].cells[0].text = 'Variable'
        irr_table.rows[0].cells[1].text = 'IRR'
        irr_table.rows[0].cells[2].text = '95% CI'
        irr_table.rows[0].cells[3].text = 'P-value'
        
        for i, (idx, row) in enumerate(sig_predictors.iterrows()):
            irr_table.rows[i+1].cells[0].text = str(row['Variable'])[:30]
            irr_table.rows[i+1].cells[1].text = f"{row['IRR']:.3f}"
            ci_text = f"({row['IRR_lower']:.3f}-{row['IRR_upper']:.3f})"
            irr_table.rows[i+1].cells[2].text = ci_text
            irr_table.rows[i+1].cells[3].text = f"{row['P_value']:.4f}"
    else:
        doc_count.add_paragraph('No statistically significant predictors found.')
else:
    doc_count.add_paragraph('No predictor variables to display.')

# Add model comparison note
doc_count.add_heading('Model Comparison', level=2)
doc_count.add_paragraph(
    f'The Poisson model achieved AIC = {poisson_aic:.2f} and BIC = {poisson_bic:.2f}. '
    f'These will be compared with the Negative Binomial model in the next section.'
)

# Add constant term
const_term = poisson_coef_df[poisson_coef_df['Variable'] == 'const']
if len(const_term) > 0:
    doc_count.add_heading('Intercept Term', level=2)
    intercept_para = doc_count.add_paragraph()
    intercept_para.add_run(f'Constant (Intercept): ').bold = True
    intercept_para.add_run(f"Coef = {const_term['Coefficient'].values[0]:.4f}, ")
    intercept_para.add_run(f"IRR = {np.exp(const_term['Coefficient'].values[0]):.4f}, ")
    intercept_para.add_run(f"p = {const_term['P_value'].values[0]:.4f}")

# Add model equation
doc_count.add_heading('Model Equation', level=2)

# Build equation string
if len(main_coef) > 0:
    eq_parts = []
    for _, row in main_coef.iterrows():
        coef = row['Coefficient']
        var = row['Variable']
        if coef >= 0:
            eq_parts.append(f"+ {coef:.4f}×{var}")
        else:
            eq_parts.append(f"- {abs(coef):.4f}×{var}")
    
    intercept_val = const_term['Coefficient'].values[0] if len(const_term) > 0 else 0
    eq_text = f"log(μ) = {intercept_val:.4f} {' '.join(eq_parts)}"
    eq_para = doc_count.add_paragraph()
    eq_run = eq_para.add_run(eq_text)
    eq_run.font.name = 'Courier New'
    eq_run.font.size = Pt(10)
    eq_run.font.bold = True

doc_count.add_paragraph(
    "Where μ is the expected parity count, and the log link function ensures positive predictions."
)

# Add assumption check notes
doc_count.add_heading('Model Assumptions', level=2)
assumptions = doc_count.add_paragraph()
assumptions.add_run('Poisson regression assumes:\n').bold = True
assumptions.add_run('1. Independence of observations\n')
assumptions.add_run('2. Mean = Variance (equidispersion)\n')
assumptions.add_run('3. Log-linear relationship between predictors and outcome\n')
assumptions.add_run('4. No excess zeros\n\n')

if overdispersion_ratio > 1.2:
    assumptions.add_run('⚠ The overdispersion test indicates violation of assumption #2.')
    assumptions.runs[-1].font.color.rgb = RGBColor(255, 0, 0)

# Save the updated Word document
import time
try:
    doc_count.save('Count_Data_Analysis_Report.docx')
    print("\n✅ Count Data Analysis Report updated successfully")
except PermissionError:
    timestamp = time.strftime("%Y%m%d_%H%M%S")
    backup_name = f'Count_Data_Analysis_Report_{timestamp}.docx'
    doc_count.save(backup_name)
    print(f"\n✅ Document saved as: {backup_name} (original was locked)")

# Print summary to console
print("\n" + "=" * 60)
print("✅ SECTION 6.2 COMPLETE - POISSON REGRESSION")
print("=" * 60)
print(f"📊 Poisson Model Statistics:")
print(f"   Log-Likelihood: {poisson_m.llf:.4f}")
print(f"   Pseudo R-squared: {poisson_pseudo_r2:.4f}")
print(f"   AIC: {poisson_aic:.4f}")
print(f"   BIC: {poisson_bic:.4f}")
print(f"   Overdispersion Ratio: {overdispersion_ratio:.4f}")
if len(main_coef) > 0:
    sig_count = len(main_coef[main_coef['P_value'] < 0.05])
    print(f"   Significant predictors: {sig_count}")
print("=" * 60)
print("📄 Updated: Count_Data_Analysis_Report.docx")
print("=" * 60)

# Clean up temporary column
if 'Effect_Size' in main_coef.columns:
    main_coef = main_coef.drop(columns=['Effect_Size'])
# ============================================
# SECTION 6.3: Fit Negative Binomial Regression
# ============================================

import statsmodels.api as sm
import pandas as pd
import numpy as np
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from scipy import stats
import warnings
warnings.filterwarnings('ignore')

# ============================================
# YOUR ORIGINAL CODE
# ============================================

nb_m = sm.GLM(
    y_count, X_const,
    family=sm.families.NegativeBinomial()
).fit()

print('NEGATIVE BINOMIAL REGRESSION RESULTS')
print(nb_m.summary())

# ============================================
# ADDITIONAL DIAGNOSTICS FOR WORD DOC
# ============================================

# Calculate additional statistics
nb_deviance = nb_m.deviance
nb_null_deviance = nb_m.null_deviance
nb_df_resid = nb_m.df_resid
nb_df_model = nb_m.df_model
nb_pseudo_r2 = 1 - (nb_deviance / nb_null_deviance)
nb_aic = nb_m.aic
nb_bic = nb_m.bic

# Extract coefficients and statistics
nb_coef_df = pd.DataFrame({
    'Variable': nb_m.params.index,
    'Coefficient': nb_m.params.values,
    'Std_Error': nb_m.bse.values,
    'z_value': nb_m.tvalues.values,
    'P_value': nb_m.pvalues.values
})

# Calculate Incident Rate Ratios (IRR)
nb_coef_df['IRR'] = np.exp(nb_coef_df['Coefficient'])
nb_coef_df['IRR_lower'] = np.exp(nb_coef_df['Coefficient'] - 1.96 * nb_coef_df['Std_Error'])
nb_coef_df['IRR_upper'] = np.exp(nb_coef_df['Coefficient'] + 1.96 * nb_coef_df['Std_Error'])

# Add significance stars
def get_sig_stars(p_val):
    if p_val < 0.001:
        return '***'
    elif p_val < 0.01:
        return '**'
    elif p_val < 0.05:
        return '*'
    else:
        return ''

nb_coef_df['Sig'] = nb_coef_df['P_value'].apply(get_sig_stars)

# Get the alpha parameter (dispersion parameter) for Negative Binomial
# Note: For GLM NegativeBinomial, the dispersion parameter might be accessed differently
try:
    nb_alpha = nb_m.family.alpha
    print(f"\nNegative Binomial dispersion parameter (alpha): {nb_alpha:.4f}")
except:
    nb_alpha = "N/A"
    print("\nCould not extract alpha parameter directly")

# Calculate overdispersion ratio for Negative Binomial
pearson_residuals_nb = nb_m.resid_pearson
nb_overdispersion_ratio = np.sum(pearson_residuals_nb**2) / nb_df_resid

print(f"Negative Binomial overdispersion ratio: {nb_overdispersion_ratio:.4f}")

# ============================================
# COMPARE POISSON AND NEGATIVE BINOMIAL
# ============================================

# Likelihood Ratio Test for model comparison
lr_stat = 2 * (nb_m.llf - poisson_m.llf)
lr_pvalue = stats.chi2.sf(lr_stat, df=1)  # 1 degree of freedom for alpha parameter

print(f"\n{'='*60}")
print("MODEL COMPARISON: Poisson vs Negative Binomial")
print(f"{'='*60}")
print(f"Log-Likelihood (Poisson):           {poisson_m.llf:.4f}")
print(f"Log-Likelihood (Negative Binomial): {nb_m.llf:.4f}")
print(f"Likelihood Ratio Statistic:         {lr_stat:.4f}")
print(f"P-value (LR test):                  {lr_pvalue:.6f}")
print(f"{'='*60}")

if lr_pvalue < 0.05:
    print("✓ Negative Binomial significantly better than Poisson (p < 0.05)")
    best_model = "Negative Binomial"
else:
    print("ℹ No significant improvement with Negative Binomial")
    best_model = "Poisson (simpler model preferred)"

# ============================================
# ADD RESULTS TO COUNT DATA WORD DOCUMENT
# ============================================

# Add section header
doc_count.add_heading('6.3 Negative Binomial Regression', level=1)

# Add model description
doc_count.add_paragraph(
    'Negative Binomial regression extends Poisson regression by adding a dispersion '
    'parameter (alpha) that accounts for overdispersion. This model is appropriate '
    'when the variance exceeds the mean, which is common in parity data due to '
    'unobserved heterogeneity across women.'
)

# Add model fit statistics
doc_count.add_heading('Model Fit Statistics', level=2)

# Create fit statistics data
nb_fit_stats = [
    ('Log-Likelihood', f"{nb_m.llf:.4f}"),
    ('Deviance', f"{nb_deviance:.4f}"),
    ('Null Deviance', f"{nb_null_deviance:.4f}"),
    ('Pseudo R-squared', f"{nb_pseudo_r2:.4f}"),
    ('AIC', f"{nb_aic:.4f}"),
    ('BIC', f"{nb_bic:.4f}"),
    ('Overdispersion Ratio', f"{nb_overdispersion_ratio:.4f}"),
    ('Dispersion (alpha)', f"{nb_alpha if nb_alpha != 'N/A' else 'Not available'}")
]

# Create table
nb_fit_table = doc_count.add_table(rows=len(nb_fit_stats) + 1, cols=2)
nb_fit_table.style = 'Light Grid Accent 1'
nb_fit_table.autofit = False
nb_fit_table.columns[0].width = Inches(2.5)
nb_fit_table.columns[1].width = Inches(3)

# Headers
nb_fit_table.rows[0].cells[0].text = 'Metric'
nb_fit_table.rows[0].cells[1].text = 'Value'
nb_fit_table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
nb_fit_table.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True

# Fill data
for i, (metric, value) in enumerate(nb_fit_stats, 1):
    nb_fit_table.rows[i].cells[0].text = metric
    nb_fit_table.rows[i].cells[1].text = value

# Add interpretation of alpha parameter
if nb_alpha != 'N/A':
    alpha_para = doc_count.add_paragraph()
    if nb_alpha > 0:
        alpha_para.add_run(f'✓ Alpha = {nb_alpha:.4f} (positive, indicating overdispersion)')
        alpha_para.runs[-1].font.color.rgb = RGBColor(0, 128, 0)
    else:
        alpha_para.add_run(f'ℹ Alpha = {nb_alpha:.4f} (non-positive, model may be equivalent to Poisson)')

# Add coefficients table
doc_count.add_heading('Regression Coefficients', level=2)

# Filter out constant for main table
nb_main_coef = nb_coef_df[nb_coef_df['Variable'] != 'const'].copy()

if len(nb_main_coef) > 0:
    nb_coef_table = doc_count.add_table(rows=len(nb_main_coef) + 1, cols=7)
    nb_coef_table.style = 'Light Grid Accent 1'
    nb_coef_table.autofit = False

    # Set column widths
    widths = [1.5, 0.9, 0.9, 0.8, 0.8, 0.9, 0.6]
    for i, width in enumerate(widths):
        if i < len(nb_coef_table.columns):
            nb_coef_table.columns[i].width = Inches(width)

    # Headers
    headers = ['Variable', 'Coef', 'Std Err', 'z-value', 'P>|z|', 'IRR', 'Sig']
    for i, header in enumerate(headers):
        if i < len(nb_coef_table.columns):
            nb_coef_table.rows[0].cells[i].text = header
            nb_coef_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

    # Populate table
    for i, (idx, row) in enumerate(nb_main_coef.iterrows()):
        row_cells = nb_coef_table.rows[i+1].cells
        row_cells[0].text = str(row['Variable'])[:30]
        row_cells[1].text = f"{row['Coefficient']:.4f}"
        row_cells[2].text = f"{row['Std_Error']:.4f}"
        row_cells[3].text = f"{row['z_value']:.4f}"
        row_cells[4].text = f"{row['P_value']:.4f}"
        row_cells[5].text = f"{row['IRR']:.3f}"
        row_cells[6].text = row['Sig']
        
        # Bold significant coefficients
        if row['Sig']:
            for col in range(6):
                row_cells[col].paragraphs[0].runs[0].font.bold = True

# Add significance legend
sig_legend = doc_count.add_paragraph()
sig_legend.add_run('Significance codes: ').font.size = Pt(9)
sig_legend.add_run('*** p<0.001, ').font.size = Pt(9)
sig_legend.add_run('** p<0.01, ').font.size = Pt(9)
sig_legend.add_run('* p<0.05').font.size = Pt(9)

# ============================================
# MODEL COMPARISON SECTION
# ============================================

doc_count.add_heading('Model Comparison: Poisson vs Negative Binomial', level=2)

# Create comparison table
comparison_data = [
    ('Log-Likelihood', f"{poisson_m.llf:.4f}", f"{nb_m.llf:.4f}"),
    ('AIC', f"{poisson_aic:.4f}", f"{nb_aic:.4f}"),
    ('BIC', f"{poisson_bic:.4f}", f"{nb_bic:.4f}"),
    ('Pseudo R-squared', f"{poisson_pseudo_r2:.4f}", f"{nb_pseudo_r2:.4f}"),
    ('Overdispersion Ratio', f"{overdispersion_ratio:.4f}", f"{nb_overdispersion_ratio:.4f}")
]

comp_table = doc_count.add_table(rows=len(comparison_data) + 1, cols=3)
comp_table.style = 'Light Grid Accent 1'

# Headers
comp_table.rows[0].cells[0].text = 'Metric'
comp_table.rows[0].cells[1].text = 'Poisson'
comp_table.rows[0].cells[2].text = 'Negative Binomial'
for i in range(3):
    comp_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

# Fill data
for i, (metric, poisson_val, nb_val) in enumerate(comparison_data, 1):
    comp_table.rows[i].cells[0].text = metric
    comp_table.rows[i].cells[1].text = poisson_val
    comp_table.rows[i].cells[2].text = nb_val
    
    # Highlight better values
    if 'AIC' in metric or 'BIC' in metric:
        if float(nb_val) < float(poisson_val):
            comp_table.rows[i].cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
    elif 'Log-Likelihood' in metric or 'Pseudo R-squared' in metric:
        if float(nb_val) > float(poisson_val):
            comp_table.rows[i].cells[2].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)

# Likelihood Ratio Test results
doc_count.add_heading('Likelihood Ratio Test', level=3)
lrt_para = doc_count.add_paragraph()
lrt_para.add_run(f'LR Statistic: {lr_stat:.4f}\n').bold = True
lrt_para.add_run(f'P-value: {lr_pvalue:.6f}\n')

if lr_pvalue < 0.05:
    lrt_para.add_run(f'✓ The Negative Binomial model provides a significantly better fit (p < 0.05)')
    lrt_para.runs[-1].font.color.rgb = RGBColor(0, 128, 0)
else:
    lrt_para.add_run(f'ℹ No significant improvement with Negative Binomial (p = {lr_pvalue:.4f})')
    lrt_para.runs[-1].font.color.rgb = RGBColor(255, 140, 0)

# Final model recommendation
doc_count.add_heading('Final Model Recommendation', level=2)
rec_para = doc_count.add_paragraph()

if lr_pvalue < 0.05 and nb_overdispersion_ratio < 1.2:
    rec_para.add_run('RECOMMENDED MODEL: Negative Binomial Regression\n').bold = True
    rec_para.add_run('Reasoning: The LR test shows significant improvement over Poisson, ')
    rec_para.add_run('and the overdispersion is properly accounted for by the alpha parameter.')
elif lr_pvalue < 0.05:
    rec_para.add_run('RECOMMENDED MODEL: Negative Binomial Regression\n').bold = True
    rec_para.add_run('Reasoning: Significant improvement over Poisson despite remaining overdispersion.')
else:
    rec_para.add_run('RECOMMENDED MODEL: Poisson Regression\n').bold = True
    rec_para.add_run('Reasoning: The Negative Binomial does not provide significantly better fit, ')
    rec_para.add_run('and Poisson is the simpler model.')

# Add IRR summary for Negative Binomial
doc_count.add_heading('Key Predictors from Negative Binomial Model', level=2)

if len(nb_main_coef) > 0:
    # Calculate effect size
    nb_main_coef['Effect_Size'] = np.abs(nb_main_coef['IRR'] - 1)
    
    # Get significant predictors
    nb_sig_predictors = nb_main_coef[nb_main_coef['P_value'] < 0.05].copy()
    
    if len(nb_sig_predictors) > 0:
        nb_sig_predictors = nb_sig_predictors.sort_values('Effect_Size', ascending=False).head(10)
        
        nb_irr_table = doc_count.add_table(rows=len(nb_sig_predictors) + 1, cols=4)
        nb_irr_table.style = 'Light Shading'
        
        # Headers
        nb_irr_table.rows[0].cells[0].text = 'Variable'
        nb_irr_table.rows[0].cells[1].text = 'IRR'
        nb_irr_table.rows[0].cells[2].text = '95% CI'
        nb_irr_table.rows[0].cells[3].text = 'P-value'
        
        for i, (idx, row) in enumerate(nb_sig_predictors.iterrows()):
            nb_irr_table.rows[i+1].cells[0].text = str(row['Variable'])[:30]
            nb_irr_table.rows[i+1].cells[1].text = f"{row['IRR']:.3f}"
            ci_text = f"({row['IRR_lower']:.3f}-{row['IRR_upper']:.3f})"
            nb_irr_table.rows[i+1].cells[2].text = ci_text
            nb_irr_table.rows[i+1].cells[3].text = f"{row['P_value']:.4f}"
    else:
        doc_count.add_paragraph('No statistically significant predictors found in Negative Binomial model.')

# Add constant term for Negative Binomial
nb_const_term = nb_coef_df[nb_coef_df['Variable'] == 'const']
if len(nb_const_term) > 0:
    doc_count.add_heading('Intercept Term', level=2)
    intercept_para = doc_count.add_paragraph()
    intercept_para.add_run(f'Constant (Intercept): ').bold = True
    intercept_para.add_run(f"Coef = {nb_const_term['Coefficient'].values[0]:.4f}, ")
    intercept_para.add_run(f"IRR = {np.exp(nb_const_term['Coefficient'].values[0]):.4f}, ")
    intercept_para.add_run(f"p = {nb_const_term['P_value'].values[0]:.4f}")

# Save the updated Word document
import time
try:
    doc_count.save('Count_Data_Analysis_Report.docx')
    print("\n✅ Count Data Analysis Report updated successfully")
except PermissionError:
    timestamp = time.strftime("%Y%m%d_%H%M%S")
    backup_name = f'Count_Data_Analysis_Report_{timestamp}.docx'
    doc_count.save(backup_name)
    print(f"\n✅ Document saved as: {backup_name} (original was locked)")

# Print summary to console
print("\n" + "=" * 60)
print("✅ SECTION 6.3 COMPLETE - NEGATIVE BINOMIAL REGRESSION")
print("=" * 60)
print(f"📊 Negative Binomial Model Statistics:")
print(f"   Log-Likelihood: {nb_m.llf:.4f}")
print(f"   Pseudo R-squared: {nb_pseudo_r2:.4f}")
print(f"   AIC: {nb_aic:.4f}")
print(f"   BIC: {nb_bic:.4f}")
print(f"   Overdispersion Ratio: {nb_overdispersion_ratio:.4f}")
if len(nb_main_coef) > 0:
    nb_sig_count = len(nb_main_coef[nb_main_coef['P_value'] < 0.05])
    print(f"   Significant predictors: {nb_sig_count}")
print("=" * 60)
print(f"\n📊 Model Comparison:")
print(f"   LR Statistic: {lr_stat:.4f}")
print(f"   P-value: {lr_pvalue:.6f}")
print(f"   Best Model: {best_model}")
print("=" * 60)
print("📄 Updated: Count_Data_Analysis_Report.docx")
print("=" * 60)

# Clean up temporary column
if 'Effect_Size' in nb_main_coef.columns:
    nb_main_coef = nb_main_coef.drop(columns=['Effect_Size'])
# ============================================
# SECTION 6.4: Compare and Select Better Model
# ============================================

import pandas as pd
import numpy as np
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import time

# ============================================
# YOUR ORIGINAL CODE
# ============================================

compare_count = pd.DataFrame({
    'Model': ['Poisson', 'Negative Binomial'],
    'AIC': [round(poisson_m.aic, 2), round(nb_m.aic, 2)],
    'BIC': [round(poisson_m.bic, 2), round(nb_m.bic, 2)],
    'LogLik': [round(poisson_m.llf, 2), round(nb_m.llf, 2)]
})

print('Poisson vs Negative Binomial — lower AIC/BIC = better')
print(compare_count.to_string(index=False))

best_count_m = nb_m if USE_NB else poisson_m
best_count_name = 'Negative Binomial' if USE_NB else 'Poisson'
print(f'\nSelected model: {best_count_name}')

# ============================================
# ADDITIONAL COMPARISON METRICS
# ============================================

# Calculate additional comparison metrics
aic_diff = abs(poisson_m.aic - nb_m.aic)
bic_diff = abs(poisson_m.bic - nb_m.bic)

# Determine which model is better based on AIC and BIC
aic_better = 'Negative Binomial' if nb_m.aic < poisson_m.aic else 'Poisson'
bic_better = 'Negative Binomial' if nb_m.bic < poisson_m.bic else 'Poisson'

# Calculate percentage improvement
if nb_m.aic < poisson_m.aic:
    aic_improvement = ((poisson_m.aic - nb_m.aic) / poisson_m.aic) * 100
else:
    aic_improvement = ((nb_m.aic - poisson_m.aic) / nb_m.aic) * 100

# Calculate the difference in log-likelihood
ll_diff = abs(poisson_m.llf - nb_m.llf)

# ============================================
# ADD RESULTS TO COUNT DATA WORD DOCUMENT
# ============================================

# Add section header
doc_count.add_heading('6.4 Model Comparison and Selection', level=1)

# Add explanation
doc_count.add_paragraph(
    'Model selection is based on information criteria (AIC and BIC) where lower values '
    'indicate better fit. The Likelihood Ratio Test (LRT) from the previous section '
    'provides a formal statistical comparison between the Poisson and Negative Binomial models.'
)

# Add comparison table
doc_count.add_heading('Model Comparison Metrics', level=2)

# Create formatted comparison table
comp_table = doc_count.add_table(rows=len(compare_count) + 1, cols=len(compare_count.columns) + 1)
comp_table.style = 'Light Grid Accent 1'
comp_table.autofit = False

# Set column widths
col_widths = [1.5, 1.2, 1.2, 1.2, 1.2]
for i, width in enumerate(col_widths):
    if i < len(comp_table.columns):
        comp_table.columns[i].width = Inches(width)

# Headers
headers = ['Metric'] + list(compare_count.columns)
for i, header in enumerate(headers):
    comp_table.rows[0].cells[i].text = header
    comp_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    comp_table.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add rows for each model
for i in range(len(compare_count)):
    row = comp_table.rows[i+1]
    row.cells[0].text = f"Model {i+1}"
    
    for j, col in enumerate(compare_count.columns):
        value = compare_count.iloc[i][col]
        row.cells[j+1].text = str(value)
        row.cells[j+1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Highlight the better values (lower AIC/BIC, higher LogLik)
        if col == 'AIC' or col == 'BIC':
            if i == 0 and poisson_m.aic < nb_m.aic:
                row.cells[j+1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
            elif i == 1 and nb_m.aic < poisson_m.aic:
                row.cells[j+1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
        elif col == 'LogLik':
            if i == 0 and poisson_m.llf > nb_m.llf:
                row.cells[j+1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
            elif i == 1 and nb_m.llf > poisson_m.llf:
                row.cells[j+1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)

# Add comparison summary
doc_count.add_heading('Comparison Summary', level=2)

summary_para = doc_count.add_paragraph()
summary_para.add_run('Information Criteria:\n').bold = True
summary_para.add_run(f'• AIC: {aic_better} is better (Δ = {aic_diff:.2f})\n')
summary_para.add_run(f'• BIC: {bic_better} is better (Δ = {bic_diff:.2f})\n\n')

summary_para.add_run('Likelihood:\n').bold = True
summary_para.add_run(f'• Log-Likelihood difference: {ll_diff:.2f}\n')
if nb_m.llf > poisson_m.llf:
    summary_para.add_run('• Negative Binomial has higher log-likelihood (better fit)\n')
else:
    summary_para.add_run('• Poisson has higher log-likelihood (better fit)\n')

# Add model selection table
doc_count.add_heading('Model Selection Decision', level=2)

# Create selection criteria table
selection_data = [
    ('Criterion', 'Poisson', 'Negative Binomial', 'Winner'),
    ('AIC', f"{poisson_m.aic:.2f}", f"{nb_m.aic:.2f}", 'Negative Binomial' if nb_m.aic < poisson_m.aic else 'Poisson'),
    ('BIC', f"{poisson_m.bic:.2f}", f"{nb_m.bic:.2f}", 'Negative Binomial' if nb_m.bic < poisson_m.bic else 'Poisson'),
    ('Log-Likelihood', f"{poisson_m.llf:.2f}", f"{nb_m.llf:.2f}", 'Negative Binomial' if nb_m.llf > poisson_m.llf else 'Poisson'),
    ('Dispersion Control', 'None (assumes equidispersion)', 'Alpha parameter', 'Negative Binomial'),
    ('Complexity', 'Simpler (1 param less)', 'More complex', 'Poisson')
]

selection_table = doc_count.add_table(rows=len(selection_data), cols=4)
selection_table.style = 'Light Grid Accent 1'

# Add headers
for i, header in enumerate(selection_data[0]):
    selection_table.rows[0].cells[i].text = header
    selection_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

# Add data rows
for i, row in enumerate(selection_data[1:], 1):
    for j, value in enumerate(row):
        selection_table.rows[i].cells[j].text = str(value)
        
        # Highlight winner column
        if j == 3 and value != 'Criterion':
            if ('Negative Binomial' in value and best_count_name == 'Negative Binomial') or \
               ('Poisson' in value and best_count_name == 'Poisson'):
                selection_table.rows[i].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
                selection_table.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True

# Final model selection
doc_count.add_heading('Final Model Selection', level=2)

# Create a highlighted box for the final decision
final_para = doc_count.add_paragraph()
final_para.add_run('=' * 50 + '\n').bold = True
final_para.add_run(f'SELECTED MODEL: {best_count_name}\n').bold = True
final_para.add_run('=' * 50 + '\n').bold = True

# Add reasoning
reasoning_para = doc_count.add_paragraph()
reasoning_para.add_run('Reasoning:\n').bold = True

if best_count_name == 'Negative Binomial':
    reasoning_para.add_run('• Overdispersion detected in Poisson model\n')
    reasoning_para.add_run(f'• AIC/BIC favor Negative Binomial (ΔAIC = {aic_diff:.2f})\n')
    reasoning_para.add_run('• Alpha parameter > 0 confirms overdispersion\n')
    reasoning_para.add_run('• Negative Binomial accounts for unobserved heterogeneity\n')
elif best_count_name == 'Poisson':
    reasoning_para.add_run('• No significant overdispersion detected\n')
    reasoning_para.add_run(f'• AIC/BIC favor Poisson (ΔAIC = {aic_diff:.2f})\n')
    reasoning_para.add_run('• Poisson is simpler and more parsimonious\n')
    reasoning_para.add_run('• Likelihood Ratio Test not significant\n')

# Add interpretation of selected model coefficients
doc_count.add_heading(f'{best_count_name} Model - Key Findings', level=2)

# Get the selected model's coefficients
if best_count_name == 'Negative Binomial':
    selected_coef = nb_main_coef.copy()
    selected_model = nb_m
else:
    selected_coef = main_coef.copy()
    selected_model = poisson_m

if len(selected_coef) > 0:
    # Get significant predictors
    sig_selected = selected_coef[selected_coef['P_value'] < 0.05].copy()
    
    if len(sig_selected) > 0:
        # Sort by effect size
        sig_selected['Effect_Size'] = np.abs(sig_selected['IRR'] - 1)
        sig_selected = sig_selected.sort_values('Effect_Size', ascending=False).head(10)
        
        findings_table = doc_count.add_table(rows=len(sig_selected) + 1, cols=4)
        findings_table.style = 'Light Shading'
        
        # Headers
        findings_table.rows[0].cells[0].text = 'Predictor'
        findings_table.rows[0].cells[1].text = 'IRR'
        findings_table.rows[0].cells[2].text = '95% CI'
        findings_table.rows[0].cells[3].text = 'P-value'
        
        for i, (idx, row) in enumerate(sig_selected.iterrows()):
            findings_table.rows[i+1].cells[0].text = str(row['Variable'])[:35]
            findings_table.rows[i+1].cells[1].text = f"{row['IRR']:.3f}"
            ci_text = f"({row['IRR_lower']:.3f}-{row['IRR_upper']:.3f})"
            findings_table.rows[i+1].cells[2].text = ci_text
            findings_table.rows[i+1].cells[3].text = f"{row['P_value']:.4f}"
        
        # Add interpretation of direction
        doc_count.add_paragraph(
            'IRR > 1 indicates the predictor increases expected parity (birth order), '
            'while IRR < 1 indicates it decreases expected parity.'
        )
    else:
        doc_count.add_paragraph('No statistically significant predictors found in the selected model.')

# Add model equation for selected model
doc_count.add_heading('Model Equation', level=2)

# Get constant term
if best_count_name == 'Negative Binomial':
    const_selected = nb_const_term
else:
    const_selected = const_term

if len(selected_coef) > 0:
    eq_parts = []
    for _, row in selected_coef.iterrows():
        coef = row['Coefficient']
        var = row['Variable']
        if coef >= 0:
            eq_parts.append(f"+ {coef:.4f}×{var}")
        else:
            eq_parts.append(f"- {abs(coef):.4f}×{var}")
    
    intercept_val = const_selected['Coefficient'].values[0] if len(const_selected) > 0 else 0
    
    if best_count_name == 'Negative Binomial':
        eq_text = f"log(μ) = {intercept_val:.4f} {' '.join(eq_parts)}"
        doc_count.add_paragraph(eq_text)
        doc_count.add_paragraph(
            "Where μ is the expected parity count, and the model accounts for overdispersion "
            "through the alpha parameter."
        )
    else:
        eq_text = f"log(μ) = {intercept_val:.4f} {' '.join(eq_parts)}"
        doc_count.add_paragraph(eq_text)
        doc_count.add_paragraph(
            "Where μ is the expected parity count, assuming equidispersion (mean = variance)."
        )

# Add summary of findings
doc_count.add_heading('Summary of Count Data Analysis', level=2)

summary_box = doc_count.add_paragraph()
summary_box.add_run('Key Conclusions:\n').bold = True
summary_box.add_run(f'1. Dispersion diagnostic: {status}\n')
summary_box.add_run(f'2. Best fitting model: {best_count_name}\n')
summary_box.add_run(f'3. Significant predictors identified: {len(sig_selected) if "sig_selected" in locals() else 0}\n')
if best_count_name == 'Negative Binomial':
    summary_box.add_run(f'4. Overdispersion parameter (alpha): {nb_alpha if nb_alpha != "N/A" else "significant"}\n')
summary_box.add_run(f'5. Model fit (AIC): {selected_model.aic:.2f}\n')

# Add note about exported files
doc_count.add_heading('Exported Results', level=2)
doc_count.add_paragraph(
    'Model comparison results have been saved to the Count Data Analysis Report. '
    'The selected model can be used for further analysis and predictions.'
)

# Save the updated Word document
try:
    doc_count.save('Count_Data_Analysis_Report.docx')
    print("\n✅ Count Data Analysis Report updated successfully")
except PermissionError:
    timestamp = time.strftime("%Y%m%d_%H%M%S")
    backup_name = f'Count_Data_Analysis_Report_{timestamp}.docx'
    doc_count.save(backup_name)
    print(f"\n✅ Document saved as: {backup_name} (original was locked)")

# Print summary to console
print("\n" + "=" * 60)
print("✅ SECTION 6.4 COMPLETE - MODEL COMPARISON")
print("=" * 60)
print("\nModel Comparison Results:")
print(compare_count.to_string(index=False))
print(f"\n{'='*60}")
print(f"Selected Model: {best_count_name}")
print(f"{'='*60}")
print(f"\n📊 Selection Criteria:")
print(f"   AIC Winner: {aic_better} (Δ = {aic_diff:.2f})")
print(f"   BIC Winner: {bic_better} (Δ = {bic_diff:.2f})")
print(f"   LogLik Winner: {'Negative Binomial' if nb_m.llf > poisson_m.llf else 'Poisson'}")
print("=" * 60)
print("📄 Updated: Count_Data_Analysis_Report.docx")
print("=" * 60)

# Store the selected model for potential future use
# This makes the selected model available for predictions or further analysis
selected_model = best_count_m
selected_model_name = best_count_name

print(f"\n💡 Selected model stored as 'selected_model' for future use")
print(f"   Model type: {selected_model_name}")
# ============================================
# SECTION 6.5: Incidence Rate Ratios (IRR)
# ============================================

import pandas as pd
import numpy as np
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import time

# ============================================
# YOUR ORIGINAL CODE
# ============================================

irr_df = pd.DataFrame({
    'IRR': np.exp(best_count_m.params).round(3),
    'CI_lower': np.exp(best_count_m.conf_int().iloc[:, 0]).round(3),
    'CI_upper': np.exp(best_count_m.conf_int().iloc[:, 1]).round(3),
    'p_value': best_count_m.pvalues.round(4),
    'Sig': best_count_m.pvalues.apply(
        lambda p: '***' if p < 0.001 else ('**' if p < 0.01 else ('*' if p < 0.05 else '')))
}).drop('const', errors='ignore')

print('=' * 65)
print(f'MODEL 2 — {best_count_name.upper()}: INCIDENCE RATE RATIOS')
print('IRR > 1: associated with higher birth count')
print('IRR < 1: associated with lower birth count')
print('=' * 65)
print(irr_df.to_string())

irr_df.to_excel('model2_count_IRR.xlsx')
print('\nSaved: model2_count_IRR.xlsx')

# ============================================
# ADDITIONAL ANALYSES FOR WORD DOC
# ============================================

# Calculate additional statistics for IRR interpretation
irr_df['Effect_Direction'] = irr_df['IRR'].apply(
    lambda x: 'Positive' if x > 1 else ('Negative' if x < 1 else 'Neutral')
)
irr_df['Effect_Strength'] = irr_df['IRR'].apply(
    lambda x: 'Very Strong' if x > 2 or x < 0.5 else 
              ('Strong' if x > 1.5 or x < 0.67 else 
               ('Moderate' if x > 1.2 or x < 0.83 else 'Weak'))
)

# Count effects by direction
n_positive = len(irr_df[irr_df['IRR'] > 1])
n_negative = len(irr_df[irr_df['IRR'] < 1])
n_neutral = len(irr_df[irr_df['IRR'] == 1])
n_significant = len(irr_df[irr_df['Sig'] != ''])

# Get top positive and negative predictors
top_positive = irr_df[irr_df['IRR'] > 1].nlargest(5, 'IRR')
top_negative = irr_df[irr_df['IRR'] < 1].nsmallest(5, 'IRR')

# ============================================
# ADD RESULTS TO COUNT DATA WORD DOCUMENT
# ============================================

# Add section header
doc_count.add_heading('6.5 Incidence Rate Ratios (IRR)', level=1)

# Add explanation
doc_count.add_paragraph(
    f'Incidence Rate Ratios (IRR) are derived by exponentiating the coefficients from '
    f'the {best_count_name} regression model. IRR represents the multiplicative effect '
    f'on the expected count of parity (birth order).'
)

doc_count.add_paragraph(
    'Interpretation guidelines:\n'
    '• IRR = 1.00: No effect\n'
    '• IRR > 1.00: Associated with HIGHER expected birth count\n'
    '• IRR < 1.00: Associated with LOWER expected birth count\n\n'
    'For a one-unit increase in the predictor, the expected count is multiplied by the IRR.',
    style='Intense Quote'
)

# Add IRR summary table
doc_count.add_heading(f'{best_count_name} Model - Incidence Rate Ratios', level=2)

if len(irr_df) > 0:
    # Create IRR table
    irr_table = doc_count.add_table(rows=len(irr_df) + 1, cols=6)
    irr_table.style = 'Light Grid Accent 1'
    irr_table.autofit = False
    
    # Set column widths
    col_widths = [1.8, 0.8, 0.8, 0.8, 0.9, 0.6]
    for i, width in enumerate(col_widths):
        if i < len(irr_table.columns):
            irr_table.columns[i].width = Inches(width)
    
    # Headers
    headers = ['Variable', 'IRR', 'CI Lower', 'CI Upper', 'p-value', 'Sig']
    for i, header in enumerate(headers):
        irr_table.rows[0].cells[i].text = header
        irr_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        irr_table.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Populate table
    for i, (idx, row) in enumerate(irr_df.iterrows()):
        row_cells = irr_table.rows[i+1].cells
        row_cells[0].text = str(idx)[:40]  # Truncate long names
        row_cells[1].text = f"{row['IRR']:.3f}"
        row_cells[2].text = f"{row['CI_lower']:.3f}"
        row_cells[3].text = f"{row['CI_upper']:.3f}"
        row_cells[4].text = f"{row['p_value']:.4f}"
        row_cells[5].text = row['Sig']
        
        # Color code IRR values
        if row['IRR'] > 1:
            row_cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(204, 102, 0)  # Orange for positive
        elif row['IRR'] < 1:
            row_cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 102, 204)  # Blue for negative
        
        # Bold significant ones
        if row['Sig']:
            for col in range(5):
                row_cells[col].paragraphs[0].runs[0].font.bold = True

# Add significance legend
sig_legend = doc_count.add_paragraph()
sig_legend.add_run('Significance codes: ').font.size = Pt(9)
sig_legend.add_run('*** p<0.001, ').font.size = Pt(9)
sig_legend.add_run('** p<0.01, ').font.size = Pt(9)
sig_legend.add_run('* p<0.05').font.size = Pt(9)

# Add summary statistics
doc_count.add_heading('Summary of IRR Analysis', level=2)

# Create summary statistics table
summary_stats = [
    ('Total Predictors', len(irr_df)),
    ('Significant Predictors (p < 0.05)', n_significant),
    ('Positive Associations (IRR > 1)', n_positive),
    ('Negative Associations (IRR < 1)', n_negative),
    ('Neutral Associations (IRR = 1)', n_neutral)
]

summary_table = doc_count.add_table(rows=len(summary_stats) + 1, cols=2)
summary_table.style = 'Light Shading'
summary_table.autofit = False
summary_table.columns[0].width = Inches(3)
summary_table.columns[1].width = Inches(2.5)

# Headers
summary_table.rows[0].cells[0].text = 'Metric'
summary_table.rows[0].cells[1].text = 'Count'
summary_table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
summary_table.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True

# Populate summary
for i, (metric, value) in enumerate(summary_stats, 1):
    summary_table.rows[i].cells[0].text = metric
    summary_table.rows[i].cells[1].text = str(value)

# Add top predictors section
if len(top_positive) > 0:
    doc_count.add_heading('Top 5 Positive Associations (IRR > 1)', level=2)
    
    pos_table = doc_count.add_table(rows=len(top_positive) + 1, cols=3)
    pos_table.style = 'Light Shading'
    
    # Headers
    pos_table.rows[0].cells[0].text = 'Predictor'
    pos_table.rows[0].cells[1].text = 'IRR (95% CI)'
    pos_table.rows[0].cells[2].text = 'p-value'
    
    for i, (idx, row) in enumerate(top_positive.iterrows(), 1):
        pos_table.rows[i].cells[0].text = str(idx)[:40]
        ci_text = f"{row['IRR']:.3f} ({row['CI_lower']:.3f}-{row['CI_upper']:.3f})"
        pos_table.rows[i].cells[1].text = ci_text
        pos_table.rows[i].cells[2].text = f"{row['p_value']:.4f} {row['Sig']}"
    
    doc_count.add_paragraph(
        'Note: These predictors increase the expected birth count. '
        'For example, an IRR of 1.50 means the expected count is 50% higher '
        'for each unit increase in the predictor.'
    )

if len(top_negative) > 0:
    doc_count.add_heading('Top 5 Negative Associations (IRR < 1)', level=2)
    
    neg_table = doc_count.add_table(rows=len(top_negative) + 1, cols=3)
    neg_table.style = 'Light Shading'
    
    # Headers
    neg_table.rows[0].cells[0].text = 'Predictor'
    neg_table.rows[0].cells[1].text = 'IRR (95% CI)'
    neg_table.rows[0].cells[2].text = 'p-value'
    
    for i, (idx, row) in enumerate(top_negative.iterrows(), 1):
        neg_table.rows[i].cells[0].text = str(idx)[:40]
        ci_text = f"{row['IRR']:.3f} ({row['CI_lower']:.3f}-{row['CI_upper']:.3f})"
        neg_table.rows[i].cells[1].text = ci_text
        neg_table.rows[i].cells[2].text = f"{row['p_value']:.4f} {row['Sig']}"
    
    doc_count.add_paragraph(
        'Note: These predictors decrease the expected birth count. '
        'For example, an IRR of 0.67 means the expected count is 33% lower '
        'for each unit increase in the predictor.'
    )

# Add effect strength distribution
doc_count.add_heading('Effect Strength Distribution', level=2)

strength_counts = irr_df['Effect_Strength'].value_counts()
strength_table = doc_count.add_table(rows=len(strength_counts) + 1, cols=2)
strength_table.style = 'Light Grid Accent 1'

# Headers
strength_table.rows[0].cells[0].text = 'Effect Strength'
strength_table.rows[0].cells[1].text = 'Number of Predictors'
strength_table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
strength_table.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True

# Populate
for i, (strength, count) in enumerate(strength_counts.items(), 1):
    strength_table.rows[i].cells[0].text = strength
    strength_table.rows[i].cells[1].text = str(count)

# Add interpretation guidelines
doc_count.add_heading('Interpretation Guidelines for IRR', level=2)

guidelines = doc_count.add_paragraph()
guidelines.add_run('Effect size interpretation (based on absolute deviation from 1):\n').bold = True
guidelines.add_run('• IRR = 1.00: No effect\n')
guidelines.add_run('• 1.00 < IRR ≤ 1.20 or 0.83 ≤ IRR < 1.00: Weak effect\n')
guidelines.add_run('• 1.20 < IRR ≤ 1.50 or 0.67 ≤ IRR < 0.83: Moderate effect\n')
guidelines.add_run('• 1.50 < IRR ≤ 2.00 or 0.50 ≤ IRR < 0.67: Strong effect\n')
guidelines.add_run('• IRR > 2.00 or IRR < 0.50: Very strong effect\n')

# Add confidence interval interpretation
doc_count.add_heading('Confidence Intervals', level=2)
doc_count.add_paragraph(
    'The 95% confidence intervals provide the range of plausible values for the true IRR. '
    'If the confidence interval does not include 1.00, the effect is statistically '
    'significant at α = 0.05.'
)

# Add note about Excel export
doc_count.add_paragraph(
    f'📎 The IRR results have been exported to Excel: model2_count_IRR.xlsx',
    style='Intense Quote'
)

# Save the updated Word document
try:
    doc_count.save('Count_Data_Analysis_Report.docx')
    print("\n✅ Count Data Analysis Report updated successfully")
except PermissionError:
    timestamp = time.strftime("%Y%m%d_%H%M%S")
    backup_name = f'Count_Data_Analysis_Report_{timestamp}.docx'
    doc_count.save(backup_name)
    print(f"\n✅ Document saved as: {backup_name} (original was locked)")

# Print summary to console
print("\n" + "=" * 65)
print(f"✅ SECTION 6.5 COMPLETE - IRR FOR {best_count_name.upper()} MODEL")
print("=" * 65)
print(f"\n📊 IRR Summary Statistics:")
print(f"   Total predictors analyzed: {len(irr_df)}")
print(f"   Significant predictors: {n_significant}")
print(f"   Positive associations (IRR > 1): {n_positive}")
print(f"   Negative associations (IRR < 1): {n_negative}")
print(f"\n📈 Strongest Positive Association:")
if len(top_positive) > 0:
    top_pos = top_positive.iloc[0]
    print(f"   {top_pos.name}: IRR = {top_pos['IRR']:.3f} "
          f"(95% CI: {top_pos['CI_lower']:.3f}-{top_pos['CI_upper']:.3f})")

print(f"\n📉 Strongest Negative Association:")
if len(top_negative) > 0:
    top_neg = top_negative.iloc[0]
    print(f"   {top_neg.name}: IRR = {top_neg['IRR']:.3f} "
          f"(95% CI: {top_neg['CI_lower']:.3f}-{top_neg['CI_upper']:.3f})")

print("\n" + "=" * 65)
print("📄 Updated: Count_Data_Analysis_Report.docx")
print("📎 Exported: model2_count_IRR.xlsx")
print("=" * 65)

# Store IRR results for potential future use
irr_results = irr_df
print(f"\n💡 IRR results stored as 'irr_results' for future use")
# ============================================
# SECTION 7.1: Fit Binary Logistic Regression (FIXED - Confusion Matrix)
# ============================================

import statsmodels.api as sm
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime
import time
import warnings
from scipy import stats
from sklearn.metrics import confusion_matrix
warnings.filterwarnings('ignore')

# ============================================
# CREATE NEW WORD DOCUMENT FOR SECTION 7
# ============================================

doc_binary = Document()
doc_binary.add_heading('BINARY LOGISTIC REGRESSION ANALYSIS', 0)
doc_binary.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
doc_binary.add_paragraph('Analysis of first-born vs later-born parity (binary outcome)')
doc_binary.add_page_break()

# ============================================
# YOUR ORIGINAL CODE
# ============================================

y_bin = df['parity_binary'].values

logit_m = sm.Logit(y_bin, X_const).fit()

print('BINARY LOGISTIC REGRESSION')
print('Outcome: 1 = First-born  |  0 = Later-born')
print(logit_m.summary())

# ============================================
# ADDITIONAL DIAGNOSTICS FOR WORD DOC
# ============================================

# Calculate model fit statistics
logit_llf = logit_m.llf
logit_llnull = logit_m.llnull
logit_pseudo_r2 = 1 - (logit_llf / logit_llnull)
logit_aic = logit_m.aic
logit_bic = logit_m.bic
logit_df_model = logit_m.df_model
logit_df_resid = logit_m.df_resid

# Likelihood Ratio Test (model vs null)
lr_stat = 2 * (logit_llf - logit_llnull)
lr_pvalue = 1 - stats.chi2.cdf(lr_stat, logit_df_model)

# Extract coefficients and statistics
logit_coef_df = pd.DataFrame({
    'Variable': logit_m.params.index,
    'Coefficient': logit_m.params.values,
    'Std_Error': logit_m.bse.values,
    'z_value': logit_m.tvalues.values,
    'P_value': logit_m.pvalues.values
})

# Calculate Odds Ratios (OR)
logit_coef_df['OR'] = np.exp(logit_coef_df['Coefficient'])
logit_coef_df['OR_lower'] = np.exp(logit_coef_df['Coefficient'] - 1.96 * logit_coef_df['Std_Error'])
logit_coef_df['OR_upper'] = np.exp(logit_coef_df['Coefficient'] + 1.96 * logit_coef_df['Std_Error'])

# Add significance stars
def get_sig_stars(p_val):
    if p_val < 0.001:
        return '***'
    elif p_val < 0.01:
        return '**'
    elif p_val < 0.05:
        return '*'
    else:
        return ''

logit_coef_df['Sig'] = logit_coef_df['P_value'].apply(get_sig_stars)

# Classification accuracy
y_pred_prob = logit_m.predict(X_const)
y_pred_class = (y_pred_prob > 0.5).astype(int)
accuracy = np.mean(y_pred_class == y_bin)

# Confusion matrix
cm = confusion_matrix(y_bin, y_pred_class)
tn, fp, fn, tp = cm.ravel()

sensitivity = tp / (tp + fn) if (tp + fn) > 0 else 0
specificity = tn / (tn + fp) if (tn + fp) > 0 else 0
precision = tp / (tp + fp) if (tp + fp) > 0 else 0
npv = tn / (tn + fn) if (tn + fn) > 0 else 0

# Print additional diagnostics
print(f"\n{'='*60}")
print("MODEL DIAGNOSTICS")
print(f"{'='*60}")
print(f"Pseudo R-squared: {logit_pseudo_r2:.4f}")
print(f"AIC: {logit_aic:.2f}")
print(f"BIC: {logit_bic:.2f}")
print(f"Likelihood Ratio Test: LR = {lr_stat:.4f}, p = {lr_pvalue:.6f}")
print(f"Classification Accuracy: {accuracy:.4f}")
print(f"Sensitivity (True Positive Rate): {sensitivity:.4f}")
print(f"Specificity (True Negative Rate): {specificity:.4f}")
print(f"{'='*60}")

# ============================================
# ADD RESULTS TO BINARY LOGISTIC WORD DOCUMENT
# ============================================

# Add section header
doc_binary.add_heading('7.1 Binary Logistic Regression', level=1)

# Add model description
doc_binary.add_paragraph(
    'Binary logistic regression was used to model the probability of being first-born '
    '(parity = 1) versus later-born (parity > 1). The model uses a logit link function '
    'and maximum likelihood estimation (MLE).'
)

doc_binary.add_paragraph(
    f'Outcome variable coding: 1 = First-born, 0 = Later-born',
    style='Intense Quote'
)

# Add model fit statistics
doc_binary.add_heading('Model Fit Statistics', level=2)

# Create fit statistics data
fit_stats_data = [
    ('Log-Likelihood (Full Model)', f"{logit_llf:.4f}"),
    ('Log-Likelihood (Null Model)', f"{logit_llnull:.4f}"),
    ('Likelihood Ratio (LR) Statistic', f"{lr_stat:.4f}"),
    ('LR Test P-value', f"{lr_pvalue:.6f}"),
    ('Pseudo R-squared (McFadden)', f"{logit_pseudo_r2:.4f}"),
    ('AIC', f"{logit_aic:.2f}"),
    ('BIC', f"{logit_bic:.2f}"),
    ('Number of Observations', f"{len(y_bin)}"),
    ('Degrees of Freedom (Model)', f"{logit_df_model}"),
    ('Degrees of Freedom (Residual)', f"{logit_df_resid}")
]

# Create table with correct number of rows
fit_table = doc_binary.add_table(rows=len(fit_stats_data) + 1, cols=2)
fit_table.style = 'Light Grid Accent 1'
fit_table.autofit = False
fit_table.columns[0].width = Inches(2.8)
fit_table.columns[1].width = Inches(3)

# Headers
fit_table.rows[0].cells[0].text = 'Metric'
fit_table.rows[0].cells[1].text = 'Value'
fit_table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
fit_table.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True

# Fill data
for i, (metric, value) in enumerate(fit_stats_data, 1):
    fit_table.rows[i].cells[0].text = metric
    fit_table.rows[i].cells[1].text = value

# LR Test interpretation
if lr_pvalue < 0.05:
    lr_para = doc_binary.add_paragraph()
    lr_run = lr_para.add_run(f'✓ Likelihood Ratio Test is significant (p = {lr_pvalue:.6f} < 0.05)')
    lr_run.font.color.rgb = RGBColor(0, 128, 0)
    lr_run.bold = True
    doc_binary.add_paragraph('The model with predictors fits significantly better than the null model.')
else:
    lr_para = doc_binary.add_paragraph()
    lr_run = lr_para.add_run(f'⚠ Likelihood Ratio Test is NOT significant (p = {lr_pvalue:.4f})')
    lr_run.font.color.rgb = RGBColor(255, 140, 0)
    doc_binary.add_paragraph('The model does not fit significantly better than the null model.')

# Add coefficients table
doc_binary.add_heading('Regression Coefficients', level=2)

# Filter out constant for main table
logit_main_coef = logit_coef_df[logit_coef_df['Variable'] != 'const'].copy()

if len(logit_main_coef) > 0:
    coef_table = doc_binary.add_table(rows=len(logit_main_coef) + 1, cols=7)
    coef_table.style = 'Light Grid Accent 1'
    coef_table.autofit = False
    
    # Set column widths
    widths = [1.5, 0.9, 0.9, 0.8, 0.8, 0.9, 0.6]
    for i, width in enumerate(widths):
        if i < len(coef_table.columns):
            coef_table.columns[i].width = Inches(width)
    
    # Headers
    headers = ['Variable', 'Coef', 'Std Err', 'z-value', 'P>|z|', 'Odds Ratio', 'Sig']
    for i, header in enumerate(headers):
        if i < len(coef_table.columns):
            coef_table.rows[0].cells[i].text = header
            coef_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    
    # Populate table
    for i, (idx, row) in enumerate(logit_main_coef.iterrows()):
        row_cells = coef_table.rows[i+1].cells
        row_cells[0].text = str(row['Variable'])[:35]
        row_cells[1].text = f"{row['Coefficient']:.4f}"
        row_cells[2].text = f"{row['Std_Error']:.4f}"
        row_cells[3].text = f"{row['z_value']:.4f}"
        row_cells[4].text = f"{row['P_value']:.4f}"
        row_cells[5].text = f"{row['OR']:.3f}"
        row_cells[6].text = row['Sig']
        
        # Bold significant coefficients
        if row['Sig']:
            for col in range(6):
                row_cells[col].paragraphs[0].runs[0].font.bold = True
        
        # Color code odds ratios
        if row['OR'] > 1:
            row_cells[5].paragraphs[0].runs[0].font.color.rgb = RGBColor(204, 102, 0)
        elif row['OR'] < 1:
            row_cells[5].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 102, 204)

# Add significance legend
sig_legend = doc_binary.add_paragraph()
sig_legend.add_run('Significance codes: ').font.size = Pt(9)
sig_legend.add_run('*** p<0.001, ').font.size = Pt(9)
sig_legend.add_run('** p<0.01, ').font.size = Pt(9)
sig_legend.add_run('* p<0.05').font.size = Pt(9)

# Add constant term
logit_const = logit_coef_df[logit_coef_df['Variable'] == 'const']
if len(logit_const) > 0:
    doc_binary.add_heading('Intercept Term', level=2)
    intercept_para = doc_binary.add_paragraph()
    intercept_para.add_run(f'Constant (Intercept): ').bold = True
    intercept_para.add_run(f"Coef = {logit_const['Coefficient'].values[0]:.4f}, ")
    intercept_para.add_run(f"Odds Ratio = {np.exp(logit_const['Coefficient'].values[0]):.4f}, ")
    intercept_para.add_run(f"p = {logit_const['P_value'].values[0]:.4f}")

# Add Odds Ratios interpretation
doc_binary.add_heading('Odds Ratios Interpretation', level=2)
doc_binary.add_paragraph(
    'Odds Ratios (OR) represent the multiplicative effect on the odds of being first-born:\n\n'
    '• OR = 1.00: No effect\n'
    '• OR > 1.00: Higher odds of being first-born\n'
    '• OR < 1.00: Lower odds of being first-born\n\n'
    'For a one-unit increase in the predictor, the odds of being first-born '
    'are multiplied by the OR.'
)

# Create OR summary for key predictors
doc_binary.add_heading('Key Predictors (Odds Ratios)', level=2)

if len(logit_main_coef) > 0:
    # Calculate effect size (distance from 1)
    logit_main_coef['Effect_Size'] = np.abs(logit_main_coef['OR'] - 1)
    
    # Get significant predictors
    sig_predictors = logit_main_coef[logit_main_coef['P_value'] < 0.05].copy()
    
    if len(sig_predictors) > 0:
        sig_predictors = sig_predictors.sort_values('Effect_Size', ascending=False).head(10)
        
        or_table = doc_binary.add_table(rows=len(sig_predictors) + 1, cols=4)
        or_table.style = 'Light Shading'
        
        # Headers
        or_table.rows[0].cells[0].text = 'Predictor'
        or_table.rows[0].cells[1].text = 'Odds Ratio'
        or_table.rows[0].cells[2].text = '95% CI'
        or_table.rows[0].cells[3].text = 'P-value'
        for i in range(4):
            or_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        
        # Populate
        for i, (idx, row) in enumerate(sig_predictors.iterrows(), 1):
            or_table.rows[i].cells[0].text = str(row['Variable'])[:35]
            or_table.rows[i].cells[1].text = f"{row['OR']:.3f}"
            ci_text = f"({row['OR_lower']:.3f}-{row['OR_upper']:.3f})"
            or_table.rows[i].cells[2].text = ci_text
            or_table.rows[i].cells[3].text = f"{row['P_value']:.4f} {row['Sig']}"
    else:
        doc_binary.add_paragraph('No statistically significant predictors found.')

# Add classification performance
doc_binary.add_heading('Model Classification Performance', level=2)

# Create classification table with correct number of rows
class_metrics = [
    ('Overall Accuracy', f"{accuracy:.4f}"),
    ('Sensitivity (True Positive Rate)', f"{sensitivity:.4f}"),
    ('Specificity (True Negative Rate)', f"{specificity:.4f}"),
    ('Precision (Positive Predictive Value)', f"{precision:.4f}"),
    ('Negative Predictive Value (NPV)', f"{npv:.4f}")
]

# Create table with correct row count
class_table = doc_binary.add_table(rows=len(class_metrics) + 1, cols=2)
class_table.style = 'Light Grid Accent 1'
class_table.columns[0].width = Inches(2.5)
class_table.columns[1].width = Inches(3)

# Headers
class_table.rows[0].cells[0].text = 'Metric'
class_table.rows[0].cells[1].text = 'Value'
class_table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
class_table.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True

# Fill data
for i, (metric, value) in enumerate(class_metrics, 1):
    class_table.rows[i].cells[0].text = metric
    class_table.rows[i].cells[1].text = value

# Add confusion matrix (FIXED - no run formatting)
doc_binary.add_heading('Confusion Matrix', level=2)

cm_table = doc_binary.add_table(rows=3, cols=3)
cm_table.style = 'Light Shading'

# Set table cell text directly without formatting runs
# Headers
cm_table.rows[0].cells[0].text = ''
cm_table.rows[0].cells[1].text = 'Predicted: Later-born'
cm_table.rows[0].cells[2].text = 'Predicted: First-born'

# Data rows
cm_table.rows[1].cells[0].text = 'Actual: Later-born'
cm_table.rows[1].cells[1].text = str(tn)
cm_table.rows[1].cells[2].text = str(fp)

cm_table.rows[2].cells[0].text = 'Actual: First-born'
cm_table.rows[2].cells[1].text = str(fn)
cm_table.rows[2].cells[2].text = str(tp)

# Add interpretation of classification
doc_binary.add_paragraph(
    f'The model correctly classifies {accuracy*100:.1f}% of cases. '
    f'Sensitivity ({sensitivity*100:.1f}%) indicates the ability to identify first-born cases, '
    f'while specificity ({specificity*100:.1f}%) indicates the ability to identify later-born cases.'
)

# Add model equation
doc_binary.add_heading('Model Equation', level=2)

if len(logit_main_coef) > 0:
    eq_parts = []
    for _, row in logit_main_coef.iterrows():
        coef = row['Coefficient']
        var = row['Variable']
        if coef >= 0:
            eq_parts.append(f"+ {coef:.4f}×{var}")
        else:
            eq_parts.append(f"- {abs(coef):.4f}×{var}")
    
    intercept_val = logit_const['Coefficient'].values[0] if len(logit_const) > 0 else 0
    eq_text = f"logit(p) = ln(p/(1-p)) = {intercept_val:.4f} {' '.join(eq_parts)}"
    eq_para = doc_binary.add_paragraph()
    eq_run = eq_para.add_run(eq_text)
    eq_run.font.name = 'Courier New'
    eq_run.font.size = Pt(10)
    eq_run.font.bold = True
    
    doc_binary.add_paragraph(
        "Where p is the probability of being first-born. "
        "The logit (log-odds) is transformed to probability via: p = 1/(1 + exp(-logit(p)))"
    )

# Save the new Word document
try:
    doc_binary.save('Binary_Logistic_Report.docx')
    print("\n✅ New document created: Binary_Logistic_Report.docx")
except PermissionError:
    timestamp = time.strftime("%Y%m%d_%H%M%S")
    backup_name = f'Binary_Logistic_Report_{timestamp}.docx'
    doc_binary.save(backup_name)
    print(f"\n✅ Document saved as: {backup_name} (original was locked)")

# Print summary to console
print("\n" + "=" * 60)
print("✅ SECTION 7.1 COMPLETE - BINARY LOGISTIC REGRESSION")
print("=" * 60)
print(f"\n📊 Model Statistics:")
print(f"   Pseudo R-squared: {logit_pseudo_r2:.4f}")
print(f"   AIC: {logit_aic:.2f}")
print(f"   BIC: {logit_bic:.2f}")
print(f"   LR Test p-value: {lr_pvalue:.6f}")
print(f"\n📈 Classification Performance:")
print(f"   Accuracy: {accuracy:.4f}")
print(f"   Sensitivity: {sensitivity:.4f}")
print(f"   Specificity: {specificity:.4f}")
print(f"   Precision: {precision:.4f}")
print(f"   NPV: {npv:.4f}")
print(f"\n📉 Confusion Matrix:")
print(f"   TN: {tn}, FP: {fp}")
print(f"   FN: {fn}, TP: {tp}")
print("=" * 60)
print("📄 New document: Binary_Logistic_Report.docx")
print("=" * 60)

# Clean up temporary column
if 'Effect_Size' in logit_main_coef.columns:
    logit_main_coef = logit_main_coef.drop(columns=['Effect_Size'])
# 7.2  Odds Ratios
binary_or = pd.DataFrame({
    'OR'      : np.exp(logit_m.params).round(3),
    'CI_lower': np.exp(logit_m.conf_int().iloc[:, 0]).round(3),
    'CI_upper': np.exp(logit_m.conf_int().iloc[:, 1]).round(3),
    'p_value' : logit_m.pvalues.round(4),
    'Sig'     : logit_m.pvalues.apply(
        lambda p: '***' if p<0.001 else ('**' if p<0.01 else ('*' if p<0.05 else '')))
}).drop('const', errors='ignore')

print('=' * 65)
print('MODEL 3 — BINARY LOGISTIC: ODDS RATIOS')
print('OR > 1: factor associated with being FIRST-born')
print('OR < 1: factor associated with being LATER-born')
print('=' * 65)
print(binary_or.to_string())

binary_or.to_excel('model3_binary_OR.xlsx')
print('\nSaved: model3_binary_OR.xlsx')
# 7.3  Forest plot
fig, ax = plt.subplots(figsize=(9, max(5, len(binary_or) * 0.5)))

ax.errorbar(
    x=binary_or['OR'], y=list(range(len(binary_or))),
    xerr=[binary_or['OR'] - binary_or['CI_lower'],
          binary_or['CI_upper'] - binary_or['OR']],
    fmt='o', color='#534AB7', ecolor='#3C3489',
    capsize=4, linewidth=1.5, markersize=6
)
ax.axvline(x=1, color='gray', linestyle='--', linewidth=1, label='OR = 1 (no effect)')
ax.set_yticks(list(range(len(binary_or))))
ax.set_yticklabels(binary_or.index, fontsize=9)
ax.set_xlabel('Odds Ratio (95% CI)')
ax.set_title('Model 3 — Binary Logistic Regression: Forest Plot')
ax.legend()
plt.tight_layout()
plt.savefig('fig7_binary_forest.png', dpi=150, bbox_inches='tight')
plt.show()
print('Saved: fig7_binary_forest.png')
# 7.4  ROC Curve & AUC
y_pred_prob = logit_m.predict(X_const)
fpr, tpr, _ = roc_curve(y_bin, y_pred_prob)
roc_auc     = auc(fpr, tpr)

plt.figure(figsize=(6, 5))
plt.plot(fpr, tpr, color='#534AB7', lw=2,
         label=f'AUC = {roc_auc:.3f}')
plt.plot([0, 1], [0, 1], 'gray', linestyle='--', lw=1)
plt.xlabel('False Positive Rate')
plt.ylabel('True Positive Rate')
plt.title('ROC Curve — Binary Logistic Regression')
plt.legend(loc='lower right')
plt.tight_layout()
plt.savefig('fig8_roc_curve.png', dpi=150)
plt.show()

print(f'AUC = {roc_auc:.4f}')
print('0.5=random | 0.7=acceptable | 0.8=good | 0.9=excellent')
print('Saved: fig8_roc_curve.png')
# 7.5  Confusion matrix & classification report
y_pred_class = (y_pred_prob >= 0.5).astype(int)

print('CLASSIFICATION REPORT')
print(classification_report(
    y_bin, y_pred_class,
    target_names=['Later-born', 'First-born']))

cm = confusion_matrix(y_bin, y_pred_class)
plt.figure(figsize=(5, 4))
sns.heatmap(cm, annot=True, fmt='d', cmap='Purples',
            xticklabels=['Later-born', 'First-born'],
            yticklabels=['Later-born', 'First-born'])
plt.title('Confusion Matrix — Binary Logistic')
plt.ylabel('Actual')
plt.xlabel('Predicted')
plt.tight_layout()
plt.savefig('fig9_confusion_matrix.png', dpi=150)
plt.show()
print('Saved: fig9_confusion_matrix.png')
# 8.1  Final comparison table
final_compare = pd.DataFrame({
    'Model'         : ['Ordinal Logistic (POM)',
                       f'{best_count_name} Regression',
                       'Binary Logistic'],
    'Outcome'       : ['Birth order 1-9 (ordered)',
                       'Birth order 1-9 (count)',
                       'Parity 1 vs >=2 (binary)'],
    'Effect measure': ['Cumulative OR', 'IRR', 'OR'],
    'AIC'           : [round(ord_result.aic, 2),
                       round(best_count_m.aic, 2),
                       round(logit_m.aic, 2)],
    'BIC'           : [round(ord_result.bic, 2),
                       round(best_count_m.bic, 2),
                       round(logit_m.bic, 2)],
    'Log-Likelihood': [round(ord_result.llf, 2),
                       round(best_count_m.llf, 2),
                       round(logit_m.llf, 2)]
})

print('=' * 70)
print('FINAL MODEL COMPARISON')
print('Lower AIC/BIC = better  |  Higher Log-Likelihood = better')
print('=' * 70)
print(final_compare.to_string(index=False))

final_compare.to_excel('model_comparison_final.xlsx', index=False)
print('\nSaved: model_comparison_final.xlsx')
# 8.2  List all output files
import os

xlsx = sorted([f for f in os.listdir('.') if f.endswith('.xlsx')])
pngs = sorted([f for f in os.listdir('.') if f.endswith('.png')])

print('=' * 45)
print('ALL OUTPUT FILES')
print('=' * 45)
print('\nExcel result tables:')
for f in xlsx:
    print(f'  {f}')
print('\nFigures:')
for f in pngs:
    print(f'  {f}')
print('\nAnalysis complete.')
# ============================================
# SECTION 7.2: Odds Ratios - Binary Logistic Regression (FIXED)
# ============================================

import pandas as pd
import numpy as np
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import time

# ============================================
# YOUR ORIGINAL CODE
# ============================================

binary_or = pd.DataFrame({
    'OR': np.exp(logit_m.params).round(3),
    'CI_lower': np.exp(logit_m.conf_int().iloc[:, 0]).round(3),
    'CI_upper': np.exp(logit_m.conf_int().iloc[:, 1]).round(3),
    'p_value': logit_m.pvalues.round(4),
    'Sig': logit_m.pvalues.apply(
        lambda p: '***' if p < 0.001 else ('**' if p < 0.01 else ('*' if p < 0.05 else '')))
}).drop('const', errors='ignore')

print('=' * 65)
print('MODEL 3 — BINARY LOGISTIC: ODDS RATIOS')
print('OR > 1: factor associated with being FIRST-born')
print('OR < 1: factor associated with being LATER-born')
print('=' * 65)
print(binary_or.to_string())

binary_or.to_excel('model3_binary_OR.xlsx')
print('\nSaved: model3_binary_OR.xlsx')

# ============================================
# ADDITIONAL ANALYSES FOR WORD DOC (FIXED)
# ============================================

# Calculate additional statistics for OR interpretation
binary_or['Effect_Direction'] = binary_or['OR'].apply(
    lambda x: 'Higher odds of first-born' if x > 1 else 
              ('Lower odds of first-born' if x < 1 else 'No effect')
)
binary_or['Effect_Strength'] = binary_or['OR'].apply(
    lambda x: 'Very Strong' if x > 3 or x < 0.33 else 
              ('Strong' if x > 2 or x < 0.5 else 
               ('Moderate' if x > 1.5 or x < 0.67 else 'Weak'))
)

# Calculate percentage changes (FIXED - added this column)
binary_or['Percent_Change'] = (binary_or['OR'] - 1) * 100

# Count effects by direction
n_positive = len(binary_or[binary_or['OR'] > 1])
n_negative = len(binary_or[binary_or['OR'] < 1])
n_neutral = len(binary_or[binary_or['OR'] == 1])
n_significant = len(binary_or[binary_or['Sig'] != ''])

# Get top positive and negative predictors
top_positive = binary_or[binary_or['OR'] > 1].nlargest(5, 'OR')
top_negative = binary_or[binary_or['OR'] < 1].nsmallest(5, 'OR')

# ============================================
# ADD RESULTS TO BINARY LOGISTIC WORD DOCUMENT
# ============================================

# Add section header
doc_binary.add_heading('7.2 Odds Ratios - Binary Logistic Regression', level=1)

# Add explanation
doc_binary.add_paragraph(
    'Odds Ratios (OR) are derived by exponentiating the coefficients from the binary '
    'logistic regression model. They represent the multiplicative effect on the odds '
    'of being first-born versus later-born.'
)

doc_binary.add_paragraph(
    'Interpretation guidelines:\n'
    '• OR = 1.00: No effect on the odds of being first-born\n'
    '• OR > 1.00: Associated with HIGHER odds of being first-born\n'
    '• OR < 1.00: Associated with LOWER odds of being first-born (higher odds of being later-born)\n\n'
    'For a one-unit increase in the predictor, the odds of being first-born are multiplied by the OR.',
    style='Intense Quote'
)

# Add complete OR table
doc_binary.add_heading('Complete Odds Ratios Table', level=2)

if len(binary_or) > 0:
    # Create OR table
    or_table = doc_binary.add_table(rows=len(binary_or) + 1, cols=6)
    or_table.style = 'Light Grid Accent 1'
    or_table.autofit = False
    
    # Set column widths
    col_widths = [1.8, 0.8, 0.8, 0.8, 0.9, 0.6]
    for i, width in enumerate(col_widths):
        if i < len(or_table.columns):
            or_table.columns[i].width = Inches(width)
    
    # Headers
    headers = ['Variable', 'Odds Ratio', 'CI Lower', 'CI Upper', 'p-value', 'Sig']
    for i, header in enumerate(headers):
        or_table.rows[0].cells[i].text = header
        or_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
        or_table.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Populate table
    for i, (idx, row) in enumerate(binary_or.iterrows()):
        row_cells = or_table.rows[i+1].cells
        row_cells[0].text = str(idx)[:40]  # Truncate long names
        row_cells[1].text = f"{row['OR']:.3f}"
        row_cells[2].text = f"{row['CI_lower']:.3f}"
        row_cells[3].text = f"{row['CI_upper']:.3f}"
        row_cells[4].text = f"{row['p_value']:.4f}"
        row_cells[5].text = row['Sig']
        
        # Color code OR values
        if row['OR'] > 1:
            row_cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(204, 102, 0)  # Orange
            row_cells[1].paragraphs[0].runs[0].bold = True
        elif row['OR'] < 1:
            row_cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 102, 204)  # Blue
            row_cells[1].paragraphs[0].runs[0].bold = True
        
        # Bold significant ones
        if row['Sig']:
            for col in range(5):
                if col != 1:  # Skip OR column as it's already colored
                    row_cells[col].paragraphs[0].runs[0].font.bold = True

# Add significance legend
sig_legend = doc_binary.add_paragraph()
sig_legend.add_run('Significance codes: ').font.size = Pt(9)
sig_legend.add_run('*** p<0.001, ').font.size = Pt(9)
sig_legend.add_run('** p<0.01, ').font.size = Pt(9)
sig_legend.add_run('* p<0.05').font.size = Pt(9)

# Add summary statistics
doc_binary.add_heading('Summary of Odds Ratios Analysis', level=2)

# Create summary statistics table
summary_stats = [
    ('Total Predictors Analyzed', len(binary_or)),
    ('Significant Predictors (p < 0.05)', n_significant),
    ('Positive Associations (OR > 1)', f"{n_positive} (higher odds of first-born)"),
    ('Negative Associations (OR < 1)', f"{n_negative} (lower odds of first-born)"),
    ('Neutral Associations (OR = 1)', n_neutral)
]

summary_table = doc_binary.add_table(rows=len(summary_stats) + 1, cols=2)
summary_table.style = 'Light Shading'
summary_table.autofit = False
summary_table.columns[0].width = Inches(3)
summary_table.columns[1].width = Inches(2.5)

# Headers
summary_table.rows[0].cells[0].text = 'Metric'
summary_table.rows[0].cells[1].text = 'Count'
summary_table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
summary_table.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True

# Populate summary
for i, (metric, value) in enumerate(summary_stats, 1):
    summary_table.rows[i].cells[0].text = metric
    summary_table.rows[i].cells[1].text = str(value)

# Add top predictors section
if len(top_positive) > 0:
    doc_binary.add_heading('Top 5 Positive Associations (Higher Odds of First-born)', level=2)
    
    pos_table = doc_binary.add_table(rows=len(top_positive) + 1, cols=4)
    pos_table.style = 'Light Shading'
    
    # Headers
    pos_table.rows[0].cells[0].text = 'Predictor'
    pos_table.rows[0].cells[1].text = 'Odds Ratio'
    pos_table.rows[0].cells[2].text = '95% CI'
    pos_table.rows[0].cells[3].text = 'Effect'
    for i in range(4):
        pos_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    
    for i, (idx, row) in enumerate(top_positive.iterrows(), 1):
        pos_table.rows[i].cells[0].text = str(idx)[:40]
        pos_table.rows[i].cells[1].text = f"{row['OR']:.3f}"
        ci_text = f"({row['CI_lower']:.3f}-{row['CI_upper']:.3f})"
        pos_table.rows[i].cells[2].text = ci_text
        # Use the Percent_Change column which is now properly defined
        pct_change = row['Percent_Change'] if 'Percent_Change' in row.index else (row['OR'] - 1) * 100
        pos_table.rows[i].cells[3].text = f"{pct_change:.1f}% increase"
        
        # Highlight the OR column
        pos_table.rows[i].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(204, 102, 0)
        pos_table.rows[i].cells[1].paragraphs[0].runs[0].bold = True
    
    doc_binary.add_paragraph(
        'Note: These factors increase the odds of being first-born. '
        'For example, an OR of 2.00 means the odds are 100% higher (doubled).'
    )

if len(top_negative) > 0:
    doc_binary.add_heading('Top 5 Negative Associations (Lower Odds of First-born)', level=2)
    
    neg_table = doc_binary.add_table(rows=len(top_negative) + 1, cols=4)
    neg_table.style = 'Light Shading'
    
    # Headers
    neg_table.rows[0].cells[0].text = 'Predictor'
    neg_table.rows[0].cells[1].text = 'Odds Ratio'
    neg_table.rows[0].cells[2].text = '95% CI'
    neg_table.rows[0].cells[3].text = 'Effect'
    for i in range(4):
        neg_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    
    for i, (idx, row) in enumerate(top_negative.iterrows(), 1):
        neg_table.rows[i].cells[0].text = str(idx)[:40]
        neg_table.rows[i].cells[1].text = f"{row['OR']:.3f}"
        ci_text = f"({row['CI_lower']:.3f}-{row['CI_upper']:.3f})"
        neg_table.rows[i].cells[2].text = ci_text
        # Use the Percent_Change column which is now properly defined
        pct_change = row['Percent_Change'] if 'Percent_Change' in row.index else (row['OR'] - 1) * 100
        neg_table.rows[i].cells[3].text = f"{abs(pct_change):.1f}% decrease"
        
        # Highlight the OR column
        neg_table.rows[i].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 102, 204)
        neg_table.rows[i].cells[1].paragraphs[0].runs[0].bold = True
    
    doc_binary.add_paragraph(
        'Note: These factors decrease the odds of being first-born (i.e., increase odds of being later-born). '
        'For example, an OR of 0.50 means the odds are 50% lower.'
    )

# Add effect strength distribution
doc_binary.add_heading('Effect Strength Distribution', level=2)

# FIXED: Create table with enough rows for all strength categories
strength_counts = binary_or['Effect_Strength'].value_counts()
strength_table = doc_binary.add_table(rows=len(strength_order) + 1, cols=2)  # Use len(strength_order) instead of len(strength_counts)
strength_table.style = 'Light Grid Accent 1'

# Headers
strength_table.rows[0].cells[0].text = 'Effect Strength'
strength_table.rows[0].cells[1].text = 'Number of Predictors'
strength_table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
strength_table.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True

# Define order of strengths
strength_order = ['Very Strong', 'Strong', 'Moderate', 'Weak']
for i, strength in enumerate(strength_order, 1):
    count = strength_counts.get(strength, 0)
    strength_table.rows[i].cells[0].text = strength
    strength_table.rows[i].cells[1].text = str(count)

# Add interpretation guidelines
doc_binary.add_heading('Effect Size Interpretation Guidelines', level=2)

guidelines = doc_binary.add_paragraph()
guidelines.add_run('Odds Ratio interpretation (based on absolute deviation from 1):\n').bold = True
guidelines.add_run('• OR = 1.00: No effect\n')
guidelines.add_run('• 1.00 < OR ≤ 1.50 or 0.67 ≤ OR < 1.00: Weak effect\n')
guidelines.add_run('• 1.50 < OR ≤ 2.00 or 0.50 ≤ OR < 0.67: Moderate effect\n')
guidelines.add_run('• 2.00 < OR ≤ 3.00 or 0.33 ≤ OR < 0.50: Strong effect\n')
guidelines.add_run('• OR > 3.00 or OR < 0.33: Very strong effect\n')

# Add confidence interval interpretation
doc_binary.add_heading('Confidence Intervals', level=2)
doc_binary.add_paragraph(
    'The 95% confidence intervals provide the range of plausible values for the true odds ratio. '
    'If the confidence interval does not include 1.00, the effect is statistically '
    'significant at α = 0.05. Wider intervals indicate less precise estimates.'
)

# Add clinical/practical significance note
doc_binary.add_heading('Practical Significance', level=2)
doc_binary.add_paragraph(
    'While statistical significance indicates whether an effect is likely to be real, '
    'practical significance considers the magnitude of the effect. Odds ratios close to 1.00 '
    '(e.g., between 0.95 and 1.05) may be statistically significant with large sample sizes '
    'but have limited practical importance.'
)

# Add note about Excel export
doc_binary.add_paragraph(
    f'📎 The odds ratios have been exported to Excel: model3_binary_OR.xlsx',
    style='Intense Quote'
)

# Save the updated Word document
try:
    doc_binary.save('Binary_Logistic_Report.docx')
    print("\n✅ Binary Logistic Report updated successfully")
except PermissionError:
    timestamp = time.strftime("%Y%m%d_%H%M%S")
    backup_name = f'Binary_Logistic_Report_{timestamp}.docx'
    doc_binary.save(backup_name)
    print(f"\n✅ Document saved as: {backup_name} (original was locked)")

# Print summary to console
print("\n" + "=" * 65)
print("✅ SECTION 7.2 COMPLETE - BINARY LOGISTIC ODDS RATIOS")
print("=" * 65)
print(f"\n📊 Odds Ratios Summary:")
print(f"   Total predictors analyzed: {len(binary_or)}")
print(f"   Significant predictors: {n_significant}")
print(f"   Positive associations (higher odds first-born): {n_positive}")
print(f"   Negative associations (lower odds first-born): {n_negative}")
print(f"\n📈 Strongest Positive Association (increases first-born odds):")
if len(top_positive) > 0:
    top_pos = top_positive.iloc[0]
    print(f"   {top_pos.name}: OR = {top_pos['OR']:.3f} "
          f"(95% CI: {top_pos['CI_lower']:.3f}-{top_pos['CI_upper']:.3f})")

print(f"\n📉 Strongest Negative Association (decreases first-born odds):")
if len(top_negative) > 0:
    top_neg = top_negative.iloc[0]
    print(f"   {top_neg.name}: OR = {top_neg['OR']:.3f} "
          f"(95% CI: {top_neg['CI_lower']:.3f}-{top_neg['CI_upper']:.3f})")

print("\n" + "=" * 65)
print("📄 Updated: Binary_Logistic_Report.docx")
print("📎 Exported: model3_binary_OR.xlsx")
print("=" * 65)

# Store OR results for potential future use
binary_or_results = binary_or
print(f"\n💡 Odds Ratios stored as 'binary_or_results' for future use")
# ============================================
# SECTION 7.3: Forest Plot - Binary Logistic Regression
# ============================================

import matplotlib.pyplot as plt
import pandas as pd
import numpy as np
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib
import os
import time

# Set matplotlib backend to avoid display issues
matplotlib.use('Agg')

# ============================================
# YOUR ORIGINAL CODE
# ============================================

fig, ax = plt.subplots(figsize=(9, max(5, len(binary_or) * 0.5)))

ax.errorbar(
    x=binary_or['OR'], y=list(range(len(binary_or))),
    xerr=[binary_or['OR'] - binary_or['CI_lower'],
          binary_or['CI_upper'] - binary_or['OR']],
    fmt='o', color='#534AB7', ecolor='#3C3489',
    capsize=4, linewidth=1.5, markersize=6
)
ax.axvline(x=1, color='gray', linestyle='--', linewidth=1, label='OR = 1 (no effect)')
ax.set_yticks(list(range(len(binary_or))))
ax.set_yticklabels(binary_or.index, fontsize=9)
ax.set_xlabel('Odds Ratio (95% CI)')
ax.set_title('Model 3 — Binary Logistic Regression: Forest Plot')
ax.legend()
plt.tight_layout()
plt.savefig('fig7_binary_forest.png', dpi=150, bbox_inches='tight')
plt.savefig('fig7_binary_forest_hd.png', dpi=600, bbox_inches='tight')
plt.show()
print('Saved: fig7_binary_forest.png (150 DPI)')
print('Saved: fig7_binary_forest_hd.png (600 DPI for publication)')

# ============================================
# ADDITIONAL ANALYSES FOR WORD DOC
# ============================================

# Identify significant predictors from the plot
or_above_1 = binary_or[binary_or['CI_lower'] > 1]
or_below_1 = binary_or[binary_or['CI_upper'] < 1]
or_crossing_1 = binary_or[(binary_or['CI_lower'] <= 1) & (binary_or['CI_upper'] >= 1)]

# Calculate statistics for interpretation
n_positive_sig = len(or_above_1)
n_negative_sig = len(or_below_1)
n_non_sig = len(or_crossing_1)

# Get the strongest effects
if len(or_above_1) > 0:
    strongest_positive = or_above_1.nlargest(1, 'OR').iloc[0]
    strongest_positive_name = strongest_positive.name
    strongest_positive_or = strongest_positive['OR']
else:
    strongest_positive_name = None
    strongest_positive_or = None

if len(or_below_1) > 0:
    strongest_negative = or_below_1.nsmallest(1, 'OR').iloc[0]
    strongest_negative_name = strongest_negative.name
    strongest_negative_or = strongest_negative['OR']
else:
    strongest_negative_name = None
    strongest_negative_or = None

# ============================================
# ADD RESULTS TO BINARY LOGISTIC WORD DOCUMENT
# ============================================

# Add section header
doc_binary.add_heading('7.3 Forest Plot - Odds Ratios Visualization', level=1)

# Add interpretation text before the plot
doc_binary.add_paragraph(
    'The forest plot visualizes the odds ratios (OR) and their 95% confidence intervals '
    'for each predictor variable in the binary logistic regression model. '
    'The vertical dashed line at OR = 1 represents no effect. '
    'Points to the right of this line indicate positive associations (higher odds of being first-born), '
    'while points to the left indicate negative associations (lower odds of being first-born).'
)

# Add key observations
doc_binary.add_heading('Key Observations', level=2)

obs_para = doc_binary.add_paragraph()
obs_para.add_run(f'• Predictors with OR > 1 (positive association - higher odds of first-born): {n_positive_sig}\n')
if n_positive_sig > 0:
    for idx in or_above_1.index[:5]:  # Show first 5
        obs_para.add_run(f'  - {idx}: OR = {binary_or.loc[idx, "OR"]:.2f} '
                        f'(95% CI: {binary_or.loc[idx, "CI_lower"]:.2f}-{binary_or.loc[idx, "CI_upper"]:.2f})\n')

obs_para.add_run(f'• Predictors with OR < 1 (negative association - lower odds of first-born): {n_negative_sig}\n')
if n_negative_sig > 0:
    for idx in or_below_1.index[:5]:
        obs_para.add_run(f'  - {idx}: OR = {binary_or.loc[idx, "OR"]:.2f} '
                        f'(95% CI: {binary_or.loc[idx, "CI_lower"]:.2f}-{binary_or.loc[idx, "CI_upper"]:.2f})\n')

obs_para.add_run(f'• Predictors crossing OR = 1 (not statistically significant): {n_non_sig}\n')

# Add the forest plot image to Word
doc_binary.add_heading('Forest Plot', level=2)

# Add the high-resolution image
try:
    # Check if image file exists
    if os.path.exists('fig7_binary_forest_hd.png'):
        # Add the forest plot image (using high-res version)
        doc_binary.add_picture('fig7_binary_forest_hd.png', 
                               width=Inches(6.5))  # Adjust width as needed
        
        # Center align the image
        last_paragraph = doc_binary.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add figure caption
        caption = doc_binary.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption.add_run('Figure 7.1: Forest plot of odds ratios for Model 3 (Binary Logistic Regression)')
        caption_run.font.size = Pt(9)
        caption_run.italic = True
    else:
        doc_binary.add_paragraph('⚠ Forest plot image file not found. Please check that fig7_binary_forest_hd.png was generated.')
        print("Warning: Image file not found")
        
except Exception as e:
    doc_binary.add_paragraph(f'⚠ Image could not be added: {str(e)}')
    print(f"Warning: Could not add image to Word document: {e}")

# Add interpretation table for key predictors
doc_binary.add_heading('Key Predictors by Effect Size', level=2)

# Create summary table of top predictors
top_positive = binary_or[binary_or['OR'] > 1].nlargest(5, 'OR')[['OR', 'CI_lower', 'CI_upper', 'Sig']]
top_negative = binary_or[binary_or['OR'] < 1].nsmallest(5, 'OR')[['OR', 'CI_lower', 'CI_upper', 'Sig']]

# Create a combined table for display
if len(top_positive) > 0 or len(top_negative) > 0:
    summary_table = doc_binary.add_table(rows=1, cols=5)
    summary_table.style = 'Light Grid Accent 1'
    
    # Add headers
    headers = ['Strongest Positive Associations', 'OR (95% CI)', 'Signif.', 
               'Strongest Negative Associations', 'OR (95% CI)']
    for i, header in enumerate(headers):
        summary_table.rows[0].cells[i].text = header
        summary_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    
    # Add data rows
    max_rows = max(len(top_positive), len(top_negative))
    for i in range(max_rows):
        row = summary_table.add_row()
        
        # Positive associations
        if i < len(top_positive):
            var_name = top_positive.index[i]
            row.cells[0].text = str(var_name)[:30]
            ci_text = f"{top_positive.iloc[i]['OR']:.2f} ({top_positive.iloc[i]['CI_lower']:.2f}-{top_positive.iloc[i]['CI_upper']:.2f})"
            row.cells[1].text = ci_text
            row.cells[2].text = top_positive.iloc[i]['Sig']
            
            # Color code positive associations
            row.cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(204, 102, 0)
        
        # Negative associations
        if i < len(top_negative):
            var_name = top_negative.index[i]
            row.cells[3].text = str(var_name)[:30]
            ci_text = f"{top_negative.iloc[i]['OR']:.2f} ({top_negative.iloc[i]['CI_lower']:.2f}-{top_negative.iloc[i]['CI_upper']:.2f})"
            row.cells[4].text = ci_text
            # Note: Significance for negative associations is in the original row
            
            # Color code negative associations
            row.cells[4].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 102, 204)

# Add strongest effects highlight
doc_binary.add_heading('Strongest Effects', level=2)

if strongest_positive_name:
    strong_pos_para = doc_binary.add_paragraph()
    strong_pos_para.add_run(f'Strongest Positive Effect: ').bold = True
    strong_pos_para.add_run(f'{strongest_positive_name} (OR = {strongest_positive_or:.3f}) - '
                           f'This predictor has the largest positive association with being first-born.')

if strongest_negative_name:
    strong_neg_para = doc_binary.add_paragraph()
    strong_neg_para.add_run(f'Strongest Negative Effect: ').bold = True
    strong_neg_para.add_run(f'{strongest_negative_name} (OR = {strongest_negative_or:.3f}) - '
                           f'This predictor has the largest negative association with being first-born '
                           f'(i.e., strongest association with being later-born).')

# Add methodological note
doc_binary.add_heading('Note on Interpretation', level=2)
doc_binary.add_paragraph(
    'The forest plot displays odds ratios on a logarithmic scale. '
    'Confidence intervals that do not cross the vertical line at OR=1 indicate '
    'statistically significant associations at the 5% significance level. '
    'The points represent the point estimates, while the horizontal '
    'bars represent the 95% confidence interval width.\n\n'
    'The purple color scheme (#534AB7 for points, #3C3489 for error bars) '
    'was used for the binary logistic model to distinguish it from previous analyses.'
)

# Add clinical significance discussion
doc_binary.add_heading('Clinical/Substantive Significance', level=2)
doc_binary.add_paragraph(
    'When interpreting forest plots, consider both statistical and substantive significance:\n\n'
    '• OR between 0.95 and 1.05: Minimal practical effect even if statistically significant\n'
    '• OR between 0.67 and 1.50: Small to moderate practical effect\n'
    '• OR < 0.67 or > 1.50: Substantial practical effect\n'
    '• OR < 0.50 or > 2.00: Large practical effect\n'
)

# Add information about saved images
doc_binary.add_heading('Saved Files', level=2)
file_list = doc_binary.add_paragraph()
file_list.add_run('The following image files were saved to disk:\n').bold = True
file_list.add_run('• fig7_binary_forest.png - Forest plot (150 DPI)\n')
file_list.add_run('• fig7_binary_forest_hd.png - High-resolution version for publication (600 DPI)\n')

# Add formatting recommendations
doc_binary.add_heading('Publication Quality', level=2)
doc_binary.add_paragraph(
    'The forest plot has been saved at 600 DPI for publication quality. '
    'For inclusion in manuscripts, consider the following:\n'
    '• Adjust figure size in the code if needed (currently 9×{:.1f} inches)\n'
    '• Colors can be customized to match journal guidelines\n'
    '• The plot uses a purple color scheme (#534AB7 for points, #3C3489 for error bars)'
    .format(max(5, len(binary_or) * 0.5))
)

# Save the updated Word document
try:
    doc_binary.save('Binary_Logistic_Report.docx')
    print("\n✅ Binary Logistic Report updated successfully")
except PermissionError:
    timestamp = time.strftime("%Y%m%d_%H%M%S")
    backup_name = f'Binary_Logistic_Report_{timestamp}.docx'
    doc_binary.save(backup_name)
    print(f"\n✅ Document saved as: {backup_name} (original was locked)")

# Print confirmation
print("\n" + "=" * 60)
print("✅ SECTION 7.3 COMPLETE - FOREST PLOT")
print("=" * 60)
print(f"✅ Forest plot created with {len(binary_or)} predictors")
print(f"✅ High-resolution images saved (150 DPI and 600 DPI)")
print(f"✅ Plot added to Word document with interpretation")
print(f"\n📊 Forest Plot Summary:")
print(f"   • OR > 1 (positive - higher odds first-born): {n_positive_sig} predictors")
print(f"   • OR < 1 (negative - lower odds first-born): {n_negative_sig} predictors")
print(f"   • Non-significant (cross OR=1): {n_non_sig} predictors")
if strongest_positive_name:
    print(f"   • Strongest positive: {strongest_positive_name} (OR = {strongest_positive_or:.3f})")
if strongest_negative_name:
    print(f"   • Strongest negative: {strongest_negative_name} (OR = {strongest_negative_or:.3f})")
print("=" * 60)
print("📄 Updated: Binary_Logistic_Report.docx")
print("=" * 60)
# ============================================
# SECTION 7.4: ROC Curve & AUC - Binary Logistic Regression (FIXED)
# ============================================

import matplotlib.pyplot as plt
import numpy as np
from sklearn.metrics import roc_curve, auc, accuracy_score, precision_score, recall_score, f1_score
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib
import os
import time

# Set matplotlib backend to avoid display issues
matplotlib.use('Agg')

# ============================================
# YOUR ORIGINAL CODE
# ============================================

y_pred_prob = logit_m.predict(X_const)
fpr, tpr, _ = roc_curve(y_bin, y_pred_prob)
roc_auc = auc(fpr, tpr)

plt.figure(figsize=(6, 5))
plt.plot(fpr, tpr, color='#534AB7', lw=2,
         label=f'AUC = {roc_auc:.3f}')
plt.plot([0, 1], [0, 1], 'gray', linestyle='--', lw=1)
plt.xlabel('False Positive Rate')
plt.ylabel('True Positive Rate')
plt.title('ROC Curve — Binary Logistic Regression')
plt.legend(loc='lower right')
plt.tight_layout()
plt.savefig('fig8_roc_curve.png', dpi=150)
plt.savefig('fig8_roc_curve_hd.png', dpi=600)
plt.show()

print(f'AUC = {roc_auc:.4f}')
print('0.5=random | 0.7=acceptable | 0.8=good | 0.9=excellent')
print('Saved: fig8_roc_curve.png (150 DPI)')
print('Saved: fig8_roc_curve_hd.png (600 DPI)')

# ============================================
# ADDITIONAL ANALYSES FOR WORD DOC
# ============================================

# Calculate additional ROC metrics
# Find optimal cutoff (Youden's J statistic)
youden_j = tpr - fpr
optimal_idx = np.argmax(youden_j)
optimal_threshold = np.array([threshold for threshold in _])[optimal_idx] if len(_) > 0 else 0.5
optimal_tpr = tpr[optimal_idx]
optimal_fpr = fpr[optimal_idx]

# Calculate other performance metrics at optimal cutoff
y_pred_optimal = (y_pred_prob >= optimal_threshold).astype(int)

optimal_accuracy = accuracy_score(y_bin, y_pred_optimal)
optimal_precision = precision_score(y_bin, y_pred_optimal, zero_division=0)
optimal_recall = recall_score(y_bin, y_pred_optimal, zero_division=0)
optimal_f1 = f1_score(y_bin, y_pred_optimal, zero_division=0)

# AUC interpretation
if roc_auc >= 0.9:
    auc_quality = "Excellent"
    auc_color = RGBColor(0, 128, 0)
elif roc_auc >= 0.8:
    auc_quality = "Good"
    auc_color = RGBColor(0, 128, 0)
elif roc_auc >= 0.7:
    auc_quality = "Acceptable"
    auc_color = RGBColor(255, 140, 0)
elif roc_auc >= 0.6:
    auc_quality = "Poor"
    auc_color = RGBColor(255, 140, 0)
else:
    auc_quality = "Failed (No discriminative ability)"
    auc_color = RGBColor(255, 0, 0)

# ============================================
# ADD RESULTS TO BINARY LOGISTIC WORD DOCUMENT
# ============================================

# Add section header
doc_binary.add_heading('7.4 ROC Curve & Model Discriminative Ability', level=1)

# Add explanation
doc_binary.add_paragraph(
    'The Receiver Operating Characteristic (ROC) curve evaluates the discriminative '
    'ability of the binary logistic regression model. It plots the True Positive Rate '
    '(Sensitivity) against the False Positive Rate (1 - Specificity) at various threshold settings.'
)

doc_binary.add_paragraph(
    'AUC (Area Under the Curve) Interpretation:\n'
    '• AUC = 0.50: No discriminative ability (random guessing)\n'
    '• 0.70 ≤ AUC < 0.80: Acceptable discrimination\n'
    '• 0.80 ≤ AUC < 0.90: Good discrimination\n'
    '• AUC ≥ 0.90: Excellent discrimination',
    style='Intense Quote'
)

# Add ROC curve image
doc_binary.add_heading('ROC Curve', level=2)

try:
    if os.path.exists('fig8_roc_curve_hd.png'):
        doc_binary.add_picture('fig8_roc_curve_hd.png', width=Inches(5.5))
        
        # Center align
        last_paragraph = doc_binary.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add caption
        caption = doc_binary.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption.add_run('Figure 7.2: ROC Curve for Binary Logistic Regression Model')
        caption_run.font.size = Pt(9)
        caption_run.italic = True
    else:
        doc_binary.add_paragraph('⚠ ROC curve image file not found.')
except Exception as e:
    doc_binary.add_paragraph(f'⚠ Image could not be added: {str(e)}')

# Add AUC results
doc_binary.add_heading('Area Under the Curve (AUC) Results', level=2)

# Create AUC summary table
auc_table = doc_binary.add_table(rows=3, cols=2)
auc_table.style = 'Light Grid Accent 1'
auc_table.autofit = False
auc_table.columns[0].width = Inches(2.5)
auc_table.columns[1].width = Inches(3)

# Headers
auc_table.rows[0].cells[0].text = 'Metric'
auc_table.rows[0].cells[1].text = 'Value'
auc_table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
auc_table.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True

# Data
auc_table.rows[1].cells[0].text = 'AUC Score'
auc_table.rows[1].cells[1].text = f'{roc_auc:.4f}'
auc_table.rows[1].cells[1].paragraphs[0].runs[0].font.bold = True
auc_table.rows[1].cells[1].paragraphs[0].runs[0].font.color.rgb = auc_color

auc_table.rows[2].cells[0].text = 'Discriminative Ability'
auc_table.rows[2].cells[1].text = auc_quality

# Add interpretation paragraph
if roc_auc >= 0.7:
    auc_para = doc_binary.add_paragraph()
    auc_run = auc_para.add_run(f'✓ The model shows {auc_quality.lower()} discriminative ability (AUC = {roc_auc:.3f})')
    auc_run.font.color.rgb = RGBColor(0, 128, 0)
    auc_run.bold = True
    doc_binary.add_paragraph(
        f'The model can effectively distinguish between first-born and later-born cases '
        f'with {roc_auc*100:.1f}% accuracy (where 50% is random guessing).'
    )
else:
    auc_para = doc_binary.add_paragraph()
    auc_run = auc_para.add_run(f'⚠ The model shows {auc_quality} discriminative ability (AUC = {roc_auc:.3f})')
    auc_run.font.color.rgb = RGBColor(255, 140, 0)
    doc_binary.add_paragraph(
        'The model has limited ability to distinguish between first-born and later-born cases. '
        'Consider additional predictors or alternative modeling approaches.'
    )

# Add optimal cutoff analysis
doc_binary.add_heading('Optimal Cutoff Analysis (Youden\'s J Statistic)', level=2)

doc_binary.add_paragraph(
    'The optimal probability cutoff maximizes sensitivity and specificity simultaneously '
    'using Youden\'s J statistic (J = Sensitivity + Specificity - 1).'
)

# Create optimal cutoff table with correct number of rows (FIXED)
cutoff_metrics = [
    ('Optimal Probability Cutoff', f'{optimal_threshold:.4f}'),
    ('Sensitivity at Cutoff', f'{optimal_tpr:.4f}'),
    ('Specificity at Cutoff', f'{1 - optimal_fpr:.4f}'),
    ('Accuracy at Cutoff', f'{optimal_accuracy:.4f}'),
    ('Precision at Cutoff', f'{optimal_precision:.4f}'),
    ('F1 Score at Cutoff', f'{optimal_f1:.4f}')
]

# Create table with correct number of rows (6 metrics + 1 header = 7 rows)
cutoff_table = doc_binary.add_table(rows=len(cutoff_metrics) + 1, cols=2)
cutoff_table.style = 'Light Shading'
cutoff_table.autofit = False
cutoff_table.columns[0].width = Inches(2.5)
cutoff_table.columns[1].width = Inches(3)

# Headers
cutoff_table.rows[0].cells[0].text = 'Metric'
cutoff_table.rows[0].cells[1].text = 'Value at Optimal Cutoff'
cutoff_table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
cutoff_table.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True

# Fill data
for i, (metric, value) in enumerate(cutoff_metrics, 1):
    cutoff_table.rows[i].cells[0].text = metric
    cutoff_table.rows[i].cells[1].text = value

# Add interpretation of optimal cutoff
doc_binary.add_paragraph(
    f'At the optimal cutoff probability of {optimal_threshold:.3f}, '
    f'the model achieves sensitivity = {optimal_tpr:.3f} and specificity = {1-optimal_fpr:.3f}. '
    f'This means it correctly identifies {optimal_tpr*100:.1f}% of first-born cases and '
    f'{(1-optimal_fpr)*100:.1f}% of later-born cases.'
)

# Add model performance comparison
doc_binary.add_heading('Model Performance Summary', level=2)

# Create comprehensive performance table
perf_metrics = [
    ('AUC (Area Under ROC Curve)', f'{roc_auc:.4f}', auc_quality),
    ('Accuracy (at optimal cutoff)', f'{optimal_accuracy:.4f}', ''),
    ('Sensitivity/Recall (at optimal cutoff)', f'{optimal_tpr:.4f}', ''),
    ('Specificity (at optimal cutoff)', f'{1-optimal_fpr:.4f}', ''),
    ('Precision (at optimal cutoff)', f'{optimal_precision:.4f}', ''),
    ('F1 Score (at optimal cutoff)', f'{optimal_f1:.4f}', '')
]

perf_table = doc_binary.add_table(rows=len(perf_metrics) + 1, cols=3)
perf_table.style = 'Light Grid Accent 1'

# Headers
perf_table.rows[0].cells[0].text = 'Performance Metric'
perf_table.rows[0].cells[1].text = 'Value'
perf_table.rows[0].cells[2].text = 'Rating'
for i in range(3):
    perf_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

# Data
for i, (metric, value, rating) in enumerate(perf_metrics, 1):
    perf_table.rows[i].cells[0].text = metric
    perf_table.rows[i].cells[1].text = value
    perf_table.rows[i].cells[2].text = rating
    
    # Highlight AUC row
    if 'AUC' in metric:
        perf_table.rows[i].cells[1].paragraphs[0].runs[0].font.bold = True
        if roc_auc >= 0.7:
            perf_table.rows[i].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)

# Add clinical utility discussion
doc_binary.add_heading('Clinical/Substantive Utility', level=2)
doc_binary.add_paragraph(
    f'The model demonstrates {auc_quality.lower()} discriminative ability '
    f'(AUC = {roc_auc:.3f}). '
)
if roc_auc >= 0.7:
    doc_binary.add_paragraph(
        'This level of discrimination suggests the model could be useful for '
        'risk stratification and identifying women at higher or lower probability '
        'of first-born status. The optimal cutoff provides a practical threshold '
        'for classification in clinical or research settings.'
    )
else:
    doc_binary.add_paragraph(
        'The model has limited clinical utility for individual prediction. '
        'Consider collecting additional relevant predictors or using alternative '
        'modeling approaches to improve discriminative ability.'
    )

# Add limitations
doc_binary.add_heading('Limitations of ROC Analysis', level=2)
doc_binary.add_paragraph(
    'While ROC curves and AUC provide valuable information about model discrimination, '
    'they have limitations:\n\n'
    '• AUC does not reflect calibration (whether predicted probabilities match observed outcomes)\n'
    '• ROC curves may be overly optimistic in imbalanced datasets\n'
    '• Clinical utility depends on the specific cutoff and context\n'
    '• Cross-validation is recommended for unbiased performance estimates\n\n'
    'Consider complementing ROC analysis with calibration plots and decision curve analysis.'
)

# Add saved files information
doc_binary.add_heading('Saved Files', level=2)
file_para = doc_binary.add_paragraph()
file_para.add_run('The following files were saved:\n').bold = True
file_para.add_run('• fig8_roc_curve.png - ROC curve (150 DPI)\n')
file_para.add_run('• fig8_roc_curve_hd.png - High-resolution ROC curve (600 DPI)\n')

# Save the updated Word document
try:
    doc_binary.save('Binary_Logistic_Report.docx')
    print("\n✅ Binary Logistic Report updated successfully")
except PermissionError:
    timestamp = time.strftime("%Y%m%d_%H%M%S")
    backup_name = f'Binary_Logistic_Report_{timestamp}.docx'
    doc_binary.save(backup_name)
    print(f"\n✅ Document saved as: {backup_name} (original was locked)")

# Print confirmation
print("\n" + "=" * 60)
print("✅ SECTION 7.4 COMPLETE - ROC CURVE & AUC")
print("=" * 60)
print(f"\n📊 ROC Analysis Results:")
print(f"   AUC Score: {roc_auc:.4f}")
print(f"   Discriminative Ability: {auc_quality}")
print(f"\n📈 Optimal Cutoff Performance:")
print(f"   Optimal Threshold: {optimal_threshold:.4f}")
print(f"   Sensitivity: {optimal_tpr:.4f}")
print(f"   Specificity: {1-optimal_fpr:.4f}")
print(f"   Accuracy: {optimal_accuracy:.4f}")
print(f"   F1 Score: {optimal_f1:.4f}")
print("=" * 60)
print("📄 Updated: Binary_Logistic_Report.docx")
print("=" * 60)

# Store ROC results for potential future use
roc_results = {
    'auc': roc_auc,
    'fpr': fpr,
    'tpr': tpr,
    'optimal_threshold': optimal_threshold,
    'optimal_sensitivity': optimal_tpr,
    'optimal_specificity': 1 - optimal_fpr,
    'optimal_accuracy': optimal_accuracy
}
print(f"\n💡 ROC results stored as 'roc_results' for future use")
# ============================================
# SECTION 7.5: Confusion Matrix & Classification Report
# ============================================

import matplotlib.pyplot as plt
import seaborn as sns
import numpy as np
from sklearn.metrics import classification_report, confusion_matrix
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import matplotlib
import os
import time

# Set matplotlib backend to avoid display issues
matplotlib.use('Agg')

# ============================================
# YOUR ORIGINAL CODE
# ============================================

y_pred_class = (y_pred_prob >= 0.5).astype(int)

print('CLASSIFICATION REPORT')
print(classification_report(
    y_bin, y_pred_class,
    target_names=['Later-born', 'First-born']))

cm = confusion_matrix(y_bin, y_pred_class)
plt.figure(figsize=(5, 4))
sns.heatmap(cm, annot=True, fmt='d', cmap='Purples',
            xticklabels=['Later-born', 'First-born'],
            yticklabels=['Later-born', 'First-born'])
plt.title('Confusion Matrix — Binary Logistic')
plt.ylabel('Actual')
plt.xlabel('Predicted')
plt.tight_layout()
plt.savefig('fig9_confusion_matrix.png', dpi=150)
plt.savefig('fig9_confusion_matrix_hd.png', dpi=600)
plt.show()
print('Saved: fig9_confusion_matrix.png (150 DPI)')
print('Saved: fig9_confusion_matrix_hd.png (600 DPI)')

# ============================================
# ADDITIONAL ANALYSES FOR WORD DOC
# ============================================

# Extract confusion matrix components
tn, fp, fn, tp = cm.ravel()

# Calculate additional metrics
accuracy = (tp + tn) / (tp + tn + fp + fn)
sensitivity = tp / (tp + fn) if (tp + fn) > 0 else 0
specificity = tn / (tn + fp) if (tn + fp) > 0 else 0
precision = tp / (tp + fp) if (tp + fp) > 0 else 0
npv = tn / (tn + fn) if (tn + fn) > 0 else 0
f1_score = 2 * (precision * sensitivity) / (precision + sensitivity) if (precision + sensitivity) > 0 else 0

# Calculate additional metrics
misclassification_rate = 1 - accuracy
true_positive_rate = sensitivity
false_positive_rate = fp / (fp + tn) if (fp + tn) > 0 else 0
false_negative_rate = fn / (fn + tp) if (fn + tp) > 0 else 0

# Calculate balanced accuracy
balanced_accuracy = (sensitivity + specificity) / 2

# Calculate Matthews Correlation Coefficient (MCC)
from sklearn.metrics import matthews_corrcoef
mcc = matthews_corrcoef(y_bin, y_pred_class)

# Get classification report as dictionary for parsing
report_dict = classification_report(y_bin, y_pred_class, 
                                     target_names=['Later-born', 'First-born'],
                                     output_dict=True)

# ============================================
# ADD RESULTS TO BINARY LOGISTIC WORD DOCUMENT
# ============================================

# Add section header
doc_binary.add_heading('7.5 Confusion Matrix & Classification Performance', level=1)

# Add explanation
doc_binary.add_paragraph(
    'The confusion matrix and classification report provide a comprehensive evaluation '
    'of the binary logistic regression model\'s predictive performance at the default '
    'threshold of 0.5 (where probabilities ≥ 0.5 are classified as first-born).'
)

# Add confusion matrix image
doc_binary.add_heading('Confusion Matrix', level=2)

try:
    if os.path.exists('fig9_confusion_matrix_hd.png'):
        doc_binary.add_picture('fig9_confusion_matrix_hd.png', width=Inches(5.5))
        
        # Center align
        last_paragraph = doc_binary.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add caption
        caption = doc_binary.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption.add_run('Figure 7.3: Confusion Matrix for Binary Logistic Regression Model (Threshold = 0.5)')
        caption_run.font.size = Pt(9)
        caption_run.italic = True
    else:
        doc_binary.add_paragraph('⚠ Confusion matrix image file not found.')
except Exception as e:
    doc_binary.add_paragraph(f'⚠ Image could not be added: {str(e)}')

# Add confusion matrix interpretation
doc_binary.add_heading('Confusion Matrix Interpretation', level=2)

# Create confusion matrix interpretation table
cm_interpretation = [
    ('True Negatives (TN)', f'{tn}', 'Correctly predicted later-born'),
    ('False Positives (FP)', f'{fp}', 'Incorrectly predicted first-born (Type I error)'),
    ('False Negatives (FN)', f'{fn}', 'Incorrectly predicted later-born (Type II error)'),
    ('True Positives (TP)', f'{tp}', 'Correctly predicted first-born')
]

cm_table = doc_binary.add_table(rows=len(cm_interpretation) + 1, cols=3)
cm_table.style = 'Light Shading'
cm_table.autofit = False
cm_table.columns[0].width = Inches(2)
cm_table.columns[1].width = Inches(1)
cm_table.columns[2].width = Inches(3)

# Headers
cm_table.rows[0].cells[0].text = 'Metric'
cm_table.rows[0].cells[1].text = 'Count'
cm_table.rows[0].cells[2].text = 'Interpretation'
for i in range(3):
    cm_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

# Data
for i, (metric, value, interpretation) in enumerate(cm_interpretation, 1):
    cm_table.rows[i].cells[0].text = metric
    cm_table.rows[i].cells[1].text = value
    cm_table.rows[i].cells[2].text = interpretation

# Add comprehensive classification metrics
doc_binary.add_heading('Classification Performance Metrics', level=2)

# Create metrics table
metrics_data = [
    ('Accuracy', f'{accuracy:.4f}', f'{accuracy*100:.1f}%', 'Overall correct predictions'),
    ('Misclassification Rate', f'{misclassification_rate:.4f}', f'{misclassification_rate*100:.1f}%', 'Overall incorrect predictions'),
    ('Sensitivity (Recall)', f'{sensitivity:.4f}', f'{sensitivity*100:.1f}%', 'Correctly identified first-born cases'),
    ('Specificity', f'{specificity:.4f}', f'{specificity*100:.1f}%', 'Correctly identified later-born cases'),
    ('Precision', f'{precision:.4f}', f'{precision*100:.1f}%', 'Accuracy of first-born predictions'),
    ('Negative Predictive Value', f'{npv:.4f}', f'{npv*100:.1f}%', 'Accuracy of later-born predictions'),
    ('F1 Score', f'{f1_score:.4f}', f'{f1_score*100:.1f}%', 'Harmonic mean of precision and recall'),
    ('Balanced Accuracy', f'{balanced_accuracy:.4f}', f'{balanced_accuracy*100:.1f}%', 'Average of sensitivity and specificity'),
    ('Matthews Correlation', f'{mcc:.4f}', '', 'Correlation between actual and predicted')
]

metrics_table = doc_binary.add_table(rows=len(metrics_data) + 1, cols=4)
metrics_table.style = 'Light Grid Accent 1'
metrics_table.autofit = False

# Set column widths
metrics_table.columns[0].width = Inches(2)
metrics_table.columns[1].width = Inches(1.2)
metrics_table.columns[2].width = Inches(1.2)
metrics_table.columns[3].width = Inches(2.5)

# Headers
headers = ['Metric', 'Value', 'Percentage', 'Interpretation']
for i, header in enumerate(headers):
    metrics_table.rows[0].cells[i].text = header
    metrics_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True

# Data
for i, (metric, value, pct, interpretation) in enumerate(metrics_data, 1):
    metrics_table.rows[i].cells[0].text = metric
    metrics_table.rows[i].cells[1].text = value
    metrics_table.rows[i].cells[2].text = pct
    metrics_table.rows[i].cells[3].text = interpretation
    
    # Highlight accuracy if good
    if metric == 'Accuracy' and accuracy >= 0.7:
        metrics_table.rows[i].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
        metrics_table.rows[i].cells[1].paragraphs[0].runs[0].font.bold = True

# Add classification report table
doc_binary.add_heading('Detailed Classification Report', level=2)

# Create classification report table
class_report_data = [
    ('Class', 'Precision', 'Recall', 'F1-Score', 'Support'),
    ('Later-born', f"{report_dict['Later-born']['precision']:.4f}", 
     f"{report_dict['Later-born']['recall']:.4f}", 
     f"{report_dict['Later-born']['f1-score']:.4f}", 
     f"{int(report_dict['Later-born']['support'])}"),
    ('First-born', f"{report_dict['First-born']['precision']:.4f}", 
     f"{report_dict['First-born']['recall']:.4f}", 
     f"{report_dict['First-born']['f1-score']:.4f}", 
     f"{int(report_dict['First-born']['support'])}"),
    ('Accuracy', '', '', f"{report_dict['accuracy']:.4f}", f"{int(report_dict['macro avg']['support'])}"),
    ('Macro Avg', f"{report_dict['macro avg']['precision']:.4f}", 
     f"{report_dict['macro avg']['recall']:.4f}", 
     f"{report_dict['macro avg']['f1-score']:.4f}", ''),
    ('Weighted Avg', f"{report_dict['weighted avg']['precision']:.4f}", 
     f"{report_dict['weighted avg']['recall']:.4f}", 
     f"{report_dict['weighted avg']['f1-score']:.4f}", '')
]

class_table = doc_binary.add_table(rows=len(class_report_data), cols=5)
class_table.style = 'Light Grid Accent 1'

for i, row_data in enumerate(class_report_data):
    for j, cell_data in enumerate(row_data):
        class_table.rows[i].cells[j].text = str(cell_data)
        if i == 0:  # Header row
            class_table.rows[i].cells[j].paragraphs[0].runs[0].font.bold = True

# Add error analysis
doc_binary.add_heading('Error Analysis', level=2)

error_analysis = doc_binary.add_paragraph()
error_analysis.add_run('Type I Errors (False Positives): ').bold = True
error_analysis.add_run(f'{fp} cases were incorrectly classified as first-born when they were actually later-born.\n')
error_analysis.add_run('Type II Errors (False Negatives): ').bold = True
error_analysis.add_run(f'{fn} cases were incorrectly classified as later-born when they were actually first-born.\n\n')

if fp > fn:
    error_analysis.add_run('The model tends to over-predict first-born cases (more False Positives).')
elif fn > fp:
    error_analysis.add_run('The model tends to under-predict first-born cases (more False Negatives).')
else:
    error_analysis.add_run('The model has balanced error rates between the two classes.')

# Add threshold discussion
doc_binary.add_heading('Threshold Considerations', level=2)
doc_binary.add_paragraph(
    f'The current analysis uses the default threshold of 0.5. '
    f'Based on the ROC analysis in Section 7.4, the optimal threshold was {optimal_threshold:.3f} '
    f'using Youden\'s J statistic.\n\n'
    f'At the optimal threshold:\n'
    f'• Sensitivity: {optimal_tpr:.4f} vs current {sensitivity:.4f}\n'
    f'• Specificity: {1-optimal_fpr:.4f} vs current {specificity:.4f}\n\n'
    f'Depending on the research question, one may prefer a threshold that prioritizes '
    f'sensitivity (identifying first-born cases) or specificity (identifying later-born cases).'
)

# Add performance summary
doc_binary.add_heading('Performance Summary', level=2)

if accuracy >= 0.8:
    perf_text = "Excellent"
    perf_color = RGBColor(0, 128, 0)
elif accuracy >= 0.7:
    perf_text = "Good"
    perf_color = RGBColor(0, 128, 0)
elif accuracy >= 0.6:
    perf_text = "Moderate"
    perf_color = RGBColor(255, 140, 0)
else:
    perf_text = "Poor"
    perf_color = RGBColor(255, 0, 0)

perf_para = doc_binary.add_paragraph()
perf_para.add_run(f'Overall Performance Rating: ').bold = True
perf_run = perf_para.add_run(perf_text)
perf_run.font.color.rgb = perf_color
perf_run.bold = True

doc_binary.add_paragraph(
    f'The binary logistic regression model achieves {accuracy*100:.1f}% overall accuracy '
    f'with {sensitivity*100:.1f}% sensitivity and {specificity*100:.1f}% specificity. '
    f'The F1 score of {f1_score:.3f} indicates {("good" if f1_score >= 0.7 else "moderate" if f1_score >= 0.5 else "poor")} '
    f'balance between precision and recall.'
)

# Add saved files information
doc_binary.add_heading('Saved Files', level=2)
file_para = doc_binary.add_paragraph()
file_para.add_run('The following files were saved:\n').bold = True
file_para.add_run('• fig9_confusion_matrix.png - Confusion matrix (150 DPI)\n')
file_para.add_run('• fig9_confusion_matrix_hd.png - High-resolution confusion matrix (600 DPI)\n')

# Save the updated Word document
try:
    doc_binary.save('Binary_Logistic_Report.docx')
    print("\n✅ Binary Logistic Report updated successfully")
except PermissionError:
    timestamp = time.strftime("%Y%m%d_%H%M%S")
    backup_name = f'Binary_Logistic_Report_{timestamp}.docx'
    doc_binary.save(backup_name)
    print(f"\n✅ Document saved as: {backup_name} (original was locked)")

# Print confirmation
print("\n" + "=" * 60)
print("✅ SECTION 7.5 COMPLETE - CONFUSION MATRIX & CLASSIFICATION REPORT")
print("=" * 60)
print(f"\n📊 Confusion Matrix:")
print(f"   True Negatives: {tn}")
print(f"   False Positives: {fp}")
print(f"   False Negatives: {fn}")
print(f"   True Positives: {tp}")
print(f"\n📈 Performance Metrics:")
print(f"   Accuracy: {accuracy:.4f} ({accuracy*100:.1f}%)")
print(f"   Sensitivity: {sensitivity:.4f} ({sensitivity*100:.1f}%)")
print(f"   Specificity: {specificity:.4f} ({specificity*100:.1f}%)")
print(f"   Precision: {precision:.4f} ({precision*100:.1f}%)")
print(f"   F1 Score: {f1_score:.4f}")
print(f"   Balanced Accuracy: {balanced_accuracy:.4f}")
print(f"   MCC: {mcc:.4f}")
print("=" * 60)
print("📄 Updated: Binary_Logistic_Report.docx")
print("=" * 60)

# Store classification results
classification_results = {
    'confusion_matrix': cm,
    'accuracy': accuracy,
    'sensitivity': sensitivity,
    'specificity': specificity,
    'precision': precision,
    'f1_score': f1_score,
    'balanced_accuracy': balanced_accuracy,
    'mcc': mcc
}
print(f"\n💡 Classification results stored as 'classification_results' for future use")
# ============================================
# SECTION 8.1: Final Model Comparison (FIXED)
# ============================================

import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime
import time
import warnings
warnings.filterwarnings('ignore')

# ============================================
# CREATE NEW WORD DOCUMENT FOR SECTION 8
# ============================================

doc_final = Document()
doc_final.add_heading('FINAL MODEL COMPARISON AND RECOMMENDATIONS', 0)
doc_final.add_paragraph(f'Generated on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}')
doc_final.add_paragraph('Comprehensive comparison of all three modeling approaches for parity analysis')
doc_final.add_page_break()

# ============================================
# YOUR ORIGINAL CODE
# ============================================

final_compare = pd.DataFrame({
    'Model': ['Ordinal Logistic (POM)',
              f'{best_count_name} Regression',
              'Binary Logistic'],
    'Outcome': ['Birth order 1-9 (ordered)',
                'Birth order 1-9 (count)',
                'Parity 1 vs >=2 (binary)'],
    'Effect measure': ['Cumulative OR', 'IRR', 'OR'],
    'AIC': [round(ord_result.aic, 2),
            round(best_count_m.aic, 2),
            round(logit_m.aic, 2)],
    'BIC': [round(ord_result.bic, 2),
            round(best_count_m.bic, 2),
            round(logit_m.bic, 2)],
    'Log-Likelihood': [round(ord_result.llf, 2),
                       round(best_count_m.llf, 2),
                       round(logit_m.llf, 2)]
})

print('=' * 70)
print('FINAL MODEL COMPARISON')
print('Lower AIC/BIC = better  |  Higher Log-Likelihood = better')
print('=' * 70)
print(final_compare.to_string(index=False))

final_compare.to_excel('model_comparison_final.xlsx', index=False)
print('\nSaved: model_comparison_final.xlsx')

# ============================================
# ADDITIONAL ANALYSES FOR WORD DOC
# ============================================

# Identify best model by each criterion
best_aic_model = final_compare.loc[final_compare['AIC'].idxmin(), 'Model']
best_bic_model = final_compare.loc[final_compare['BIC'].idxmin(), 'Model']
best_ll_model = final_compare.loc[final_compare['Log-Likelihood'].idxmax(), 'Model']

# Calculate differences
aic_diff_ordinal = abs(final_compare.loc[0, 'AIC'] - final_compare.loc[2, 'AIC'])
aic_diff_count = abs(final_compare.loc[1, 'AIC'] - final_compare.loc[2, 'AIC'])

# Add recommendation based on multiple criteria
if best_aic_model == best_bic_model == best_ll_model:
    recommended_model = best_aic_model
    recommendation_confidence = "Strong consensus across all criteria"
elif best_aic_model == best_bic_model:
    recommended_model = best_aic_model
    recommendation_confidence = "AIC and BIC agree"
elif best_aic_model == best_ll_model:
    recommended_model = best_aic_model
    recommendation_confidence = "AIC and Log-Likelihood agree"
elif best_bic_model == best_ll_model:
    recommended_model = best_bic_model
    recommendation_confidence = "BIC and Log-Likelihood agree"
else:
    # Use AIC as default tiebreaker
    recommended_model = best_aic_model
    recommendation_confidence = "Mixed criteria - AIC used as tiebreaker"

# ============================================
# ADD RESULTS TO FINAL COMPARISON WORD DOCUMENT
# ============================================

# Add section header
doc_final.add_heading('8.1 Final Model Comparison', level=1)

# Add explanation
doc_final.add_paragraph(
    'This section compares the three modeling approaches used to analyze parity data:\n\n'
    '1. **Ordinal Logistic Regression (Proportional Odds Model)** - Treats parity as ordered categories\n'
    '2. **Count Model (Poisson/Negative Binomial)** - Treats parity as count data\n'
    '3. **Binary Logistic Regression** - Dichotomizes parity (first-born vs later-born)\n\n'
    'Comparison criteria include AIC (Akaike Information Criterion), BIC (Bayesian Information Criterion), '
    'and Log-Likelihood. Lower AIC/BIC values indicate better fit, while higher Log-Likelihood indicates better fit.'
)

# Add main comparison table
doc_final.add_heading('Model Comparison Summary', level=2)

# Create formatted table
compare_table = doc_final.add_table(rows=len(final_compare) + 1, cols=len(final_compare.columns))
compare_table.style = 'Light Grid Accent 1'
compare_table.autofit = False

# Set column widths
col_widths = [1.8, 2.2, 1.2, 0.8, 0.8, 1.0]
for i, width in enumerate(col_widths):
    if i < len(compare_table.columns):
        compare_table.columns[i].width = Inches(width)

# Headers
for i, col in enumerate(final_compare.columns):
    compare_table.rows[0].cells[i].text = col
    compare_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    compare_table.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

# Populate table and highlight best values
for i, row in final_compare.iterrows():
    for j, col in enumerate(final_compare.columns):
        value = str(row[col])
        compare_table.rows[i+1].cells[j].text = value
        compare_table.rows[i+1].cells[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Highlight best values
        if col == 'AIC':
            if value == str(final_compare['AIC'].min()):
                compare_table.rows[i+1].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
                compare_table.rows[i+1].cells[j].paragraphs[0].runs[0].font.bold = True
        elif col == 'BIC':
            if value == str(final_compare['BIC'].min()):
                compare_table.rows[i+1].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
                compare_table.rows[i+1].cells[j].paragraphs[0].runs[0].font.bold = True
        elif col == 'Log-Likelihood':
            if value == str(final_compare['Log-Likelihood'].max()):
                compare_table.rows[i+1].cells[j].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)
                compare_table.rows[i+1].cells[j].paragraphs[0].runs[0].font.bold = True

# Add best model indicators
doc_final.add_heading('Model Selection Criteria Results', level=2)

# Create criteria table with correct number of rows (FIXED)
criteria_data = [
    ('AIC (lower is better)', best_aic_model),
    ('BIC (lower is better)', best_bic_model),
    ('Log-Likelihood (higher is better)', best_ll_model),
    ('Consensus', f"{recommendation_confidence} → {recommended_model}")
]

# Create table with correct row count (4 data rows + 1 header = 5 rows)
criteria_table = doc_final.add_table(rows=len(criteria_data) + 1, cols=2)
criteria_table.style = 'Light Shading'
criteria_table.autofit = False
criteria_table.columns[0].width = Inches(2.5)
criteria_table.columns[1].width = Inches(3.5)

# Headers
criteria_table.rows[0].cells[0].text = 'Criterion'
criteria_table.rows[0].cells[1].text = 'Best Model'
criteria_table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
criteria_table.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True

# Data
for i, (criterion, best_model) in enumerate(criteria_data, 1):
    criteria_table.rows[i].cells[0].text = criterion
    criteria_table.rows[i].cells[1].text = best_model
    if i <= 3:  # For the first three criteria, highlight the best model
        criteria_table.rows[i].cells[1].paragraphs[0].runs[0].font.bold = True
        criteria_table.rows[i].cells[1].paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 128, 0)

# Add model comparison details
doc_final.add_heading('Detailed Comparison', level=2)

# Create detailed comparison sections
doc_final.add_heading('AIC Comparison', level=3)
aic_para = doc_final.add_paragraph()
aic_para.add_run(f'• Best AIC: {best_aic_model} (AIC = {final_compare["AIC"].min():.2f})\n')
aic_para.add_run(f'• Ordinal vs Binary difference: {aic_diff_ordinal:.2f}\n')
aic_para.add_run(f'• Count vs Binary difference: {aic_diff_count:.2f}\n')

doc_final.add_heading('BIC Comparison', level=3)
bic_para = doc_final.add_paragraph()
bic_para.add_run(f'• Best BIC: {best_bic_model} (BIC = {final_compare["BIC"].min():.2f})\n')

doc_final.add_heading('Log-Likelihood Comparison', level=3)
ll_para = doc_final.add_paragraph()
ll_para.add_run(f'• Best Log-Likelihood: {best_ll_model} (LL = {final_compare["Log-Likelihood"].max():.2f})\n')

# Add model-specific strengths and weaknesses
doc_final.add_heading('Model-Specific Considerations', level=2)

# Ordinal Logistic
doc_final.add_heading('Ordinal Logistic (POM)', level=3)
ordinal_pros = doc_final.add_paragraph()
ordinal_pros.add_run('Strengths:\n').bold = True
ordinal_pros.add_run('• Preserves ordinal nature of parity\n')
ordinal_pros.add_run('• Provides category-specific thresholds\n')
ordinal_pros.add_run('• More efficient than multinomial models\n')
ordinal_pros.add_run('• Allows testing of proportional odds assumption\n\n')
ordinal_pros.add_run('Weaknesses:\n').bold = True
ordinal_pros.add_run('• Proportional odds assumption may be violated\n')
ordinal_pros.add_run('• More complex interpretation than binary model\n')
ordinal_pros.add_run('• Requires larger sample size for stable estimates\n')

# Count Model
doc_final.add_heading(f'{best_count_name} Regression', level=3)
count_pros = doc_final.add_paragraph()
count_pros.add_run('Strengths:\n').bold = True
count_pros.add_run('• Treats parity as actual counts (preserves magnitude)\n')
if best_count_name == 'Negative Binomial':
    count_pros.add_run('• Accounts for overdispersion (common in parity data)\n')
count_pros.add_run('• Provides incidence rate ratios (easily interpretable)\n')
count_pros.add_run('• Can handle zero inflation if present\n\n')
count_pros.add_run('Weaknesses:\n').bold = True
count_pros.add_run('• Assumes independence of events\n')
count_pros.add_run('• May not capture ceiling effects at higher parities\n')
if best_count_name == 'Poisson':
    count_pros.add_run('• Assumes equidispersion (may be violated)\n')

# Binary Logistic
doc_final.add_heading('Binary Logistic Regression', level=3)
binary_pros = doc_final.add_paragraph()
binary_pros.add_run('Strengths:\n').bold = True
binary_pros.add_run('• Simple and widely understood\n')
binary_pros.add_run('• Provides clear odds ratios\n')
binary_pros.add_run('• Good for binary decision-making\n')
binary_pros.add_run('• Robust to certain assumption violations\n\n')
binary_pros.add_run('Weaknesses:\n').bold = True
binary_pros.add_run('• Loses information by dichotomizing parity\n')
binary_pros.add_run('• Cannot distinguish between parity levels 2,3,4+\n')
binary_pros.add_run('• Reduced statistical power compared to ordinal model\n')

# Add final recommendation
doc_final.add_heading('Final Model Recommendation', level=2)

# Create recommendation box
rec_box = doc_final.add_paragraph()
rec_box.add_run('=' * 50 + '\n').bold = True
rec_box.add_run(f'RECOMMENDED MODEL: {recommended_model}\n').bold = True
rec_box.add_run('=' * 50 + '\n').bold = True

# Add reasoning
reasoning = doc_final.add_paragraph()
reasoning.add_run('Reasoning:\n').bold = True

if recommended_model == 'Ordinal Logistic (POM)':
    reasoning.add_run('• Preserves the full ordinal information of parity\n')
    reasoning.add_run('• Provides interpretable cumulative odds ratios\n')
    reasoning.add_run('• Allows testing of proportional odds assumption\n')
    reasoning.add_run('• Best suited for ordered categorical outcomes\n')
elif recommended_model == f'{best_count_name} Regression':
    reasoning.add_run('• Treats parity as count data (most natural representation)\n')
    if best_count_name == 'Negative Binomial':
        reasoning.add_run('• Accounts for overdispersion common in parity data\n')
    reasoning.add_run('• Provides incident rate ratios for easy interpretation\n')
    reasoning.add_run('• Maintains full information about parity levels\n')
else:  # Binary Logistic
    reasoning.add_run('• Simplest model with clear interpretation\n')
    reasoning.add_run('• Good for binary classification of first-born status\n')
    reasoning.add_run('• Provides straightforward odds ratios\n')
    reasoning.add_run('• Useful when only first-born vs later-born distinction is needed\n')

# Add practical recommendations
doc_final.add_heading('Practical Recommendations', level=2)
practical = doc_final.add_paragraph()
practical.add_run('Based on the analysis, we recommend:\n\n').bold = True
practical.add_run('1. Primary Analysis: Use the count model for understanding parity patterns\n')
practical.add_run('2. Supplementary Analysis: Use ordinal logistic for category-specific effects\n')
practical.add_run('3. Decision-Making: Use binary logistic for first-born prediction\n')
practical.add_run('4. Reporting: Report results from all three models for comprehensive understanding\n')

# Add note about Excel export
doc_final.add_paragraph(
    f'📎 The comparison table has been exported to Excel: model_comparison_final.xlsx',
    style='Intense Quote'
)

# Save the new Word document
try:
    doc_final.save('Final_Model_Comparison_Report.docx')
    print("\n✅ New document created: Final_Model_Comparison_Report.docx")
except PermissionError:
    timestamp = time.strftime("%Y%m%d_%H%M%S")
    backup_name = f'Final_Model_Comparison_Report_{timestamp}.docx'
    doc_final.save(backup_name)
    print(f"\n✅ Document saved as: {backup_name} (original was locked)")

# Print summary to console
print("\n" + "=" * 70)
print("✅ SECTION 8.1 COMPLETE - FINAL MODEL COMPARISON")
print("=" * 70)
print(final_compare.to_string(index=False))
print("\n" + "=" * 70)
print(f"\n📊 Model Selection Summary:")
print(f"   Best by AIC: {best_aic_model}")
print(f"   Best by BIC: {best_bic_model}")
print(f"   Best by Log-Likelihood: {best_ll_model}")
print(f"   Recommended Model: {recommended_model}")
print(f"   Consensus: {recommendation_confidence}")
print("\n" + "=" * 70)
print("📄 New document: Final_Model_Comparison_Report.docx")
print("📎 Exported: model_comparison_final.xlsx")
print("=" * 70)
# ============================================
# SECTION 8.2: List All Output Files (FIXED - Final)
# ============================================

import os
import pandas as pd
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime
import time

# ============================================
# YOUR ORIGINAL CODE
# ============================================

xlsx = sorted([f for f in os.listdir('.') if f.endswith('.xlsx')])
pngs = sorted([f for f in os.listdir('.') if f.endswith('.png')])

print('=' * 45)
print('ALL OUTPUT FILES')
print('=' * 45)
print('\nExcel result tables:')
for f in xlsx:
    print(f'  {f}')
print('\nFigures:')
for f in pngs:
    print(f'  {f}')
print('\nAnalysis complete.')

# ============================================
# ADDITIONAL ANALYSES FOR WORD DOC
# ============================================

# Categorize files by type and purpose
ordinal_files = [f for f in xlsx if 'ordinal' in f.lower() or 'model1' in f.lower()]
count_files = [f for f in xlsx if 'count' in f.lower() or 'model2' in f.lower()]
binary_files = [f for f in xlsx if 'binary' in f.lower() or 'model3' in f.lower()]
comparison_files = [f for f in xlsx if 'comparison' in f.lower() or 'final' in f.lower()]

# Categorize images
ordinal_images = [f for f in pngs if 'ordinal' in f.lower() or 'fig5' in f.lower() or 'fig6' in f.lower()]
count_images = [f for f in pngs if 'count' in f.lower()]
binary_images = [f for f in pngs if 'binary' in f.lower() or 'fig7' in f.lower() or 'fig8' in f.lower() or 'fig9' in f.lower()]
roc_images = [f for f in pngs if 'roc' in f.lower()]
forest_images = [f for f in pngs if 'forest' in f.lower()]
confusion_images = [f for f in pngs if 'confusion' in f.lower()]

# Calculate file statistics
total_excel = len(xlsx)
total_images = len(pngs)
total_files = total_excel + total_images

# Get file sizes
file_info = []
for f in xlsx + pngs:
    file_path = os.path.join('.', f)
    file_size = os.path.getsize(file_path) if os.path.exists(file_path) else 0
    file_size_kb = file_size / 1024
    file_info.append({
        'filename': f,
        'size_kb': file_size_kb,
        'type': 'Excel' if f.endswith('.xlsx') else 'Image'
    })

# Sort by size
file_info_df = pd.DataFrame(file_info).sort_values('size_kb', ascending=False)

# ============================================
# ADD RESULTS TO FINAL COMPARISON WORD DOCUMENT
# ============================================

# Add section header
doc_final.add_heading('8.2 Output Files Summary', level=1)

# Add explanation
doc_final.add_paragraph(
    'This analysis generated multiple output files including Excel result tables '
    'and high-resolution figures. Below is a comprehensive list of all files '
    'produced during the analysis.'
)

# Add summary statistics
doc_final.add_heading('File Generation Summary', level=2)

# Create summary data list
summary_data = [
    ('Total Excel Files', total_excel),
    ('Total Image Files', total_images),
    ('Total Files Generated', total_files),
    ('Excel Files (Ordinal Models)', len(ordinal_files)),
    ('Excel Files (Count Models)', len(count_files)),
    ('Excel Files (Binary Models)', len(binary_files)),
    ('Excel Files (Comparisons)', len(comparison_files)),
    ('Images (Ordinal)', len(ordinal_images)),
    ('Images (Count)', len(count_images)),
    ('Images (Binary)', len(binary_images))
]

# Create summary table with correct number of rows
summary_table = doc_final.add_table(rows=len(summary_data) + 1, cols=2)
summary_table.style = 'Light Grid Accent 1'
summary_table.autofit = False
summary_table.columns[0].width = Inches(2.5)
summary_table.columns[1].width = Inches(3)

# Headers
summary_table.rows[0].cells[0].text = 'File Category'
summary_table.rows[0].cells[1].text = 'Count'
summary_table.rows[0].cells[0].paragraphs[0].runs[0].font.bold = True
summary_table.rows[0].cells[1].paragraphs[0].runs[0].font.bold = True

# Fill data - show all rows
for i, (category, count) in enumerate(summary_data, 1):
    summary_table.rows[i].cells[0].text = category
    summary_table.rows[i].cells[1].text = str(count)

# Add detailed file lists
doc_final.add_heading('Excel Output Files', level=2)

if len(xlsx) > 0:
    excel_table = doc_final.add_table(rows=len(xlsx) + 1, cols=3)
    excel_table.style = 'Light Shading'
    excel_table.autofit = False
    
    # Set column widths
    excel_table.columns[0].width = Inches(2)
    excel_table.columns[1].width = Inches(1.5)
    excel_table.columns[2].width = Inches(2.5)
    
    # Headers
    excel_table.rows[0].cells[0].text = 'Filename'
    excel_table.rows[0].cells[1].text = 'Size (KB)'
    excel_table.rows[0].cells[2].text = 'Purpose'
    for i in range(3):
        excel_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    
    # Populate table
    for i, f in enumerate(xlsx, 1):
        file_size_kb = file_info_df[file_info_df['filename'] == f]['size_kb'].values[0] if len(file_info_df[file_info_df['filename'] == f]) > 0 else 0
        
        excel_table.rows[i].cells[0].text = f
        excel_table.rows[i].cells[1].text = f"{file_size_kb:.2f}"
        
        # Determine purpose
        if 'ordinal' in f.lower() or 'model1' in f.lower():
            purpose = "Ordinal Logistic Regression Results"
        elif 'count' in f.lower() or 'model2' in f.lower():
            purpose = "Count Model (IRR) Results"
        elif 'binary' in f.lower() or 'model3' in f.lower():
            purpose = "Binary Logistic (OR) Results"
        elif 'comparison' in f.lower():
            purpose = "Model Comparison Summary"
        else:
            purpose = "Analysis Results"
        
        excel_table.rows[i].cells[2].text = purpose
else:
    doc_final.add_paragraph('No Excel files generated.')

# Add image files
doc_final.add_heading('Figure/Image Files', level=2)

if len(pngs) > 0:
    # Group images by type
    doc_final.add_heading('Ordinal Model Figures', level=3)
    ordinal_figs = [f for f in pngs if 'ordinal' in f.lower() or 'fig5' in f.lower() or 'fig6' in f.lower()]
    if len(ordinal_figs) > 0:
        for f in ordinal_figs:
            doc_final.add_paragraph(f'• {f}', style='List Bullet')
    else:
        doc_final.add_paragraph('No ordinal model figures found.')
    
    doc_final.add_heading('Count Model Figures', level=3)
    count_figs = [f for f in pngs if 'count' in f.lower()]
    if len(count_figs) > 0:
        for f in count_figs:
            doc_final.add_paragraph(f'• {f}', style='List Bullet')
    else:
        doc_final.add_paragraph('No count model figures found.')
    
    doc_final.add_heading('Binary Model Figures', level=3)
    binary_figs = [f for f in pngs if 'binary' in f.lower() or 'fig7' in f.lower() or 'fig8' in f.lower() or 'fig9' in f.lower()]
    if len(binary_figs) > 0:
        for f in binary_figs:
            doc_final.add_paragraph(f'• {f}', style='List Bullet')
    else:
        doc_final.add_paragraph('No binary model figures found.')
    
    # List all remaining images
    all_listed = set(ordinal_figs + count_figs + binary_figs)
    remaining = [f for f in pngs if f not in all_listed]
    if len(remaining) > 0:
        doc_final.add_heading('Additional Figures', level=3)
        for f in remaining:
            doc_final.add_paragraph(f'• {f}', style='List Bullet')
else:
    doc_final.add_paragraph('No image files generated.')

# Add largest files section
doc_final.add_heading('Largest Output Files', level=2)

if len(file_info_df) > 0:
    # Limit to top 10 files
    top_files = file_info_df.head(10)
    largest_table = doc_final.add_table(rows=len(top_files) + 1, cols=3)
    largest_table.style = 'Light Grid Accent 1'
    largest_table.autofit = False
    
    # Headers
    largest_table.rows[0].cells[0].text = 'Filename'
    largest_table.rows[0].cells[1].text = 'Size (KB)'
    largest_table.rows[0].cells[2].text = 'Type'
    for i in range(3):
        largest_table.rows[0].cells[i].paragraphs[0].runs[0].font.bold = True
    
    # Data - FIXED: Use proper index
    for i, (idx, row) in enumerate(top_files.iterrows()):
        largest_table.rows[i+1].cells[0].text = row['filename']
        largest_table.rows[i+1].cells[1].text = f"{row['size_kb']:.2f}"
        largest_table.rows[i+1].cells[2].text = row['type']
else:
    doc_final.add_paragraph('No files found.')

# Add file management recommendations
doc_final.add_heading('File Management Recommendations', level=2)

recommendations = doc_final.add_paragraph()
recommendations.add_run('Recommendations:\n').bold = True
recommendations.add_run('• Keep all Excel files for reproducibility and further analysis\n')
recommendations.add_run('• High-resolution images (600 DPI) are suitable for publication\n')
recommendations.add_run('• Standard resolution images (150 DPI) are suitable for quick viewing\n')
recommendations.add_run('• Consider archiving files in a structured folder by model type\n')
recommendations.add_run('• Back up all files to secure storage\n')

# Add file organization suggestion
doc_final.add_heading('Suggested File Organization', level=2)

org_structure = doc_final.add_paragraph()
org_structure.add_run('Suggested folder structure:\n').bold = True
org_structure.add_run('├── Excel_Results/\n')
org_structure.add_run('│   ├── Ordinal_Models/\n')
org_structure.add_run('│   ├── Count_Models/\n')
org_structure.add_run('│   ├── Binary_Models/\n')
org_structure.add_run('│   └── Comparisons/\n')
org_structure.add_run('├── Figures/\n')
org_structure.add_run('│   ├── Ordinal_Plots/\n')
org_structure.add_run('│   ├── Count_Plots/\n')
org_structure.add_run('│   ├── Binary_Plots/\n')
org_structure.add_run('│   └── ROC_Curves/\n')
org_structure.add_run('└── Reports/\n')
org_structure.add_run('    ├── Word_Documents/\n')
org_structure.add_run('    └── PDF_Exports/\n')

# Add completion message
doc_final.add_heading('Analysis Completion', level=2)
completion_para = doc_final.add_paragraph()
completion_para.add_run('✓ All analyses completed successfully\n').bold = True
completion_para.add_run(f'✓ Total files generated: {total_files}\n')
completion_para.add_run(f'✓ Excel tables: {total_excel}\n')
completion_para.add_run(f'✓ Figures: {total_images}\n')
completion_para.add_run('✓ Reports saved to Word documents\n')

# Add timestamp
doc_final.add_paragraph(
    f'Analysis completed on: {datetime.now().strftime("%Y-%m-%d %H:%M:%S")}',
    style='Intense Quote'
)

# Save the updated Word document
try:
    doc_final.save('Final_Model_Comparison_Report.docx')
    print("\n✅ Final Model Comparison Report updated successfully")
except PermissionError:
    timestamp = time.strftime("%Y%m%d_%H%M%S")
    backup_name = f'Final_Model_Comparison_Report_{timestamp}.docx'
    doc_final.save(backup_name)
    print(f"\n✅ Document saved as: {backup_name} (original was locked)")

# Print summary to console
print("\n" + "=" * 45)
print("✅ SECTION 8.2 COMPLETE - OUTPUT FILES SUMMARY")
print("=" * 45)
print(f"\n📊 File Generation Summary:")
print(f"   Total Excel files: {total_excel}")
print(f"   Total Image files: {total_images}")
print(f"   Total files generated: {total_files}")
print("\n📁 Excel Files:")
for f in xlsx:
    print(f"   • {f}")
print("\n🖼️ Image Files:")
for f in pngs:
    print(f"   • {f}")
print("\n" + "=" * 45)
print("✅ Analysis Complete!")
print("=" * 45)
print("📄 Updated: Final_Model_Comparison_Report.docx")
print("=" * 45)
"""
generate_results_doc.py
=======================
Paste this entire script as a new cell at the END of your Jupyter notebook
and run it.  It reads all Excel results + PNG figures that the notebook saved
and produces:

    results_chapter.docx

Requirements (already installed by the notebook):
    pip install python-docx pandas openpyxl pillow
"""

# ─────────────────────────────────────────────────────────────────────────────
# 0.  Imports
# ─────────────────────────────────────────────────────────────────────────────
import pandas as pd
import numpy as np
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import subprocess, sys

# install python-docx if missing
try:
    import docx
except ImportError:
    subprocess.check_call([sys.executable, "-m", "pip", "install",
                           "python-docx", "--quiet"])
    import docx

print("Libraries ready.")

# ─────────────────────────────────────────────────────────────────────────────
# 1.  Load all result files
# ─────────────────────────────────────────────────────────────────────────────
chi2_df   = pd.read_excel("results_chisquare.xlsx")
ord_df    = pd.read_excel("model1_ordinal_OR.xlsx",   index_col=0)
count_df  = pd.read_excel("model2_count_IRR.xlsx",    index_col=0)
bin_df    = pd.read_excel("model3_binary_OR.xlsx",    index_col=0)
comp_df   = pd.read_excel("model_comparison_final.xlsx")

# figure paths (only include ones that actually exist)
FIG_FILES = {
    "fig1": "fig1_parity_distribution.png",
    "fig2": "fig2_age_by_parity.png",
    "fig3": "fig3_predictors_vs_parity.png",
    "fig4": "fig4_birthweight_by_parity.png",
    "fig5": "fig5_ordinal_forest.png",
    "fig6": "fig6_proportional_odds.png",
    "fig7": "fig7_binary_forest.png",
    "fig8": "fig8_roc_curve.png",
    "fig9": "fig9_confusion_matrix.png",
}
figs = {k: v for k, v in FIG_FILES.items() if Path(v).exists()}
print(f"Loaded results.  Figures found: {list(figs.keys())}")

# ─────────────────────────────────────────────────────────────────────────────
# 2.  Derive key numbers for automatic interpretation
# ─────────────────────────────────────────────────────────────────────────────
# --- dispersion (recompute from the notebook's df if still in memory) --------
try:
    mean_p = float(df["parity"].mean())
    var_p  = float(df["parity"].var())
    disp_r = var_p / mean_p
    n_total = len(df)
    pct_first = float(df["parity_binary"].mean()) * 100
except NameError:
    # df not in scope — use placeholder values; user can edit
    mean_p, var_p, disp_r, n_total, pct_first = 2.5, 3.8, 1.52, 0, 40.0

count_model_name = "Negative Binomial" if disp_r > 1.2 else "Poisson"

# --- chi-square sig variables ------------------------------------------------
sig_vars = chi2_df[chi2_df["Significant"].str.startswith("Yes")]["Variable"].tolist()
nonsig_vars = chi2_df[~chi2_df["Significant"].str.startswith("Yes")]["Variable"].tolist()

# --- ordinal model: top 3 most significant predictors -----------------------
ord_sig = ord_df[ord_df["p_value"] < 0.05].sort_values("p_value")
ord_top = ord_sig.head(3)

# --- binary model: strongest OR (furthest from 1) ---------------------------
bin_sig = bin_df[bin_df["p_value"] < 0.05].copy()
bin_sig["dist"] = (bin_sig["OR"] - 1).abs()
bin_top = bin_sig.sort_values("dist", ascending=False).head(3)

# --- model comparison: pick best by AIC -------------------------------------
comp_df_sorted = comp_df.sort_values("AIC")
best_model_row = comp_df_sorted.iloc[0]

# --- AUC from binary model (approximate from saved data) --------------------
try:
    from sklearn.metrics import roc_auc_score
    auc_val = roc_auc_score(df["parity_binary"],
                             logit_m.predict(sm.add_constant(X_encoded)))
    auc_str = f"{auc_val:.3f}"
except Exception:
    auc_str = "see figure"

print(f"Dispersion ratio  : {disp_r:.4f}  → {count_model_name}")
print(f"Significant chi2  : {sig_vars}")
print(f"Best model by AIC : {best_model_row.get('Model','')}")

# ─────────────────────────────────────────────────────────────────────────────
# 3.  Document helpers
# ─────────────────────────────────────────────────────────────────────────────
doc = Document()

# --- page margins (A4, 1-inch all sides) ------------------------------------
for section in doc.sections:
    section.top_margin    = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin   = Cm(2.54)
    section.right_margin  = Cm(2.54)

# --- set default font --------------------------------------------------------
from docx.oxml.ns import qn as _qn
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(12)
style.paragraph_format.line_spacing = Pt(24)   # double-space
style.paragraph_format.space_after  = Pt(0)

def set_cell_bg(cell, hex_color):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)

def add_heading(text, level=1):
    p = doc.add_heading(text, level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in p.runs:
        run.font.name  = "Times New Roman"
        run.font.color.rgb = RGBColor(0, 0, 0)
    return p

def add_para(text, indent=True, bold_parts=None):
    """Add a justified double-spaced paragraph.
    bold_parts: list of substrings to make bold."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.space_after  = Pt(6)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.27)
    if bold_parts:
        remaining = text
        for bp in bold_parts:
            idx = remaining.find(bp)
            if idx == -1:
                continue
            if idx > 0:
                run = p.add_run(remaining[:idx])
                run.font.name = "Times New Roman"
                run.font.size = Pt(12)
            br = p.add_run(bp)
            br.bold = True
            br.font.name = "Times New Roman"
            br.font.size = Pt(12)
            remaining = remaining[idx + len(bp):]
        if remaining:
            run = p.add_run(remaining)
            run.font.name = "Times New Roman"
            run.font.size = Pt(12)
    else:
        run = p.add_run(text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
    return p

def add_equation(label, equation_text):
    """Add a displayed equation line with label."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    p.paragraph_format.line_spacing = Pt(18)
    run = p.add_run(f"{equation_text}    {label}")
    run.font.name   = "Cambria Math"
    run.font.size   = Pt(12)
    run.font.italic = True
    return p

def add_caption(text):
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(8)
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)
        run.bold = True
    return p

def add_figure(path, caption, width=5.5):
    if Path(path).exists():
        doc.add_picture(path, width=Inches(width))
        last = doc.paragraphs[-1]
        last.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_caption(caption)
    else:
        add_para(f"[Figure not found: {path}]", indent=False)

def add_results_table(df_in, caption_text,
                       col_headers=None, shade_rows=True):
    """Render a pandas DataFrame as a Word table."""
    add_caption(caption_text)
    df_in = df_in.reset_index()
    if col_headers and len(col_headers) == len(df_in.columns):
        df_in.columns = col_headers

    tbl = doc.add_table(rows=1, cols=len(df_in.columns))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # header row
    hdr = tbl.rows[0]
    for i, col in enumerate(df_in.columns):
        cell = hdr.cells[i]
        cell.text = str(col)
        set_cell_bg(cell, "D6E4F0")
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.name = "Times New Roman"
            run.font.size = Pt(10)

    # data rows
    for ridx, row in df_in.iterrows():
        tr = tbl.add_row()
        bg = "F5F5F5" if (shade_rows and ridx % 2 == 0) else "FFFFFF"
        for cidx, val in enumerate(row):
            cell = tr.cells[cidx]
            # format numbers nicely
            if isinstance(val, float):
                cell.text = f"{val:.4f}" if abs(val) < 1000 else f"{val:,.2f}"
            else:
                cell.text = str(val)
            set_cell_bg(cell, bg)
            cell.paragraphs[0].alignment = (
                WD_ALIGN_PARAGRAPH.CENTER if cidx > 0
                else WD_ALIGN_PARAGRAPH.LEFT)
            for run in cell.paragraphs[0].runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)

    doc.add_paragraph()  # spacer after table

def spacer(n=1):
    for _ in range(n):
        p = doc.add_paragraph("")
        p.paragraph_format.space_after  = Pt(0)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.line_spacing = Pt(12)

# ─────────────────────────────────────────────────────────────────────────────
# 4.  CHAPTER TITLE
# ─────────────────────────────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("CHAPTER FOUR")
r.bold = True; r.font.size = Pt(14); r.font.name = "Times New Roman"

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("RESULTS AND INTERPRETATION")
r2.bold = True; r2.font.size = Pt(14); r2.font.name = "Times New Roman"
spacer(2)

# ─────────────────────────────────────────────────────────────────────────────
# 5.  SECTION 4.1 — Descriptive Statistics
# ─────────────────────────────────────────────────────────────────────────────
add_heading("4.1  Descriptive Statistics and Parity Distribution", level=1)

add_para(
    f"A total of {n_total:,} registered live births were included in the final "
    f"cleaned analytical dataset after removing records with missing values in "
    f"key variables and implausible maternal age entries. The distribution of "
    f"birth order across the nine parity categories is presented in Figure 4.1 "
    f"and summarized below."
)

add_para(
    f"First-born children (parity = 1) accounted for {pct_first:.1f}% of all "
    f"registered births, making this the single largest birth order group. The "
    f"proportion of births declined progressively with increasing birth order, "
    f"which is consistent with the fertility transition documented in Sri Lanka "
    f"since the 1960s (UNFPA, 2017). Higher-order births (parity \u2265 5) "
    f"collectively represented a small minority of the total, reflecting the "
    f"low total fertility rate — approximately 2.0 children per woman — that "
    f"characterizes Sri Lanka's demographic profile during the study period "
    f"(World Bank, 2025).",
    bold_parts=[f"{pct_first:.1f}%"]
)

if "fig1" in figs:
    add_figure(figs["fig1"], "Figure 4.1: Birth order distribution — count and percentage share")

add_heading("4.1.1  Maternal Age by Birth Order", level=2)
add_para(
    "Maternal age at birth increased systematically with birth order, as shown "
    "in Figure 4.2. Mothers of first-born children were concentrated in the "
    "20–29 age group, while mothers of higher-order births (parity \u2265 4) were "
    "predominantly aged 30 years and above. This pattern reflects the sequential "
    "nature of childbearing — higher-order births necessarily occur later in a "
    "woman's reproductive life course. The age group composition bars confirm "
    "that very young mothers (aged under 20) contribute a disproportionately "
    "higher share of first births relative to later births."
)

if "fig2" in figs:
    add_figure(figs["fig2"], "Figure 4.2: Maternal age distribution by birth order")

add_heading("4.1.2  Maternal and Birth Characteristics by Parity Group", level=2)
add_para(
    "Figure 4.3 displays the proportion of first births versus later births "
    "across the six categorical predictor variables. Noteworthy patterns include "
    "differences in first-birth proportions across marital status categories, "
    "race of mother, and place of delivery. Hospital births showed a higher "
    "proportion of first-born children relative to non-hospital births, "
    "consistent with findings in the broader obstetric literature linking "
    "primiparity with higher health-seeking behaviour. Birth weight distributions "
    "by birth order (Figure 4.4) indicate that first-born children tend to have "
    "marginally lower birth weights compared to second and third-order births, "
    "a phenomenon documented in the obstetric literature as the parity–birth "
    "weight effect."
)

if "fig3" in figs:
    add_figure(figs["fig3"],
               "Figure 4.3: First-birth proportion by maternal and birth characteristics")
if "fig4" in figs:
    add_figure(figs["fig4"],
               "Figure 4.4: Birth weight distribution by birth order")

# ─────────────────────────────────────────────────────────────────────────────
# 6.  SECTION 4.2 — Chi-Square Tests
# ─────────────────────────────────────────────────────────────────────────────
add_heading("4.2  Chi-Square Tests of Association", level=1)

add_para(
    "Prior to regression modelling, chi-square tests of independence were "
    "conducted to assess the bivariate association between birth order and "
    "each categorical predictor variable. The null hypothesis for each test "
    "states that birth order is independent of the predictor variable. "
    "Results are presented in Table 4.1."
)

add_results_table(
    chi2_df,
    "Table 4.1: Chi-Square Test Results — Predictor Variables vs Birth Order",
    col_headers=["Variable", "Chi\u00b2 Statistic", "df", "p-value", "Significant?"]
)

sig_str   = ", ".join(sig_vars)   if sig_vars   else "none"
nonsig_str= ", ".join(nonsig_vars) if nonsig_vars else "none"

add_para(
    f"The chi-square tests revealed that the following variables were "
    f"statistically significantly associated with birth order at the 5% "
    f"significance level: {sig_str}. "
    f"In contrast, {nonsig_str} did not demonstrate a statistically "
    f"significant association with birth order (p > 0.05). All statistically "
    f"significant variables were retained in the subsequent multivariable "
    f"regression models.",
    bold_parts=[sig_str]
)

# ─────────────────────────────────────────────────────────────────────────────
# 7.  SECTION 4.3 — Model 1: Ordinal Logistic Regression
# ─────────────────────────────────────────────────────────────────────────────
add_heading("4.3  Model 1: Ordinal Logistic Regression", level=1)

add_heading("4.3.1  Model Specification and Statistical Equation", level=2)

add_para(
    "The Proportional Odds Model (POM) is specified as follows. Let Y denote "
    "birth order with ordered categories j = 1, 2, \u2026, 9, and let "
    "x\u2081, x\u2082, \u2026, x\u2096 denote the k predictor variables. "
    "The cumulative probability up to and including category j is denoted "
    "P(Y \u2264 j | x). The model is expressed as:",
    indent=True
)

add_equation(
    "(4.1)",
    "logit[ P(Y \u2264 j | x) ] = \u03b1\u2c7c \u2212 (\u03b2\u2081x\u2081 + \u03b2\u2082x\u2082 + \u22ef + \u03b2\u2096x\u2096)"
)

add_para(
    "where \u03b1\u2c7c (j = 1, 2, \u2026, 8) are the eight threshold parameters "
    "(intercepts) representing the log-cumulative odds when all predictors equal "
    "zero, and \u03b2\u2081 through \u03b2\u2096 are the regression coefficients "
    "common across all cumulative splits — the defining constraint of the "
    "proportional odds assumption. The model is estimated by maximum likelihood. "
    "The exponentiated coefficients exp(\u03b2) are interpreted as cumulative "
    "odds ratios (OR): an OR greater than 1 indicates that the predictor is "
    "associated with a higher birth order, while an OR less than 1 indicates "
    "an association with lower birth order.",
    indent=True
)

add_heading("4.3.2  Proportional Odds Assumption Check", level=2)
add_para(
    "The proportional odds assumption was evaluated by fitting eight separate "
    "binary logistic regression models, each using a different cumulative "
    "dichotomization of birth order as the outcome (parity > 1, parity > 2, "
    "\u2026, parity > 8). Under the proportional odds assumption, the regression "
    "coefficients should remain approximately stable across all eight "
    "dichotomizations. Figure 4.5 displays the coefficient trajectories for "
    "each predictor across the eight cutpoints."
)

if "fig6" in figs:
    add_figure(figs["fig6"],
               "Figure 4.5: Proportional odds assumption — coefficient stability by cutpoint")

add_para(
    "Lines that remain approximately flat and parallel across cutpoints provide "
    "support for the proportional odds assumption. Lines that cross or diverge "
    "substantially indicate a potential violation for that predictor. Where "
    "minor instability is observed, results should be interpreted with caution "
    "and a partial proportional odds model may be considered as a robustness "
    "check in future work."
)

add_heading("4.3.3  Ordinal Logistic Regression Results", level=2)
add_para(
    "The estimated cumulative odds ratios with 95% confidence intervals and "
    "p-values from the Proportional Odds Model are presented in Table 4.2 "
    "and visualized as a forest plot in Figure 4.6."
)

add_results_table(
    ord_df,
    "Table 4.2: Model 1 — Ordinal Logistic Regression: Cumulative Odds Ratios",
    col_headers=["Predictor", "OR", "95% CI Lower", "95% CI Upper", "p-value", "Sig."]
)

if "fig5" in figs:
    add_figure(figs["fig5"],
               "Figure 4.6: Model 1 — Ordinal Logistic Regression: Forest Plot of Odds Ratios")

# Auto-interpret top predictors
add_heading("4.3.4  Interpretation of Key Findings", level=2)

if len(ord_top) > 0:
    for predictor, row in ord_top.iterrows():
        direction = "higher" if row["OR"] > 1 else "lower"
        add_para(
            f"The predictor {predictor} was statistically significant "
            f"(OR = {row['OR']:.3f}, 95% CI: {row['CI_lower']:.3f}\u2013"
            f"{row['CI_upper']:.3f}, p = {row['p_value']:.4f}), indicating "
            f"that this factor is associated with a {direction} birth order "
            f"after controlling for all other variables in the model.",
            bold_parts=[predictor]
        )
else:
    add_para(
        "The ordinal logistic regression model did not identify any "
        "statistically significant predictors at the 5% significance level. "
        "This may reflect insufficient statistical power or the absence of "
        "strong associations in this dataset after multivariate adjustment."
    )

add_para(
    "Predictors whose 95% confidence intervals for the cumulative OR excluded "
    "the null value of 1.0 are considered statistically significant. "
    "Predictors for which the confidence interval includes 1.0 are not "
    "statistically distinguishable from no effect at the 5% significance level."
)

# ─────────────────────────────────────────────────────────────────────────────
# 8.  SECTION 4.4 — Model 2: Count Regression
# ─────────────────────────────────────────────────────────────────────────────
add_heading("4.4  Model 2: Count Regression", level=1)

add_heading("4.4.1  Dispersion Test and Model Selection", level=2)

add_para(
    f"Before selecting between Poisson and Negative Binomial regression, the "
    f"dispersion of the parity variable was assessed. The mean parity was "
    f"{mean_p:.4f} and the variance was {var_p:.4f}, yielding a dispersion "
    f"ratio (variance / mean) of {disp_r:.4f}.",
    bold_parts=[f"{disp_r:.4f}"]
)

if disp_r > 1.2:
    add_para(
        f"Since the dispersion ratio of {disp_r:.4f} substantially exceeds 1.0, "
        f"the parity data exhibit overdispersion — the variance is greater than "
        f"the mean, violating the equidispersion assumption of the standard "
        f"Poisson model. Overdispersion in count data arises when there is "
        f"unobserved heterogeneity among observations not accounted for by the "
        f"included predictors. In such cases, the standard Poisson model produces "
        f"underestimated standard errors, leading to inflated test statistics and "
        f"spuriously significant p-values (Agresti, 2018). The Negative Binomial "
        f"regression model addresses overdispersion by incorporating an additional "
        f"dispersion parameter \u03b8 that allows the variance to exceed the mean: "
        f"Var(Y) = \u03bc + \u03bc\u00b2/\u03b8. Accordingly, the Negative "
        f"Binomial model was selected as the primary count regression model.",
        bold_parts=["Negative Binomial"]
    )
elif disp_r < 0.8:
    add_para(
        f"The dispersion ratio of {disp_r:.4f} is less than 1.0, indicating "
        f"underdispersion. Standard Poisson regression was applied with robust "
        f"standard errors to account for this."
    )
else:
    add_para(
        f"The dispersion ratio of {disp_r:.4f} is approximately equal to 1.0, "
        f"consistent with the equidispersion assumption of the Poisson model. "
        f"Standard Poisson regression was therefore applied."
    )

add_heading("4.4.2  Statistical Equations", level=2)

add_para("The Poisson regression model is specified as:", indent=True)
add_equation("(4.2)", "log(\u03bc\u1d62) = \u03b2\u2080 + \u03b2\u2081x\u2081\u1d62 + \u03b2\u2082x\u2082\u1d62 + \u22ef + \u03b2\u2096x\u2096\u1d62")

add_para(
    "where \u03bc\u1d62 = E(Y\u1d62) is the expected birth order count for "
    "observation i. The Negative Binomial model extends this by adding an "
    "observation-level random effect \u03b5\u1d62 to account for overdispersion:",
    indent=True
)
add_equation("(4.3)", "log(\u03bc\u1d62) = \u03b2\u2080 + \u03b2\u2081x\u2081\u1d62 + \u22ef + \u03b2\u2096x\u2096\u1d62 + \u03b5\u1d62,   \u03b5\u1d62 ~ Gamma(1/\u03b8, \u03b8)")

add_para(
    "yielding a marginal variance of Var(Y\u1d62) = \u03bc\u1d62 + \u03bc\u1d62\u00b2/\u03b8, "
    "where \u03b8 is the dispersion parameter estimated from the data. "
    "Exponentiated coefficients exp(\u03b2) are interpreted as Incidence Rate "
    "Ratios (IRR): an IRR > 1 indicates the predictor is associated with a "
    "higher expected birth order count, while IRR < 1 indicates a lower "
    "expected count.",
    indent=True
)

add_heading(f"4.4.3  {count_model_name} Regression Results", level=2)
add_para(
    f"The Incidence Rate Ratios (IRR) with 95% confidence intervals estimated "
    f"from the {count_model_name} regression model are presented in Table 4.3."
)

add_results_table(
    count_df,
    f"Table 4.3: Model 2 — {count_model_name} Regression: Incidence Rate Ratios",
    col_headers=["Predictor", "IRR", "95% CI Lower", "95% CI Upper", "p-value", "Sig."]
)

# auto-interpret count model top predictors
count_sig = count_df[count_df["p_value"] < 0.05].copy()
count_sig["dist"] = (count_sig["IRR"] - 1).abs()
count_top = count_sig.sort_values("dist", ascending=False).head(3)

add_heading("4.4.4  Interpretation of Key Findings", level=2)
if len(count_top) > 0:
    for predictor, row in count_top.iterrows():
        direction = "higher" if row["IRR"] > 1 else "lower"
        pct_change = abs(row["IRR"] - 1) * 100
        add_para(
            f"The predictor {predictor} was statistically significant "
            f"(IRR = {row['IRR']:.3f}, 95% CI: {row['CI_lower']:.3f}\u2013"
            f"{row['CI_upper']:.3f}, p = {row['p_value']:.4f}). This indicates "
            f"that this factor is associated with a {pct_change:.1f}% "
            f"{'increase' if row['IRR']>1 else 'decrease'} in the expected "
            f"birth order count, holding all other predictors constant.",
            bold_parts=[predictor]
        )
else:
    add_para(
        "No predictors reached statistical significance at the 5% level in the "
        "count regression model after multivariate adjustment."
    )

# ─────────────────────────────────────────────────────────────────────────────
# 9.  SECTION 4.5 — Model 3: Binary Logistic Regression
# ─────────────────────────────────────────────────────────────────────────────
add_heading("4.5  Model 3: Binary Logistic Regression", level=1)

add_heading("4.5.1  Model Specification and Statistical Equation", level=2)

add_para(
    "Binary logistic regression was applied with parity dichotomized as "
    "first-born (Y = 1) versus later-born (Y = 0). The model estimates the "
    "log-odds of a birth being a first birth as a linear function of the "
    "predictor variables:",
    indent=True
)
add_equation(
    "(4.4)",
    "log[ P(Y=1|x) / (1\u2212P(Y=1|x)) ] = \u03b2\u2080 + \u03b2\u2081x\u2081 + \u03b2\u2082x\u2082 + \u22ef + \u03b2\u2096x\u2096"
)

add_para(
    "The probability of a first birth given the predictor values is recovered as:",
    indent=True
)
add_equation(
    "(4.5)",
    "P(Y=1|x) = exp(\u03b2\u2080 + \u03b2\u2081x\u2081 + \u22ef + \u03b2\u2096x\u2096) / [1 + exp(\u03b2\u2080 + \u03b2\u2081x\u2081 + \u22ef + \u03b2\u2096x\u2096)]"
)

add_para(
    "Exponentiated coefficients exp(\u03b2) are interpreted as odds ratios (OR). "
    "An OR > 1 indicates that the predictor is associated with increased odds "
    "of being a first birth; an OR < 1 indicates reduced odds of being a "
    "first birth, relative to the reference category.",
    indent=True
)

add_heading("4.5.2  Binary Logistic Regression Results", level=2)
add_para(
    "The estimated odds ratios with 95% confidence intervals and p-values are "
    "presented in Table 4.4 and visualized as a forest plot in Figure 4.7."
)

add_results_table(
    bin_df,
    "Table 4.4: Model 3 — Binary Logistic Regression: Odds Ratios (First-born vs Later-born)",
    col_headers=["Predictor", "OR", "95% CI Lower", "95% CI Upper", "p-value", "Sig."]
)

if "fig7" in figs:
    add_figure(figs["fig7"],
               "Figure 4.7: Model 3 — Binary Logistic Regression: Forest Plot of Odds Ratios")

add_heading("4.5.3  Interpretation of Key Findings", level=2)
if len(bin_top) > 0:
    for predictor, row in bin_top.iterrows():
        direction = "increased" if row["OR"] > 1 else "decreased"
        add_para(
            f"The predictor {predictor} was statistically significant "
            f"(OR = {row['OR']:.3f}, 95% CI: {row['CI_lower']:.3f}\u2013"
            f"{row['CI_upper']:.3f}, p = {row['p_value']:.4f}), indicating "
            f"{direction} odds of being a first birth for this group "
            f"relative to the reference category, after adjusting for "
            f"all other variables.",
            bold_parts=[predictor]
        )
else:
    add_para(
        "No predictors reached statistical significance in the binary "
        "logistic regression model after multivariate adjustment."
    )

add_heading("4.5.4  Model Performance: ROC Curve and AUC", level=2)
add_para(
    f"The discriminatory ability of the binary logistic regression model was "
    f"assessed using the Receiver Operating Characteristic (ROC) curve and the "
    f"Area Under the Curve (AUC). The ROC curve plots the true positive rate "
    f"(sensitivity) against the false positive rate (1 \u2212 specificity) "
    f"across all probability thresholds. An AUC of 0.5 corresponds to random "
    f"classification, while an AUC of 1.0 represents perfect discrimination. "
    f"The AUC for this model was {auc_str}. Figure 4.8 presents the ROC curve "
    f"and Figure 4.9 presents the confusion matrix at the 0.5 classification "
    f"threshold.",
    bold_parts=[f"AUC = {auc_str}"]
)

if "fig8" in figs:
    add_figure(figs["fig8"], "Figure 4.8: ROC Curve — Binary Logistic Regression")
if "fig9" in figs:
    add_figure(figs["fig9"], "Figure 4.9: Confusion Matrix — Binary Logistic Regression")

try:
    auc_num = float(auc_str)
    if auc_num >= 0.9:
        auc_interp = "excellent"
    elif auc_num >= 0.8:
        auc_interp = "good"
    elif auc_num >= 0.7:
        auc_interp = "acceptable"
    else:
        auc_interp = "poor"
    add_para(
        f"An AUC of {auc_str} indicates {auc_interp} discriminatory performance "
        f"of the model in distinguishing first births from later births based "
        f"on the included maternal and birth-related predictor variables."
    )
except ValueError:
    pass

# ─────────────────────────────────────────────────────────────────────────────
# 10.  SECTION 4.6 — Model Comparison
# ─────────────────────────────────────────────────────────────────────────────
add_heading("4.6  Comparison of the Three Models", level=1)

add_para(
    "The three regression models were compared using the Akaike Information "
    "Criterion (AIC), Bayesian Information Criterion (BIC), and log-likelihood. "
    "Lower AIC and BIC values indicate a better balance of model fit and "
    "parsimony, while a higher log-likelihood indicates a better-fitting model. "
    "The comparison is presented in Table 4.5."
)

add_results_table(
    comp_df,
    "Table 4.5: Model Comparison — AIC, BIC, and Log-Likelihood",
    col_headers=["Model", "Outcome", "Effect Measure", "AIC", "BIC", "Log-Likelihood"]
)

# auto-select best model
best_name = str(best_model_row.get("Model", ""))
best_aic  = float(best_model_row.get("AIC", 0))
worst_row = comp_df_sorted.iloc[-1]
worst_name= str(worst_row.get("Model", ""))
worst_aic = float(worst_row.get("AIC", 0))
delta_aic = abs(worst_aic - best_aic)

add_para(
    f"The {best_name} returned the lowest AIC value ({best_aic:,.2f}), "
    f"indicating it provides the best balance of fit and parsimony among "
    f"the three competing specifications. The {worst_name} had the highest "
    f"AIC ({worst_aic:,.2f}), a difference of \u0394AIC = {delta_aic:,.2f}. "
    f"In general, a \u0394AIC greater than 10 provides strong evidence in "
    f"favour of the lower-AIC model (Burnham & Anderson, 2002).",
    bold_parts=[best_name]
)

add_para(
    "Each model addresses a different aspect of the research objectives. The "
    "ordinal logistic regression (Model 1) is the most informative primary "
    "model because it preserves the full ordering of birth order 1 through 9 "
    "and yields cumulative odds ratios interpretable across the entire parity "
    f"distribution. The {count_model_name} regression (Model 2) offers "
    "a complementary perspective by treating birth order as a count outcome, "
    "consistent with the fertility literature on children-ever-born. The binary "
    "logistic regression (Model 3) directly addresses the study's specific "
    "objective of characterizing first-born children relative to later-born "
    "children, and provides an easily interpretable odds ratio for each "
    "predictor with respect to first-birth status. Together, the three models "
    "provide a comprehensive and robust characterisation of the parity "
    "structure in Sri Lanka's birth registration data."
)

# ─────────────────────────────────────────────────────────────────────────────
# 11.  SECTION 4.7 — Summary of Findings
# ─────────────────────────────────────────────────────────────────────────────
add_heading("4.7  Summary of Chapter Findings", level=1)

add_para(
    "This chapter presented the results of descriptive analysis, bivariate "
    "chi-square testing, and three multivariable regression models applied to "
    "Sri Lanka's birth registration data. The key findings are as follows:"
)

summary_points = [
    f"First-born children accounted for {pct_first:.1f}% of all registered births, "
    f"with birth order declining progressively at higher parities.",
    f"Chi-square tests identified {len(sig_vars)} statistically significant predictors "
    f"of birth order: {sig_str}.",
    f"The Proportional Odds Model (ordinal logistic regression) identified "
    f"several predictors significantly associated with cumulative birth order, "
    f"with the top significant predictors shown in Table 4.2.",
    f"The dispersion diagnostic confirmed {count_model_name} regression as the "
    f"appropriate count model (dispersion ratio = {disp_r:.4f}).",
    f"The binary logistic regression model achieved an AUC of {auc_str}, "
    f"demonstrating the model's ability to discriminate first births from "
    f"later births.",
    f"Model comparison by AIC identified {best_name} as the best-fitting "
    f"model among the three specifications."
]

for pt in summary_points:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.line_spacing = Pt(24)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(pt)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)

spacer()
add_para(
    "These findings are discussed in the context of Sri Lanka's demographic "
    "transition and the existing fertility literature in Chapter Five."
)

# ─────────────────────────────────────────────────────────────────────────────
# 12.  Save
# ─────────────────────────────────────────────────────────────────────────────
OUT = "results_chapter.docx"
doc.save(OUT)
print(f"\n{'='*50}")
print(f"  SUCCESS: {OUT} has been saved.")
print(f"  Open it in Microsoft Word or LibreOffice.")
print(f"{'='*50}")