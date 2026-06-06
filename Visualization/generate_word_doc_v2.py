from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def add_table(doc, data, headers):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
    for row_data in data:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)

def create_thesis_doc_v2():
    doc = Document()
    
    # Title
    title = doc.add_heading('CHAPTER FOUR', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph('RESULTS AND DISCUSSION')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True

    # 4.1 Data Cleaning
    doc.add_heading('4.1 Data Cleaning and Sample Derivation', level=1)
    doc.add_paragraph(
        "This section documents the full exclusion log arising from the data preparation procedures. "
        "The process involved outlier validation for maternal age and birth weight, followed by listwise deletion of records with missing values."
    )
    
    # Table 4.1: Exclusion Log
    doc.add_heading('Table 4.1: Structured Exclusion Log — Birth Registration Dataset (2000-2020)', level=2)
    exclusion_data = [
        ["1", "Maternal Age", "<10 or >60 (Impossible)", "6"],
        ["2", "Birth Weight", "<300g or >6000g (Impossible)", "0"],
        ["3", "Missing Values", "Listwise Deletion (Key Variables)", "0"],
        ["", "TOTAL RECORDS", "Final Analytical Sample", "1,600,173"]
    ]
    add_table(doc, exclusion_data, ["Step", "Variable", "Criterion", "n Removed"])

    # 4.2 Descriptive Statistics
    doc.add_heading('4.2 Descriptive Statistics', level=1)
    doc.add_heading('4.2.1 Distribution of Birth Order (Parity)', level=2)
    doc.add_paragraph(
        "Table 4.3 presents the frequency distribution of birth order across the analytical sample. "
        "The distribution is markedly right-skewed, reflecting the dominance of low-parity reproductive behaviour."
    )
    
    # Table 4.3: Parity Distribution
    parity_data = [
        ["First", "703,720", "43.98%", "45.39%"],
        ["Second", "541,074", "33.81%", "83.54%"],
        ["Third", "254,296", "15.89%", "100.01%"],
        ["Fourth", "68,953", "4.31%", "49.70%"],
        ["Fifth", "21,445", "1.34%", "1.41%"],
        ["6th and above", "10,685", "0.68%", "-"]
    ]
    add_table(doc, parity_data, ["Birth Order", "Frequency (n)", "Percentage (%)", "Cumulative (%)"])

    # Visuals for Parity
    img_p = 'Chapter_4_Results/parity_distribution_by_year.png'
    if os.path.exists(img_p):
        doc.add_picture(img_p, width=Inches(5))
        doc.add_paragraph('Figure 4.1: Yearly Parity Distribution Trends').alignment = 1

    # 4.3 Bivariate Analysis
    doc.add_heading('4.3 Bivariate Analysis', level=1)
    doc.add_paragraph("Chi-square tests were conducted to identify associations between parity and sociodemographic predictors.")
    
    # Table 4.4: Chi-Square Results
    chi_data = [
        ["Gender", "15.49", "0.0503", "8", "0.002"],
        ["Hospital or Not", "13,856.2", "0.0000", "16", "0.066"],
        ["Marital Status", "2,145.2", "0.0000", "8", "0.037"],
        ["Race of Mother", "46,190.4", "0.0000", "88", "0.060"]
    ]
    add_table(doc, chi_data, ["Variable", "Chi-Square", "p-value", "df", "Cramér's V"])

    # 4.4 Regression Results
    doc.add_heading('4.4 Regression Results', level=1)
    
    # Model 1: Ordinal
    doc.add_heading('4.4.1 Model 1: Ordinal Logistic Regression', level=2)
    doc.add_paragraph("Maternal age emerges as the strongest predictor of parity in the Proportional Odds Model.")
    ordinal_data = [
        ["Age < 20", "0.06", "[Ref: 25-29]"],
        ["Age 30-34", "2.54", ""],
        ["Age 35+", "5.49", ""],
        ["Moor Race", "2.48", "[Ref: Sinhalese]"],
        ["Non-Hospital", "2.10", "[Ref: Hospital]"]
    ]
    add_table(doc, ordinal_data, ["Predictor", "Cumulative OR", "Note"])

    # Model 2: Count
    doc.add_heading('4.4.2 Model 2: Poisson Regression', level=2)
    poisson_data = [
        ["Age 35+", "1.52", "Multiplicative effect"],
        ["Moor Race", "1.26", ""],
        ["Non-Hospital", "1.20", ""]
    ]
    add_table(doc, poisson_data, ["Predictor", "IRR", "Note"])

    # Model 3: Binary
    doc.add_heading('4.4.3 Model 3: Binary Logistic (First Birth)', level=2)
    binary_data = [
        ["Age < 20", "14.87", "Very high odds of 1st birth"],
        ["Age 35+", "0.25", "Very low odds of 1st birth"],
        ["Hospital Delivery", "1.66", "[Calculated as 1/0.60]"]
    ]
    add_table(doc, binary_data, ["Predictor", "Odds Ratio", "Note"])

    # 4.5 Geospatial Atlas (Incorporating the grid maps)
    doc.add_heading('4.5 Geospatial Analysis of Parity Transition', level=1)
    doc.add_paragraph("The following maps document the regional diffusion of low-parity norms.")
    
    img_g1 = 'Chapter_4_Results/parity_grid_part1.png'
    if os.path.exists(img_g1):
        doc.add_picture(img_g1, width=Inches(6))
        doc.add_paragraph('Figure 4.2: Spatial Atlas (1st-3rd Parity)').alignment = 1

    img_g2 = 'Chapter_4_Results/parity_grid_part2.png'
    if os.path.exists(img_g2):
        doc.add_picture(img_g2, width=Inches(6))
        doc.add_paragraph('Figure 4.3: Spatial Atlas (4th-6th+ Parity)').alignment = 1

    # 4.6 Discussion
    doc.add_heading('4.7 Discussion', level=1)
    doc.add_paragraph(
        "The results document a society in the final stages of demographic transition. "
        "The decline in TFR to 1.83 by 2020 is driven by a universal shift toward the first and second birth orders, "
        "regardless of ethnic background or geographic location. Maternal age remains the dominant biological driver, "
        "with delayed childbearing becoming the national standard."
    )

    doc.save('Chapter_4_Results/Chapter_4_Final_HighFidelity.docx')
    print("High-fidelity Word document created.")

if __name__ == "__main__":
    create_thesis_doc_v2()
