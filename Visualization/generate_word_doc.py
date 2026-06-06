from docx import Document
from docx.shared import Inches
import os

def create_thesis_chapter():
    doc = Document()
    
    # Title
    doc.add_heading('Chapter 4: Results and Discussion', 0)

    # 4.1 Introduction
    doc.add_heading('4.1 Introduction', level=1)
    doc.add_paragraph(
        "This chapter presents the findings of the study based on the analysis of 1,600,179 birth registration records "
        "obtained from the Sri Lankan Civil Registration and Vital Statistics (CRVS) system for the years 2000, 2005, 2010, 2015, and 2020. "
        "The analysis explores national fertility trends, the internal dynamics of birth parity, regional variations through geospatial mapping, "
        "and the associations between parity and maternal sociodemographic characteristics."
    )

    # 4.2 National Fertility Dynamics
    doc.add_heading('4.2 National Fertility Dynamics', level=1)
    doc.add_heading('4.2.1 Total Fertility Rate (TFR) and Age-Specific Fertility', level=2)
    doc.add_paragraph(
        "Sri Lanka has undergone a distinct demographic shift during the first two decades of the 21st century. "
        "As illustrated in Figure 4.1, the Total Fertility Rate (TFR) experienced an initial rise from 2.01 in 2000 to a peak of 2.27 in 2010. "
        "However, the subsequent decade documented a consistent decline, reaching 1.83 children per woman by 2020."
    )
    doc.add_paragraph(
        "This transition into sub-replacement fertility is further elucidated by the Age-Specific Fertility Rates (ASFR). "
        "The peak childbearing age consistently remains in the 25–29 and 30–34 cohorts. However, the 2020 curve shows a significant "
        "downward shift across all age groups compared to the 2010 peak, signaling a universal reduction in fertility rather than a mere delay in childbearing."
    )
    
    img_path = 'Chapter_4_Results/fertility_analysis_sri_lanka.png'
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6))
        doc.add_paragraph('Figure 4.1: Total Fertility Rate and Age-Specific Fertility Rates (2000-2020)').alignment = 1

    # 4.3 Parity Distribution and Dynamics
    doc.add_heading('4.3 Parity Distribution and Dynamics', level=1)
    doc.add_heading('4.3.1 Composition of Family Size', level=2)
    doc.add_paragraph(
        "The core of this study focuses on birth parity (birth order). Figure 4.2 shows that first and second births dominate the "
        "Sri Lankan demographic landscape, consistently accounting for over 75% of all registrations. There is a visible \"thinning\" "
        "of higher-order parity segments (3rd, 4th, 5th+) between 2000 and 2020."
    )
    
    img_path = 'Chapter_4_Results/parity_distribution_by_year.png'
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        doc.add_paragraph('Figure 4.2: Distribution of Birth Parity by Year').alignment = 1

    doc.add_heading('4.3.2 Statistical Skewness of Parity', level=2)
    doc.add_paragraph(
        "The distribution of parity is heavily right-skewed, as shown in Figure 4.3. The skewness coefficient has declined from "
        "1.682 in 2000 to 0.939 in 2020. This trend indicates that while family sizes are concentrating around 1-2 children, "
        "the \"long tail\" of very large families (6+) has nearly disappeared, leading to a more homogenized national reproductive profile."
    )
    
    img_path = 'Chapter_4_Results/parity_skewness_by_year.png'
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(5.5))
        doc.add_paragraph('Figure 4.3: Analysis of Parity Skewness and Frequency Polygons').alignment = 1

    # 4.4 Geospatial Analysis of Birth Parity
    doc.add_heading('4.4 Geospatial Analysis of Birth Parity', level=1)
    doc.add_heading('4.4.1 Regional Registration Volume', level=2)
    doc.add_paragraph(
        "Geospatial mapping (Figure 4.4) reveals a persistent concentration of birth registrations in the Western Province, "
        "driven by high population density and the presence of tertiary healthcare hubs in Colombo and Gampaha."
    )
    
    img_path = 'Chapter_4_Results/yearly_birth_maps_sri_lanka.png'
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6.5))
        doc.add_paragraph('Figure 4.4: Yearly Birth Maps of Sri Lanka (Absolute Counts)').alignment = 1

    doc.add_heading('4.4.2 The Spatial Atlas of Parity Transition', level=2)
    doc.add_paragraph(
        "The spatial diffusion of low-parity norms is documented in Figures 4.5 and 4.6. The 2020 maps show that the "
        "\"two-child standard\" has expanded from the urban Western core to nearly all districts. Regional pockets of high-order births, "
        "which were prominent in the Eastern and Northern provinces in 2000, have largely vanished by 2020."
    )
    
    img_path1 = 'Chapter_4_Results/parity_grid_part1.png'
    if os.path.exists(img_path1):
        doc.add_picture(img_path1, width=Inches(6.5))
        doc.add_paragraph('Figure 4.5: Spatial Atlas of Parity Transitions (1st to 3rd Births)').alignment = 1
        
    img_path2 = 'Chapter_4_Results/parity_grid_part2.png'
    if os.path.exists(img_path2):
        doc.add_picture(img_path2, width=Inches(6.5))
        doc.add_paragraph('Figure 4.6: Spatial Atlas of Parity Transitions (4th to 6th+ Births)').alignment = 1

    # 4.5 Sociodemographic Associations
    doc.add_heading('4.5 Sociodemographic Associations', level=1)
    doc.add_heading('4.5.1 Parity and Maternal Ethnicity', level=2)
    doc.add_paragraph(
        "The study explored how parity varies across ethnic groups (Figure 4.7). While the Sri Lankan Moor community "
        "historically maintained a higher proportion of 3rd and 4th parity births, the longitudinal data shows a convergence. "
        "By 2020, all major ethnic groups show a primary concentration in the first two birth orders."
    )
    
    img_path = 'Chapter_4_Results/parity_by_race_distribution_with_counts.png'
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6))
        doc.add_paragraph('Figure 4.7: Birth Parity Distribution by Mother\'s Race').alignment = 1

    doc.add_heading('4.5.2 Parity and Gender', level=2)
    doc.add_paragraph(
        "As shown in Figure 4.8, the sex ratio at birth remains remarkably stable across all parity levels. The biological "
        "balance between Male and Female births does not appear to be influenced by family size or birth order."
    )
    
    img_path = 'Chapter_4_Results/parity_gender_distribution_by_year.png'
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6))
        doc.add_paragraph('Figure 4.8: Gender Distribution by Birth Order').alignment = 1

    # 4.6 Maternal Age Analysis
    doc.add_heading('4.6 Maternal Age Analysis', level=1)
    doc.add_heading('4.6.1 Delayed Childbearing Trends', level=2)
    doc.add_paragraph(
        "There is a profound correlation between maternal age and the demographic transition. Figure 4.9 demonstrates "
        "that the median age for first-time mothers has risen from 24.0 (2000) to 26.0 (2020). This upward shift is universal "
        "across all parity levels, indicating a nationwide trend toward later marriage and delayed childbearing."
    )
    
    img_path = 'Chapter_4_Results/maternal_age_by_parity_boxplot.png'
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6))
        doc.add_paragraph('Figure 4.9: Maternal Age Distribution by Parity').alignment = 1

    doc.add_heading('4.6.2 Distributional Shifts and Outlier Justification', level=2)
    doc.add_paragraph(
        "The density curves in Figure 4.10 visually confirm the rightward shift of the peak childbearing age. Furthermore, "
        "Figure 4.11 provides the methodological justification for our data cleaning approach. By retaining biologically plausible "
        "cases (ages 10-14 and 50-60), the study ensures that rare but clinically valid reproductive events are captured."
    )
    
    img_path10 = 'Chapter_4_Results/maternal_age_distribution_curves.png'
    if os.path.exists(img_path10):
        doc.add_picture(img_path10, width=Inches(6))
        doc.add_paragraph('Figure 4.10: Maternal Age Density Curves and Frequency Polygons').alignment = 1
        
    img_path11 = 'Chapter_4_Results/maternal_age_outlier_comparison.png'
    if os.path.exists(img_path11):
        doc.add_picture(img_path11, width=Inches(6))
        doc.add_paragraph('Figure 4.11: Maternal Age Outlier Analysis Comparison').alignment = 1

    # 4.7 Conclusion
    doc.add_heading('4.7 Conclusion of Findings', level=1)
    doc.add_paragraph(
        "In summary, Chapter 4 documents a society in the final stages of demographic transition. Sri Lanka has moved "
        "from a peak fertility period in 2010 to a sub-replacement regime by 2020. This transition is characterized by a "
        "geographic homogenization of parity, a universal delay in maternal age, and a steady convergence of reproductive behavior."
    )

    doc.save('Chapter_4_Results/Chapter_4_Results_and_Discussion.docx')
    print("Word document created: Chapter_4_Results/Chapter_4_Results_and_Discussion.docx")

if __name__ == "__main__":
    create_thesis_chapter()
