import pandas as pd
import numpy as np
import os
import matplotlib.pyplot as plt
import seaborn as sns
from docx import Document
from docx.shared import Inches

# Set style for visualizations
sns.set_theme(style="whitegrid")
plt.rcParams['figure.figsize'] = (10, 6)

def generate_visualizations(years):
    print("Generating visualizations...")
    
    # 1. Yearly Total Births Trend
    summary_path = 'analysis_merged/yearly_summary_trends.csv'
    if os.path.exists(summary_path):
        ys = pd.read_csv(summary_path)
        plt.figure()
        sns.lineplot(data=ys, x='Year', y='TotalBirths', marker='o', linewidth=2.5)
        plt.title('Trend of Total Births in Sri Lanka (2000-2020)', fontsize=14)
        plt.xlabel('Year')
        plt.ylabel('Number of Births')
        plt.xticks(years)
        plt.savefig('analysis_merged/yearly_birth_trend.png', dpi=300, bbox_inches='tight')
        plt.close()

        # Average Mother Age Trend
        plt.figure()
        sns.lineplot(data=ys, x='Year', y='AvgMotherAge', marker='s', color='orange', linewidth=2.5)
        plt.title('Trend of Average Maternal Age (2000-2020)', fontsize=14)
        plt.xlabel('Year')
        plt.ylabel('Average Age')
        plt.xticks(years)
        plt.savefig('analysis_merged/maternal_age_trend.png', dpi=300, bbox_inches='tight')
        plt.close()

    # 2. District Trends (Top 10 Districts)
    trends_path = 'analysis_merged/district_trends_over_years.csv'
    if os.path.exists(trends_path):
        dt = pd.read_csv(trends_path)
        dt['Total_All_Years'] = dt[[f'Births_{y}' for y in years]].sum(axis=1)
        top_10 = dt.sort_values('Total_All_Years', ascending=False).head(10)
        
        # Melt for plotting
        melted = top_10.melt(id_vars='District', value_vars=[f'Births_{y}' for y in years], 
                             var_name='Year', value_name='Births')
        melted['Year'] = melted['Year'].str.extract('(\d+)').astype(int)
        
        plt.figure()
        sns.lineplot(data=melted, x='Year', y='Births', hue='District', marker='o')
        plt.title('Birth Trends in Top 10 Districts (2000-2020)', fontsize=14)
        plt.xticks(years)
        plt.legend(bbox_to_anchor=(1.05, 1), loc='upper left')
        plt.savefig('analysis_merged/district_trends_top10.png', dpi=300, bbox_inches='tight')
        plt.close()

    # 3. Year-specific Plots
    for year in years:
        race_path = f'analysis_{year}/freq_Race_of_Mother_{year}.csv'
        if os.path.exists(race_path):
            race_df = pd.read_csv(race_path)
            # The previous script saved with 'Race_of_Mother' as the first column (index)
            # or it might be 'Unnamed: 0' depending on how read_csv handled it.
            # Let's handle both.
            if 'Race_of_Mother' in race_df.columns:
                label_col = 'Race_of_Mother'
            elif 'Unnamed: 0' in race_df.columns:
                label_col = 'Unnamed: 0'
            else:
                label_col = race_df.columns[0]
                
            plt.figure()
            plt.pie(race_df['Count'], labels=race_df[label_col], autopct='%1.1f%%', startangle=140)
            plt.title(f'Distribution of Maternal Race ({year})', fontsize=14)
            plt.savefig(f'analysis_{year}/race_distribution_{year}.png', dpi=300, bbox_inches='tight')
            plt.close()

def update_word_doc(years):
    print("Updating Word document with visualizations...")
    doc = Document()
    doc.add_heading('Professional Descriptive Analysis & Visualization of Birth Data', 0)
    
    doc.add_heading('1. Longitudinal Trends (2000-2020)', level=1)
    
    # Total Births
    doc.add_heading('1.1 National Birth Volume Trends', level=2)
    doc.add_paragraph('The figure below illustrates the fluctuation in total birth registrations in Sri Lanka over the two-decade study period.')
    if os.path.exists('analysis_merged/yearly_birth_trend.png'):
        doc.add_picture('analysis_merged/yearly_birth_trend.png', width=Inches(6))
    
    # Maternal Age
    doc.add_heading('1.2 Demographic Shifts: Maternal Age', level=2)
    doc.add_paragraph('As shown in the following trend line, the average age of mothers at birth has undergone significant changes, reflecting broader socio-economic shifts.')
    if os.path.exists('analysis_merged/maternal_age_trend.png'):
        doc.add_picture('analysis_merged/maternal_age_trend.png', width=Inches(6))

    doc.add_heading('2. Regional Analysis', level=1)
    doc.add_paragraph('Spatial variation is a key factor in birth patterns across Sri Lanka.')
    if os.path.exists('analysis_merged/district_trends_top10.png'):
        doc.add_picture('analysis_merged/district_trends_top10.png', width=Inches(6))

    doc.add_heading('3. Detailed Yearly Demographic Snapshots', level=1)
    for year in years:
        doc.add_heading(f'3.{years.index(year)+1} Snapshot: Year {year}', level=2)
        race_img = f'analysis_{year}/race_distribution_{year}.png'
        # Generate race distribution for each year if not exists (minimal version)
        if os.path.exists(race_img):
            doc.add_picture(race_img, width=Inches(4))
            
    doc.add_heading('4. Conclusion for Thesis', level=1)
    doc.add_paragraph('The integration of these visualizations provides clear empirical evidence of demographic transitions. These trends support the thesis argument regarding regional healthcare demand and spatial accessibility.')

    doc.save('Analysis_Report_with_Visuals.docx')

def main():
    years = [2000, 2005, 2010, 2015, 2020]
    # We assume previous script already ran and generated CSVs
    generate_visualizations(years)
    update_word_doc(years)
    print("Visualizations and updated report complete.")

if __name__ == "__main__":
    main()
